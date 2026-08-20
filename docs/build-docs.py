"""Convert the component .docx documentation into the site's module pages.

Faithful conversion, not paraphrase: headings keep their level, lists stay lists,
tables stay tables, and block quotes keep their emphasis. The prose is Perry's.

    py site/build_docs.py

Re-run whenever documentation/*.docx changes. Output overwrites site/modules/*.html.
"""
import html
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "documentation")
OUT = os.path.join(ROOT, "docs", "components")

# source .docx -> (output slug, nav label, one-line standfirst)
PAGES = [
    ("1. DataView_Database.docx", "dataview", "DataView",
     "The database. A PPDM 3.9 derivative, reshaped for messy real-world data."),
    ("4. File_Catalog.docx", "file-catalog", "File Catalog",
     "The unstructured half. Crawl a share, catalogue what is there, extract what is useful."),
    ("2.Data_Assistant.docx", "data-assistant", "Data Assistant",
     "Bulk tabular loading, where the mapping is proposed and you say where it is wrong."),
    ("3.Document_Assistant.docx", "document-assistant", "Document Assistant",
     "Identify the table, not the document. Shape recognition that learns as it reads."),
    ("5. Mapping.docx", "mapping", "Mapping",
     "The exploration surface. For the questions you cannot phrase as a query."),
]


def esc(t):
    return html.escape(t, quote=False)


def inline(par):
    """Keep bold/italic runs; everything else is plain text."""
    out = []
    for r in par.runs:
        t = esc(r.text)
        if not t.strip():
            out.append(t)
            continue
        if r.bold:
            t = f"<strong>{t}</strong>"
        if r.italic:
            t = f"<em>{t}</em>"
        out.append(t)
    s = "".join(out)
    # `code` spans written literally in the source
    s = re.sub(r"`([^`]+)`", r"<code>\1</code>", s)
    return s


def render_table(tbl):
    rows = tbl.rows
    if not rows:
        return ""
    head = rows[0]
    body = rows[1:]
    th = "".join(f"<th>{esc(c.text.strip())}</th>" for c in head.cells)
    trs = []
    for r in body:
        tds = "".join(f"<td>{esc(c.text.strip())}</td>" for c in r.cells)
        trs.append(f"<tr>{tds}</tr>")
    return ('<div class="tablewrap"><table>\n<thead><tr>' + th + "</tr></thead>\n<tbody>\n"
            + "\n".join(trs) + "\n</tbody></table></div>")


def convert(path):
    """Walk the document body in order so tables land where they belong."""
    doc = Document(path)
    body = doc.element.body
    parts, list_open = [], None
    para_by_el = {p._p: p for p in doc.paragraphs}
    tbl_by_el = {t._tbl: t for t in doc.tables}

    def close_list():
        nonlocal list_open
        if list_open:
            parts.append(f"</{list_open}>")
            list_open = None

    for child in body.iterchildren():
        if child in tbl_by_el:
            close_list()
            parts.append(render_table(tbl_by_el[child]))
            continue
        p = para_by_el.get(child)
        if p is None:
            continue
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower()
        content = inline(p)

        if "list" in style:
            tag = "ol" if "number" in style else "ul"
            if list_open != tag:
                close_list()
                parts.append(f"<{tag}>")
                list_open = tag
            parts.append(f"<li>{content}</li>")
            continue
        close_list()

        if style.startswith("heading 1") or style == "title":
            parts.append(f"<h2>{content}</h2>")
        elif style.startswith("heading 2"):
            parts.append(f"<h2>{content}</h2>")
        elif style.startswith("heading 3"):
            parts.append(f"<h3>{content}</h3>")
        elif "quote" in style or "block text" in style:
            parts.append(f'<div class="quote">{content}</div>')
        else:
            parts.append(f"<p>{content}</p>")
    close_list()
    return "\n".join(parts)


NAV = [("../components.html", "All components"),
       ("dataview.html", "DataView"),
       ("file-catalog.html", "File Catalog"),
       ("data-assistant.html", "Data Assistant"),
       ("document-assistant.html", "Document Assistant"),
       ("mapping.html", "Mapping"),
       ("../index.html#films", "Walkthroughs")]


def page(slug, label, standfirst, inner):
    nav = "\n".join(
        f'      <a href="{h}"{" aria-current=\"page\" " if h.endswith(slug + ".html") else ""}>{t}</a>'
        for h, t in NAV)
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="color-scheme" content="dark light">
<title>{esc(label)} — Data Wrangler</title>
<meta name="description" content="{esc(standfirst)}">
<link rel="stylesheet" href="../assets/site.css">
</head>
<body>

<header class="topbar">
  <div class="wrap">
    <a class="brand" href="../index.html">DATA<span>·</span>WRANGLER</a>
    <nav class="topnav">
{nav}
    </nav>
  </div>
</header>

<main class="doc">
  <div class="wrap">
    <a class="backlink" href="../components.html">&larr; All components</a>
    <div class="eyebrow">Component documentation</div>
    <h1>{esc(label)}</h1>
    <p class="standfirst">{esc(standfirst)}</p>
{inner}
  </div>
</main>

<footer>
  <div class="wrap">
    <div>
      <strong style="font-family:var(--mono);color:var(--ink)">Data Wrangler Solutions LLC</strong><br>
      Petroleum data management · SQL Server · Oracle · Snowflake
    </div>
    <div><a href="../index.html">Overview</a> &nbsp;·&nbsp; <a href="../components.html">Components</a></div>
  </div>
</footer>

</body>
</html>
"""


def main():
    os.makedirs(OUT, exist_ok=True)
    for fname, slug, label, standfirst in PAGES:
        src = os.path.join(SRC, fname)
        if not os.path.exists(src):
            print(f"  MISSING {fname}")
            continue
        inner = convert(src)
        dst = os.path.join(OUT, slug + ".html")
        with open(dst, "w", encoding="utf-8") as fh:
            fh.write(page(slug, label, standfirst, inner))
        words = len(re.findall(r"\S+", re.sub(r"<[^>]+>", " ", inner)))
        print(f"  {slug + '.html':<26} {words:>5} words   <- {fname}")


if __name__ == "__main__":
    main()
