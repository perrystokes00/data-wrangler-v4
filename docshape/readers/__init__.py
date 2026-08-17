"""
docshape.readers
================
Getting tables and metadata out of files. Format-specific, domain-neutral.

    from docshape.readers import read_tables, SUPPORTED

    tables = read_tables("survey.pdf")     # {name: [row dicts]}

Formats split into two kinds. TABLE-BEARING documents — PDF, DOCX, XLSX — give
back named tables for the recogniser to identify. STRUCTURED files — LAS,
SEG-Y — have no tables at all and are parsed natively into their own fixed
shape, because a curve list or a binary header is not something a column
matcher can help with.

Adding a format means a module here plus an entry in DISPATCH. Nothing in the
engine or the packs changes.
"""
from __future__ import annotations

import os

#: extensions whose content is TABLES, to be identified by the recogniser
TABLE_EXTS = {".pdf", ".docx", ".xlsx"}
#: extensions parsed natively into a fixed shape
NATIVE_EXTS = {".las", ".segy", ".sgy"}
SUPPORTED = TABLE_EXTS | NATIVE_EXTS


def read_tables(path):
    """{table_name: [row dicts]} for a table-bearing document."""
    from docshape.readers.tables import raw_tables
    return raw_tables(path)


def read_native(path):
    """(kind, payload) for a natively-parsed file, or (None, None).

    kind is 'las' or 'segy'; payload is whatever that reader returns.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".las":
        from docshape.readers.las import parse_las
        return "las", parse_las(path)
    if ext in (".segy", ".sgy"):
        from docshape.readers.segy import parse_segy
        return "segy", parse_segy(path)
    return None, None


def collect(target, exts=None):
    """Every supported file under a directory, or the file itself."""
    allowed = exts or SUPPORTED
    if os.path.isfile(target):
        return [target]
    out = []
    for root, _dirs, files in os.walk(target):
        for f in sorted(files):
            if os.path.splitext(f)[1].lower() in allowed:
                out.append(os.path.join(root, f))
    return out

def read_text(path, max_pages=2, max_chars=8000):
    """First page(s) of text, for identity search. '' when unavailable.

    Deliberately shallow: a pack looks for a well API or a matter number in the
    header block, and reading a 200-page report to find it is waste. Two pages
    covers the letterhead and the summary table on every real document seen so
    far.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import pdfplumber
            out = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages[:max_pages]:
                    out.append(page.extract_text() or "")
            return "\n".join(out)[:max_chars]
        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs[:200]]
            for t in doc.tables[:5]:
                for row in t.rows[:20]:
                    parts.append(" ".join(c.text for c in row.cells))
            return "\n".join(parts)[:max_chars]
        if ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets[:2]:
                for row in ws.iter_rows(max_row=20, values_only=True):
                    parts.append(" ".join("" if v is None else str(v)
                                          for v in row))
            return "\n".join(parts)[:max_chars]
    except Exception:
        pass
    return ""
