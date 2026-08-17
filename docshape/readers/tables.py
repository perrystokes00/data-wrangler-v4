"""
docshape.readers.tables
=======================
Get tables out of a document. Format-specific, domain-neutral: this module
knows about PDFs and spreadsheets, never about wells or contracts.

TWO PDF STRATEGIES, AND THE SECOND ONE MATTERS
----------------------------------------------
pdfplumber's default table detection follows RULED LINES. Generated documents
(ReportLab, Word exports) draw them, so it works. Real vendor reports often
don't — a Baker Hughes survey has lines=0 and its whole station table collapses
into a single cell.

So when the ruled strategy finds no real grid, fall back to WORD POSITIONS:
locate the data rows (mostly numeric), take their x-centres as the columns, and
assign the header words above to the nearest one. That recovers multi-word
headers ("Meas Depth", "Dog Leg Sev") and the units line that often sits
between the names and the data.

The fallback only fires when the ruled strategy found nothing, so documents
that were already read correctly are untouched.
"""
from __future__ import annotations

import os
import re
import statistics
from html.parser import HTMLParser

def docx_tables(path):
    from docx import Document
    out = {}
    for i, tbl in enumerate(Document(path).tables, start=1):
        if not tbl.rows:
            continue
        hdr = [c.text.strip() or f"col{j}"
               for j, c in enumerate(tbl.rows[0].cells, start=1)]
        rows = [dict(zip(hdr, [c.text.strip() for c in r.cells]))
                for r in tbl.rows[1:]]
        if rows:
            out[f"table_{i:02d}"] = rows
    return out


def xlsx_sheets(path):
    from openpyxl import load_workbook
    out = {}
    for ws in load_workbook(path, read_only=True, data_only=True).worksheets:
        it = ws.iter_rows(values_only=True)
        try:
            hdr = next(it)
        except StopIteration:
            continue
        hdr = [str(h).strip() if h is not None else f"col{j}"
               for j, h in enumerate(hdr, start=1)]
        rows = [dict(zip(hdr, ["" if v is None else v for v in vals]))
                for vals in it]
        if rows:
            out[f"sheet_{ws.title}"] = rows
    return out

_NUM_RE = re.compile(r"[+\-]?[\d,]*\.?\d+")


def _text_lines(page, ytol=2.5):
    """Words grouped into visual lines, left to right."""
    rows = {}
    for w in page.extract_words(use_text_flow=False, keep_blank_chars=False):
        rows.setdefault(round(w["top"] / ytol), []).append(w)
    return [sorted(v, key=lambda w: w["x0"]) for _k, v in sorted(rows.items())]


def _is_number(tok):
    return bool(_NUM_RE.fullmatch(str(tok).replace(",", "")))


def _centre(w):
    return (w["x0"] + w["x1"]) / 2.0


def _respace(cell):
    """Put the spaces back into a cell typeset without space glyphs.

    Some PDFs draw "FinalPressure(psi)" with no space characters at all —
    the visual gap is kerning, not a glyph — so the cell arrives as one
    token and no alias can match it. Two seams are safe to split on:
    a lower->UPPER turn ("FinalPressure", "FluidType", "ftMD") and a ')'
    or digit followed by a letter.

    Deliberately conservative. It never splits UPPER->UPPER, so "H2S",
    "BS&W", "GOR" and "API" survive; and it only runs on a cell that has
    no spaces of its own, so a header that was typeset properly is left
    exactly alone.
    """
    if " " in cell or len(cell) < 6:
        return cell
    out, depth = [], 0
    for i, ch in enumerate(cell):
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        if i and ch.isupper() and (cell[i - 1].islower() or cell[i - 1] == ")"):
            # INSIDE PARENTHESES THE RULE IS STRICTER, because that is
            # where UNIT SYMBOLS live: "mD" is millidarcies and "cP" is
            # centipoise, and splitting them into "m D" and "c P"
            # destroys the unit while pretending to fix the header. So
            # inside parens split only when the run of capitals is two or
            # more — "ftMD" becomes "ft MD", "mD" is left alone. Outside
            # parens the text is CamelCase words ("FinalPressure",
            # "FluidType") and any lower->UPPER turn is a word boundary.
            run = 0
            while i + run < len(cell) and cell[i + run].isupper():
                run += 1
            if depth == 0 or run >= 2:
                out.append(" ")
        out.append(ch)
    return "".join(out)


def _char_header(page, line_words, centres, ytol=2.5):
    """Header cells rebuilt from CHARACTERS instead of extracted words.

    Real vendor PDFs sometimes carry a header line with no reliable space
    glyphs, so pdfplumber's word segmentation glues fragments across column
    boundaries — "Pre-Test (psi)Fi", "nal Pressure (ps" — and no word-level
    assignment can unscramble text that arrived pre-broken. Characters
    can't be wrongly glued: assign each char to the nearest column centre,
    sort by x, and insert a space where the gap between neighbours says the
    source had one. Returns None when the page exposes no chars, and the
    caller falls back to word assignment.
    """
    chars = getattr(page, "chars", None)
    if not chars:
        return None
    tops = {round(w["top"] / ytol) for w in line_words}
    line = [c for c in chars
            if round(c["top"] / ytol) in tops and str(c.get("text", "")).strip()]
    if not line:
        return None
    ncol = len(centres)
    # STREAM ORDER, NOT X ORDER. Sorting characters by x0 looks obviously
    # right and is the thing that corrupts an overlapping header. In one
    # real RFT the "G" of Gradient starts at x0=368.02 while the "i)" of
    # the previous column's "(psi)" runs to 373.72 — the cells overlap by
    # five points — so an x-sort interleaves them into "(psGi)radient",
    # and no downstream repair can recover the words. The PDF's own
    # content stream has them in the right order; it is only the
    # POSITIONS that overlap.
    #
    # So group in stream order, and treat a NEGATIVE gap — a character
    # that begins before its predecessor ended — as a column boundary.
    # Nothing inside a word ever moves backwards; text that does is text
    # belonging to a different cell. Groups are sorted by x afterwards, so
    # a well-behaved PDF (where stream order already equals x order)
    # behaves exactly as before.
    widths = sorted(max(c["x1"] - c["x0"], 0.1) for c in line)
    med_w = widths[len(widths) // 2]
    groups, cur = [], [line[0]]
    for c in line[1:]:
        gap = c["x0"] - cur[-1]["x1"]
        if gap > 0.6 * med_w or gap < -0.15 * med_w:
            groups.append(cur)
            cur = [c]
        else:
            cur.append(c)
    groups.append(cur)
    groups.sort(key=lambda g: min(ch["x0"] for ch in g))

    # WHEN THE GROUPING IS UNAMBIGUOUS, TRUST IT. If gap detection alone
    # produced exactly one group per column, the header is already
    # segmented and the zone machinery below can only damage it: zones are
    # derived from the DATA rows' centres, and on an overlapping header a
    # boundary can fall inside a group, re-gluing "Pre-Test (psi)" to
    # "Final" and cutting "Pressure (psi)" in half. The zone path exists
    # for headers with NO usable gaps; this one has them.
    if len(groups) == ncol:
        out = [_respace(" ".join("".join(ch["text"] for ch in g).split()))
               for g in groups]
        if all(o.strip() for o in out):
            return out
    # A section TITLE ("Formation Pressure Measurements") has far fewer
    # visual cells than the table has columns; a header or units line has
    # about as many. Judged on CHAR groups — word segmentation is exactly
    # what can't be trusted here. Returning None hands the line to the
    # word-assign path, whose filled-cells check then stops the climb.
    if len(groups) < max(2, int(ncol * 0.55)):
        return None
    bounds = [(centres[i] + centres[i + 1]) / 2.0 for i in range(ncol - 1)]

    def _zone(x):
        z = 0
        while z < len(bounds) and x > bounds[z]:
            z += 1
        return z

    # A GAPLESS run can still span several columns — some PDFs typeset
    # "Pre-Test (psi)Final Pressure (psi)Gradient (psi/ft)" with no space
    # glyphs at all. Cut such a group at each crossed boundary, SNAPPED to
    # the nearest natural name seam — ')' followed by a letter, or a
    # lower->UPPER case turn — within half a column; a hard positional cut
    # only when no seam is near.
    def _subdivide(g):
        z0, z1 = _zone(g[0]["x0"] + 0.5 * med_w), _zone(g[-1]["x1"] - 0.5 * med_w)
        if z0 == z1:
            return [g]
        seams = []
        for k in range(1, len(g)):
            a, b = str(g[k - 1]["text"]), str(g[k]["text"])
            if (a == ")" and b.isalpha()) or (a.islower() and b.isupper()):
                seams.append(k)
        parts, start = [], 0
        for z in range(z0, z1):
            bx = bounds[z]
            near = [k for k in seams
                    if start < k and abs(g[k]["x0"] - bx) <=
                    max(6 * med_w, (bounds[1] - bounds[0]) / 2 if len(bounds) > 1 else 6 * med_w)]
            if near:
                cut = min(near, key=lambda k: abs(g[k]["x0"] - bx))
            else:
                cand = [k for k in range(start + 1, len(g))
                        if g[k]["x0"] >= bx]
                cut = cand[0] if cand else len(g)
            if cut <= start or cut >= len(g):
                continue
            parts.append(g[start:cut])
            start = cut
        parts.append(g[start:])
        return [p for p in parts if p]

    flat = []
    for g in groups:
        flat.extend(_subdivide(g))
    cols = [[] for _ in range(ncol)]
    for g in flat:
        cols[_zone(g[0]["x0"] + 0.5 * med_w)].append(
            "".join(c["text"] for c in g))
    return [" ".join(g for g in col).strip() for col in cols]


def _gap_clusters(words):
    """One line's words -> visual CELLS, grouped at gaps.

    Word spacing is a fraction of a character width; a column gutter is
    several. Clustering first means "Eagle Ford A" is ONE cell — under
    word-counting it was three, which pushed those rows below the numeric
    threshold and out of the table, leaving their neighbours to be climbed
    into the header (the RFT stacked-data-header failure)."""
    if not words:
        return []
    ws = sorted(words, key=lambda w: w["x0"])
    cw = sorted(max(w["x1"] - w["x0"], 0.1) / max(len(str(w["text"])), 1)
                for w in ws)
    med = cw[len(cw) // 2]
    out, cur = [], [ws[0]]
    for w in ws[1:]:
        if w["x0"] - cur[-1]["x1"] > 1.4 * med:
            out.append(cur)
            cur = [w]
        else:
            cur.append(w)
    out.append(cur)
    return out


def _cell_text(cluster):
    return " ".join(w["text"] for w in cluster)


def _cluster_numfrac(clusters):
    if not clusters:
        return 0.0
    return sum(1 for c in clusters if _is_number(_cell_text(c))) \
        / len(clusters)


def _flatten_pair_grid(hdr, rows):
    """A label/value grid -> one wide record, or None if it isn't one.

    A document header is written as pairs across the page:

        Operator   Devon Energy   Well Name   JONES 15-5
        UWI / API  15009202380000 Field       El Dorado
        State      KS             County      15009

    Extracted as a table that is FOUR COLUMNS WIDE AND EIGHT ROWS DEEP,
    with no header of its own — so the recogniser, which identifies a
    table from its HEADER ROW, sees col1..col4 and matches nothing. The
    labels it needs are sitting in columns 0 and 2, all the way down.

    That is why an end-of-well report could be read perfectly, show its
    formation tops and its NPT table, and still be rejected downstream
    with "no detail rows": the UWI is in this block, the block was never
    identified, so no well was ever resolved.

    Flattening turns it into what the recogniser already understands —
    header = every label, one row = every value:

        ['Operator', 'Well Name', 'UWI / API', 'Field', 'State', ...]
        ['Devon Energy', 'JONES 15-5', '15009202380000', ...]

    STRUCTURAL ONLY. It requires placeholder headers (nothing was found
    above the block), an even column count, and label columns that look
    like labels — short, non-numeric, and not repeating. Whether
    "UWI / API" means anything is the vocabulary's business, not the
    reader's.
    """
    if not rows or len(hdr) < 4 or len(hdr) % 2:
        return None

    def _pairs(r):
        """The (label, value) pairs in one row, or None if it isn't one.

        A label is short and is not a number. A real data row —
        "Chase | 4,442 | 2,366 | 360" — fails on the second cell being
        the label position of the next pair, which is what keeps genuine
        columnar tables out of here.
        """
        if len(r) != len(hdr):
            return None
        got = []
        for i in range(0, len(hdr), 2):
            lab = str(r[i] or "").strip()
            val = str(r[i + 1] or "").strip()
            if not lab:
                continue
            if _is_number(lab) or len(lab) > 40 or len(lab.split()) > 5:
                return None
            got.append((lab, val))
        return got or None

    body = [_pairs(r) for r in rows]
    if not body or any(b is None for b in body):
        return None

    # WHAT TO DO WITH THE HEADER ROW. Three cases, and the third is why
    # this is not simply gated on placeholder names:
    #
    #   col1..colN  nothing was found above the block — flatten the rows.
    #   a real pair the first pair IS the header row (the reader promoted
    #               it) — put it back and flatten everything.
    #   a TITLE     a logo or banner line sits above the block and the
    #               header climb grabbed it: "Halliburton Cementing |
    #               CASING & Services | CEMENTING | RECORD". It is not
    #               data and not a header. Discard it and flatten the
    #               rows — otherwise the block stays unrecognised and the
    #               UWI inside it never reaches the catalog.
    # A REAL TABLE HAS TYPE-CONSISTENT COLUMNS; A PAIR GRID DOES NOT.
    # This is the guard that keeps genuine four-column tables out. In
    #     Sample | Depth | Lithology | Description
    #     S1     | 8100  | Sandstone | fine grained
    #     S2     | 8250  | Shale     | dark
    # every value column is all-numeric or all-text, because it is a
    # column. In a pair grid the same position holds an operator, then a
    # UWI, then a state — mixed by nature, because it is not a column at
    # all. Needs two rows to judge; with one there is nothing to compare.
    if len(body) >= 2:
        mixed = False
        for i in range(0, len(hdr), 2):
            vals = [str(r[i + 1] or "").strip() for r in rows
                    if len(r) == len(hdr) and str(r[i + 1] or "").strip()]
            if len(vals) >= 2 and len({_is_number(v) for v in vals}) > 1:
                mixed = True
                break
        if not mixed:
            return None

    placeholder = all(str(h).strip().lower().startswith("col") for h in hdr)
    head_pairs = None if placeholder else _pairs(list(hdr))
    # A HEADER ROW IS ONLY THE FIRST PAIR IF ITS VALUES LOOK LIKE VALUES.
    # "Spud Date: | 2024-01-08 | Rig Release: | 2024-05-22" is data the
    # reader promoted. "Halliburton Cementing | CASING & Services |
    # CEMENTING | RECORD" is a banner the header climb grabbed — keeping
    # it would file the company's logo as a field.
    if head_pairs:
        vals = [v for _l, v in head_pairs if v]
        if not any(_is_number(v) or re.match(r"^\d{4}-\d{2}-\d{2}", v)
                   for v in vals):
            head_pairs = None
    blocks = ([head_pairs] if head_pairs else []) + body

    labels, values, seen = [], [], set()
    for b in blocks:
        for lab, val in b:
            key, k = lab, 2
            while key.lower() in seen:          # same label twice: keep both
                key, k = f"{lab}_{k}", k + 1
            seen.add(key.lower())
            labels.append(key)
            values.append(val)
    if len(labels) < 4:
        return None
    return labels, [values]


def text_tables(page, min_cols=3, min_rows=3, header_lookback=3,
                with_title=False):
    """[(header, rows)] for whitespace-aligned tables on one page.

    with_title=True yields (header, rows, section_title) instead. Default
    OFF: capture and shape_loader unpack two values, and a signature
    change that breaks working callers to add a nicety is a bad trade.
    """
    lines = _text_lines(page)
    cls = [_gap_clusters(l) for l in lines]

    def _is_row(c):
        """Does this line look like a table row?

        The original test was "more than 60% of the cells are numeric",
        which finds depth tables, curve readings and production — and by
        construction EXCLUDES A PAIR GRID, where the cells are
        label/value/label/value and almost nothing is a number:

            Operator | Devon Energy | Well Name | JONES 15-5
            UWI/API  | 150092023800 | Field     | El Dorado

        That block clusters perfectly into four cells on every line and
        was still dropped, so an end-of-well report's entire header —
        including its UWI — was invisible. Downstream that reads as "no
        detail rows", because with no UWI nothing can be attached to a
        well.

        The second test is STRUCTURAL, not domain-aware: an EVEN number of
        columns, at least four, and a consistent count down the run (the
        run-grouping below already enforces that). Whether the labels mean
        anything is the recogniser's business — a block that matches no
        shape simply shows up unrecognised, which is a visible outcome
        rather than a silent drop.
        """
        if not c or len(c) < min_cols:
            return False
        if _cluster_numfrac(c) > 0.6:
            return True                     # numeric table, as before
        return len(c) >= 4 and len(c) % 2 == 0

    # TWO PASSES. The numeric test stands on its own, but the pair-grid
    # test must not swallow a HEADER: "Depth (ft MD) | Depth (ft TVD) |
    # Formation | …" is eight text cells and would qualify, and absorbing
    # it into the block below leaves the table with col1..col8 and no
    # identity. A text line sitting NEXT TO numeric rows is a header; a
    # text line among other text lines is a pair grid.
    numeric = [bool(c) and len(c) >= min_cols and _cluster_numfrac(c) > 0.6
               for c in cls]
    flags = []
    for n, c in enumerate(cls):
        if numeric[n]:
            flags.append(True)
            continue
        if not (c and len(c) >= 4 and len(c) % 2 == 0):
            flags.append(False)
            continue
        near_numeric = ((n > 0 and numeric[n - 1])
                        or (n + 1 < len(cls) and numeric[n + 1]))
        flags.append(not near_numeric)

    out, i = [], 0
    while i < len(lines):
        if not flags[i]:
            i += 1
            continue
        j = i
        skipped = set()
        while j + 1 < len(lines):
            if flags[j + 1] and len(cls[j + 1]) == len(cls[i]):
                j += 1
                continue
            # A PAIR GRID SURVIVES A BLANK VALUE. When one value in the
            # block is empty the line clusters into FEWER cells —
            #     Operator  | Devon Energy | Well Name | WEST 11-1   (4)
            #     UWI / API |              | Field     | Chase-Silica (3)
            #     State     | KS           | County    | 15051       (4)
            # — and an equal-count run stops dead at that line, throwing
            # away every row above it. Here that cost the Operator and
            # Well Name of an end-of-well report because its UWI was
            # blank, which is precisely the document that most needs a
            # name to match on.
            #
            # So step OVER a single short line and carry on if the run
            # resumes immediately after. The short line's own content is
            # NOT used: with a value missing there is no way to know which
            # label lost it, and pairing positionally would file "Field"
            # as the UWI. Dropping it is the honest half-measure.
            if (not numeric[i] and j + 2 < len(lines)
                    and (j + 1) not in skipped
                    and cls[j + 1] and len(cls[j + 1]) < len(cls[i])
                    and flags[j + 2] and len(cls[j + 2]) == len(cls[i])):
                skipped.add(j + 1)
                j += 2
                continue
            break
        block = [c for n2, c in enumerate(cls[i:j + 1], start=i)
                 if n2 not in skipped]
        if len(block) < min_rows:
            i = j + 1
            continue

        # Column positions from the data itself — the median x-centre of
        # each CELL position across the block, so one ragged row can't
        # shift a column.
        ncol = len(block[0])
        centres = [statistics.median(
            [(r[c][0]["x0"] + r[c][-1]["x1"]) / 2.0 for r in block])
            for c in range(ncol)]
        bounds = [(centres[c] + centres[c + 1]) / 2.0
                  for c in range(ncol - 1)]

        def assign(words):
            cells = [[] for _ in range(ncol)]
            for w in words:
                x = (w["x0"] + w["x1"]) / 2.0
                ci = 0
                while ci < len(bounds) and x > bounds[ci]:
                    ci += 1
                cells[ci].append(w["text"])
            return [" ".join(c).strip() for c in cells]

        header = [""] * ncol
        stopped_at = None
        for back in range(1, header_lookback + 1):
            k = i - back
            if k < 0 or flags[k] or not lines[k]:
                break
            # A numeric-dominant line is DATA even when a ragged cell count
            # kept it out of the block — climbing it stacks values into the
            # header (the RFT failure). Stop at it.
            if _cluster_numfrac(cls[k]) > 0.4:
                break
            part = _char_header(page, lines[k], centres) or assign(lines[k])
            # Stop climbing at the metadata block. A real header line spans
            # most columns; "CONTRACTOR: Baker Hughes  REFERENCE: KB" sits
            # above the table, fills only a few, and carries colons. Without
            # this the column names come out as "CONTRACTOR: Meas Depth (ft)".
            filled = sum(1 for c in part if c)
            if filled < max(2, int(ncol * 0.6)):
                stopped_at = k
                break
            if any(":" in c for c in part):
                stopped_at = k
                break
            # Nearer lines are more specific; a units line sits directly above
            # the data and the names line above that, so prepend as we climb.
            header = [f"{p} {h}".strip() if p else h
                      for p, h in zip(part, header)]
        if not any(header):
            header = [f"col{c + 1}" for c in range(ncol)]
        # THE LINE THE CLIMB STOPPED AT is usually the section TITLE —
        # "Drilling Parameters", "Fluid Sample Analysis". The vendor naming
        # the table type in plain English, which is worth more than any
        # name we could infer: it becomes the proposed shape name for an
        # unrecognised table, and a human label for a recognised one. It is
        # metadata about the table, never a column, so it rides in a
        # sidecar rather than the header.
        title = _section_title(lines, cls, stopped_at if stopped_at is not None
                               else i - header_lookback - 1)
        rows_out = [[_cell_text(c) for c in r] for r in block]
        out.append((header, rows_out, title) if with_title
                   else (header, rows_out))
        i = j + 1
    return out


def _title_above(page, hdr):
    """Section title for a RULED table, located by its header text.

    The ruled path gets cells from pdfplumber and never sees the page's
    line structure, so find the line containing the header's first two
    names and walk up for the nearest title-shaped line. Best effort: a
    missing title costs a nicety, never correctness.
    """
    try:
        lines = _text_lines(page)
        cls = [_gap_clusters(l) for l in lines]
    except Exception:
        return ""
    want = [str(h).strip().lower() for h in hdr[:2] if str(h).strip()]
    if not want:
        return ""
    for idx, l in enumerate(lines):
        text = " ".join(w["text"] for w in l).lower()
        if all(w.split("(")[0].strip() in text for w in want):
            return _section_title(lines, cls, idx - 1)
    return ""


def _section_title(lines, cls, k):
    """Nearest short, non-numeric, colon-free text line at or above k.

    A title is SHORT (a few cells, unlike a header spanning every column),
    carries no colon (that is the metadata block) and no digits-dominant
    content (that is data). Searching upward a few lines survives a blank
    or a units line between title and table.
    """
    for kk in range(k, max(-1, k - 4), -1):
        if kk < 0 or not lines[kk]:
            continue
        cells = cls[kk]
        if not cells or len(cells) > 4:
            continue
        text = " ".join(_cell_text(c) for c in cells).strip()
        if not text or ":" in text or len(text) > 70:
            continue
        letters = sum(ch.isalpha() for ch in text)
        if letters < max(4, len(text) // 2):
            continue
        return text
    return ""

class _TableParser(HTMLParser):
    """Tables out of HTML, with the standard library only.

    HTML is the easiest format there is — the author already declared the
    cell boundaries — so this needs none of the layout guessing the PDF
    path lives with. No bs4, lxml or pandas dependency: those are heavy for
    a job that is a state machine over four tags.

    A heading immediately above a table is kept as its SECTION TITLE, the
    same way the PDF reader keeps the line the header climb stopped at.
    """

    _HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6", "caption", "div", "p"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.tables = []          # [(title, header, rows)]
        self._t = None            # rows of the table being read
        self._row = None
        self._cell = None
        self._is_th = False
        self._depth = 0
        self._text = []           # running text, for the title
        self._in_head = False
        self._last_text = ""

    # -- text ----------------------------------------------------------- #
    def handle_data(self, data):
        if self._cell is not None:
            self._cell.append(data)
        elif self._in_head:
            self._text.append(data)

    # -- structure ------------------------------------------------------- #
    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self._depth += 1
            if self._depth == 1:
                self._t, self._title = [], " ".join(
                    self._last_text.split())[:70]
            return
        if self._depth < 1:
            if tag in self._HEADINGS:
                self._in_head, self._text = True, []
            return
        if tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell, self._is_th = [], (tag == "th")

    def handle_endtag(self, tag):
        if tag in ("td", "th"):
            if self._cell is not None and self._row is not None:
                self._row.append(("TH" if self._is_th else "TD",
                                  " ".join("".join(self._cell).split())))
            self._cell = None
        elif tag == "tr":
            if self._row:
                self._t.append(self._row)
            self._row = None
        elif tag == "table":
            if self._depth == 1 and self._t:
                self.tables.append((self._title, self._t))
                self._t = None
            self._depth = max(0, self._depth - 1)
        elif tag in self._HEADINGS and self._in_head:
            self._in_head = False
            txt = " ".join("".join(self._text).split())
            if txt:
                self._last_text = txt


def html_tables(path, titles=None):
    """{table_name: [rowdict, ...]} from an HTML file.

    The header is the <th> row when there is one, else the first row.
    A row whose cells are all blank except the first is a PLACEHOLDER —
    "No survey data", "No DST data" — which report generators emit for
    sections with nothing in them. Those are dropped: they are the
    document saying a table is absent, and turning them into data rows
    would invent a well test that never happened.
    """
    with open(path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        return {}

    par = _TableParser()
    try:
        par.feed(text)
    except Exception:
        pass

    out = {}
    for i, (title, rows) in enumerate(par.tables, start=1):
        if not rows:
            continue
        # STACKED HEADER/VALUE BLOCKS. A summary table is often written as
        # a <th> row, its values, another <th> row, its values — four
        # blocks of four columns describing ONE subject, not four rows of
        # data. The markup says so plainly (repeated <th> rows, each
        # followed by exactly one <td> row), so the reader can flatten it
        # into the single wide record it is. Purely structural: no idea
        # what a well is required.
        th_ix = [j for j, r in enumerate(rows) if any(k == "TH" for k, _v in r)]
        if (len(th_ix) > 1
                and all(j + 1 < len(rows) and j + 1 not in th_ix
                        for j in th_ix)
                and len(th_ix) * 2 == len(rows)):
            hdr_row, vals = [], []
            for j in th_ix:
                hdr_row += rows[j]
                vals += rows[j + 1]
            body = [vals]
        else:
            hdr_row = next((r for r in rows
                            if any(k == "TH" for k, _v in r)), None)
            body = [r for r in rows if r is not hdr_row]
            if hdr_row is None:
                hdr_row, body = rows[0], rows[1:]
        header, used = [], set()
        for j, (_k, v) in enumerate(hdr_row, start=1):
            h = v or f"col{j}"
            base, n = h, 2
            while h.lower() in used:      # dict(zip) collapses duplicates
                h, n = f"{base}_{n}", n + 1
            used.add(h.lower())
            header.append(h)
        keep = []
        for r in body:
            vals = [v for _k, v in r]
            if not any(v.strip() for v in vals):
                continue
            if len(vals) > 1 and not any(v.strip() for v in vals[1:]):
                continue                  # placeholder: "No survey data"
            keep.append(vals)
        if not keep:
            continue
        name = f"t{i:02d}"
        out[name] = [dict(zip(header, v)) for v in keep]
        if titles is not None and title:
            titles[name] = title
    return out


def raw_tables(path, titles=None):
    """Every table in the file, before any extractor. No classification.

    Pass a dict as `titles` to also receive {table_name: section_title} —
    the vendor's own words for what each table IS ("Fluid Sample
    Analysis"). Opt-in via an out-parameter rather than a changed return
    type, so every existing caller keeps working untouched.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import pdfplumber
            # A ')' immediately followed by a letter inside one header cell
            # — "Pre-Test (psi)Fi" — means the detected ruling cut THROUGH
            # glyphs: the grid pdfplumber found doesn't match the visual
            # columns. No cell-level repair can unscramble that, so reject
            # the table and let the word/char fallback re-read the page.
            bad_hdr = re.compile(r"\)[A-Za-z]")
            out, n = {}, 0
            with pdfplumber.open(path) as pdf:
                for pno, page in enumerate(pdf.pages, start=1):
                    found, rejected = 0, False
                    seen_hdrs = set()
                    # …and a fingerprint of each kept table's DATA. Header
                    # text alone is not enough to dedupe: the text path
                    # re-finds a ruled table under col1..colN placeholders,
                    # so the headers differ and the same rows land twice.
                    # The values do not lie.
                    seen_data = set()
                    for t in (page.extract_tables() or []):
                        if not t or len(t) < 2:
                            continue
                        # A "table" whose every row is one cell is the ruled
                        # strategy finding a text block, not a grid — that's
                        # the borderless case, so don't count it as found.
                        if max(len(r) for r in t) < 2:
                            continue
                        hdr = [(c or "").strip() or f"col{j}"
                               for j, c in enumerate(t[0], start=1)]
                        if any(bad_hdr.search(h) for h in hdr):
                            rejected = True
                            continue
                        n += 1
                        found += 1
                        seen_hdrs.add(tuple(h.lower() for h in hdr))
                        # EVERY data row, not a sample of two: the text
                        # path can re-find a ruled table starting part-way
                        # down it (a mid-table row becomes the header), so
                        # matching only the first rows misses the overlap.
                        for _r in t[1:]:
                            seen_data.add(tuple(
                                (c or "").strip().lower()
                                for c in _r if (c or "").strip()))
                        out[f"p{pno}_t{n:02d}"] = [
                            dict(zip(hdr, [(c or "").strip() for c in r]))
                            for r in t[1:]]
                        if titles is not None:
                            ti = _title_above(page, hdr)
                            if ti:
                                titles[f"p{pno}_t{n:02d}"] = ti
                    # ALWAYS TRY THE TEXT PATH, even when the ruled path
                    # succeeded. It used to `continue` here whenever a grid
                    # was found, on the assumption that a page is either
                    # ruled or not — but a page can be BOTH, and this one
                    # is: two ruled tables plus an unruled header block
                    # above them. Skipping the text path meant the header
                    # was never looked for on any page that happened to
                    # contain a ruled table.
                    #
                    # Safe because the dedupe below already existed for
                    # exactly this: a table the ruled path kept is skipped
                    # by header match, so nothing lands twice. The cost is
                    # one extra pass over pages that gain nothing from it.
                    # Fall back to word positions (char-level header
                    # rebuild inside).
                    for hdr, rows, title in text_tables(page,
                                                        with_title=True):
                        if tuple(str(h).strip().lower()
                                 for h in hdr) in seen_hdrs:
                            continue
                        # Same ROWS as something the ruled path already
                        # kept — a re-detection, not a new table.
                        if any(tuple(str(v).strip().lower()
                                     for v in _r if str(v).strip())
                               in seen_data for _r in rows[:3]):
                            continue
                        # A pair grid carries its labels DOWN the rows,
                        # not across the top. Reshape it so the header is
                        # the labels and the single row is the values —
                        # otherwise it identifies as nothing and the well
                        # it describes is never resolved.
                        _flat = _flatten_pair_grid(hdr, rows)
                        if _flat:
                            hdr, rows = _flat
                        # dict(zip(...)) collapses duplicate or empty
                        # header names, silently eating a column's values —
                        # uniquify before the rows become dicts.
                        uniq, used = [], set()
                        for j, h in enumerate(hdr, start=1):
                            h = str(h).strip() or f"col{j}"
                            base, k = h, 2
                            while h.lower() in used:
                                h = f"{base}_{k}"
                                k += 1
                            used.add(h.lower())
                            uniq.append(h)
                        n += 1
                        out[f"p{pno}_x{n:02d}"] = [
                            dict(zip(uniq, r)) for r in rows]
                        if titles is not None and title:
                            titles[f"p{pno}_x{n:02d}"] = title
            return out
        if ext in (".html", ".htm"):
            return html_tables(path, titles)
        if ext == ".docx":
            return docx_tables(path)
        if ext in (".xlsx", ".xls"):
            return xlsx_sheets(path)
    except Exception as e:
        return {"_error": [{"error": f"{type(e).__name__}: {e}"}]}
    return {}


# --------------------------------------------------------------------------- #
# Workbook
# --------------------------------------------------------------------------- #
_BAD_SHEET = re.compile(r"[\[\]:*?/\\]")
