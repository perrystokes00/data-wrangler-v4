"""
modules/file_viewer.py
======================
Universal file viewer — dispatches to the correct viewer by extension.

Usage:
    from dataview.file_catalog.file_viewer import view
    view(file_path)          # auto-detect extension
    view(file_path, ".las")  # explicit extension

Supported:
  .pdf                -- iframe embed
  .las                -- header + curve list + multi-track log plot
  .dlis .dlf .dis     -- frame/channel browser + curve plot
  .lis                -- channel table + curve plot
  .segy .sgy .seg     -- binary header + trace header table + wiggle plot
  .p190 .p90          -- shot point table + map
  .shp .geojson .gpkg -- attribute table + folium map
  .xlsx .xls .xlsm    -- sheet selector + dataframe preview
  .csv .tsv           -- dataframe preview
  .docx .doc          -- extracted text + table list
  .tif .tiff .png
  .jpg .jpeg          -- rendered image
  *                   -- raw text fallback
"""
import streamlit as st
from pathlib import Path

# ── Extension routing ─────────────────────────────────────────────────────────

PDF_EXTS    = {".pdf"}
LAS_EXTS    = {".las"}
DLIS_EXTS   = {".dlis", ".dlf", ".dis"}
LIS_EXTS    = {".lis"}
SEGY_EXTS   = {".segy", ".sgy", ".seg"}
P190_EXTS   = {".p190", ".p90", ".p1"}
SHP_EXTS    = {".shp", ".geojson", ".gpkg", ".kml", ".kmz"}
EXCEL_EXTS  = {".xlsx", ".xls", ".xlsm"}
CSV_EXTS    = {".csv", ".tsv"}
WORD_EXTS   = {".docx", ".doc"}
IMAGE_EXTS  = {".tif", ".tiff", ".png", ".jpg", ".jpeg"}


def view(file_path: str, file_ext: str = None):
    """
    Main dispatcher. Call this from any page.
    Shows the appropriate viewer for the file type.
    """
    fpath = Path(file_path)
    ext   = (file_ext or fpath.suffix).lower()

    if not fpath.exists():
        st.error(f"File not found: `{file_path}`")
        return

    st.caption(
        f"`{file_path}` · "
        f"{fpath.stat().st_size / 1024:.1f} KB · "
        f"{ext}"
    )

    if ext in PDF_EXTS:
        _view_pdf(file_path)
    elif ext in LAS_EXTS:
        _view_las(file_path)
    elif ext in DLIS_EXTS:
        _view_dlis(file_path)
    elif ext in LIS_EXTS:
        _view_lis(file_path)
    elif ext in SEGY_EXTS:
        _view_segy(file_path)
    elif ext in P190_EXTS:
        _view_p190(file_path)
    elif ext in SHP_EXTS:
        _view_shapefile(file_path)
    elif ext in EXCEL_EXTS:
        _view_excel(file_path)
    elif ext in CSV_EXTS:
        _view_csv(file_path)
    elif ext in WORD_EXTS:
        _view_word(file_path)
    elif ext in IMAGE_EXTS:
        _view_image(file_path)
    else:
        _view_text_fallback(file_path)


# =============================================================================
# PDF
# =============================================================================


from contextlib import contextmanager as _contextmanager

@_contextmanager
def _vsection(label: str):
    """Nest-safe replacement for st.expander in embeddable viewers: a bordered
    container with a bold label. Unlike expander, containers can be nested inside
    an expander (so these viewers work when embedded in the Documents page)."""
    import streamlit as st
    box = st.container(border=True)
    with box:
        if label:
            st.markdown(f"**{label}**")
        yield box


def _view_pdf(file_path: str):
    """Render a PDF as page images via PyMuPDF. Simple and fast for the typical
    1–few page scout tickets / reports. Never base64's the file into the DOM
    (that crashed the browser) and never re-reads the whole file on every rerun.
    """
    import os
    st.caption(os.path.basename(file_path))

    try:
        import fitz  # PyMuPDF
    except Exception:
        st.warning("PyMuPDF (fitz) isn't installed, so the PDF can't be "
                   "previewed in-app. Run: pip install PyMuPDF")
        if st.button("Open in native app", key=f"pdfnat_{file_path}"):
            _open_native(file_path) if "_open_native" in globals() else None
        return

    try:
        with st.spinner("Rendering…"):
            doc = fitz.open(file_path)
            npages = doc.page_count
            # cap at 10 pages so a fat report can't stall the render
            for i in range(min(npages, 10)):
                pix = doc.load_page(i).get_pixmap(dpi=100)
                st.image(pix.tobytes("png"),
                         caption=f"Page {i+1} of {npages}",
                         use_container_width=True)
            doc.close()
        if npages > 10:
            st.caption(f"Showing first 10 of {npages} pages.")
    except Exception as e:
        st.error(f"PDF render failed: {e}")
    return


# =============================================================================
# LAS
# =============================================================================

def _view_las(file_path: str):
    import pandas as pd
    try:
        import lasio
        las = lasio.read(file_path)
    except Exception as e:
        st.error(f"LAS read failed: {e}")
        _view_text_fallback(file_path)
        return

    # ── Well info ─────────────────────────────────────────────────────────────
    well_items = [
        {"Mnemonic": item.mnemonic,
         "Unit":     item.unit,
         "Value":    str(item.value),
         "Description": item.descr}
        for item in las.well
    ]
    if well_items:
        with _vsection("📋 Well header"):
            st.dataframe(pd.DataFrame(well_items),
                         hide_index=True, use_container_width=True)

    # ── Curves ───────────────────────────────────────────────────────────────
    curve_items = [
        {"Mnemonic": c.mnemonic,
         "Unit":     c.unit,
         "Description": c.descr}
        for c in las.curves
    ]
    with _vsection(f"📈 Curves ({len(curve_items)})"):
        st.dataframe(pd.DataFrame(curve_items),
                     hide_index=True, use_container_width=True)

    # ── Log plot ──────────────────────────────────────────────────────────────
    df = las.df().reset_index()
    depth_col = df.columns[0]

    curve_names = [c for c in df.columns if c != depth_col]
    if not curve_names:
        st.info("No curve data.")
        return

    selected = st.multiselect(
        "Curves to plot",
        curve_names,
        default=curve_names[:min(4, len(curve_names))],
        key=f"las_curves_{file_path}",
    )
    if not selected:
        return

    _las_log_plot(df, depth_col, selected, file_path)

    # Download
    st.download_button(
        "⬇ Download curve data CSV",
        data=df.to_csv(index=False),
        file_name=Path(file_path).stem + "_curves.csv",
        mime="text/csv",
        key=f"las_dl_{file_path}",
    )


def _las_log_plot(df, depth_col, curves, file_path):
    try:
        import plotly.graph_objects as go
        from plotly.subplots import make_subplots
        import numpy as np

        n     = len(curves)
        depth = df[depth_col].values
        NAVY  = "#1A2B4A"
        COLORS = ["#C8922A","#1A7BBF","#2CA02C","#D62728",
                  "#9467BD","#8C564B","#E377C2","#7F7F7F"]

        fig = make_subplots(
            rows=1, cols=n,
            shared_yaxes=True,
            subplot_titles=curves,
            horizontal_spacing=0.02,
        )

        for i, curve in enumerate(curves):
            vals = df[curve].replace([np.inf, -np.inf], np.nan).values
            color = COLORS[i % len(COLORS)]
            fig.add_trace(
                go.Scatter(
                    x=vals, y=depth,
                    mode="lines",
                    name=curve,
                    line=dict(color=color, width=1),
                    hovertemplate=f"{curve}: %{{x:.3f}}<br>Depth: %{{y:.1f}}<extra></extra>",
                ),
                row=1, col=i + 1,
            )

        fig.update_yaxes(autorange="reversed", title_text="Depth", col=1)
        fig.update_layout(
            height=700,
            margin=dict(l=60, r=20, t=40, b=20),
            showlegend=False,
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
        )
        for i in range(1, n + 1):
            fig.update_xaxes(showgrid=True,
                             gridcolor="rgba(128,128,128,0.15)",
                             row=1, col=i)
        fig.update_yaxes(showgrid=True,
                         gridcolor="rgba(128,128,128,0.15)")

        st.plotly_chart(fig, use_container_width=True)

    except Exception as e:
        st.warning(f"Plot failed: {e}")
        st.dataframe(df[[depth_col] + curves].head(200),
                     hide_index=True, use_container_width=True)


# =============================================================================
# DLIS
# =============================================================================

def _view_dlis(file_path: str):
    import pandas as pd
    try:
        import dlisio
        f, *tail = dlisio.dlis.load(file_path)
    except Exception as e:
        st.error(f"DLIS read failed: {e}")
        return

    try:
        # Origins
        origins = list(f.origins)
        if origins:
            o = origins[0]
            st.markdown("**File origin**")
            orig_items = []
            for attr in ["well_name","field_name","company","country",
                         "producer_name","creation_time"]:
                v = getattr(o, attr, None)
                if v:
                    orig_items.append({"Field": attr.replace("_"," ").title(),
                                       "Value": str(v)})
            if orig_items:
                st.dataframe(pd.DataFrame(orig_items),
                             hide_index=True, use_container_width=True)

        # Frames
        frames = list(f.frames)
        if not frames:
            st.warning("No frames found.")
            return

        frame_names = [fr.name for fr in frames]
        sel_frame = st.selectbox("Frame", frame_names,
                                 key=f"dlis_frame_{file_path}")
        frame = next(fr for fr in frames if fr.name == sel_frame)

        # Channels
        channels = list(frame.channels)
        ch_info = [{"Name": ch.name, "Unit": str(ch.units),
                    "Dimension": str(ch.dimension)}
                   for ch in channels]
        with _vsection(f"📡 Channels ({len(channels)})"):
            st.dataframe(pd.DataFrame(ch_info),
                         hide_index=True, use_container_width=True)

        # Curve data
        ch_names = [ch.name for ch in channels]
        selected = st.multiselect(
            "Curves to plot", ch_names,
            default=ch_names[:min(4, len(ch_names))],
            key=f"dlis_curves_{file_path}",
        )
        if not selected:
            return

        curves = frame.curves()
        import numpy as np
        depth_ch = channels[0]
        depth    = curves[depth_ch.name].flatten()

        df = pd.DataFrame({"DEPTH": depth})
        for name in selected:
            try:
                arr = curves[name]
                if arr.ndim > 1:
                    arr = arr[:, 0]
                df[name] = arr
            except Exception:
                pass

        _las_log_plot(df, "DEPTH", selected, file_path + "_dlis")

    except Exception as e:
        st.error(f"DLIS parse error: {e}")
    finally:
        try:
            f.close()
            for t in tail:
                t.close()
        except Exception:
            pass


# =============================================================================
# LIS
# =============================================================================

def _view_lis(file_path: str):
    import pandas as pd
    st.info("LIS viewer — basic channel extraction.")
    try:
        # Try using dlisio for LIS (it supports both)
        import dlisio
        f, *tail = dlisio.lis.load(file_path)
        recs = list(f.data_format_specs())
        if not recs:
            st.warning("No data records found.")
            return

        sel_rec = st.selectbox(
            "Data record", range(len(recs)),
            format_func=lambda i: f"Record {i+1}",
            key=f"lis_rec_{file_path}",
        )
        rec = recs[sel_rec]
        channels = rec.specs

        ch_info = [{"Name": ch.mnemonic, "Unit": ch.units,
                    "Size": ch.size}
                   for ch in channels]
        with _vsection(f"📡 Channels ({len(channels)})"):
            st.dataframe(pd.DataFrame(ch_info),
                         hide_index=True, use_container_width=True)

        curves = dlisio.lis.curves(f, rec)
        ch_names = [ch.mnemonic for ch in channels]
        selected = st.multiselect(
            "Curves to plot", ch_names,
            default=ch_names[:min(4, len(ch_names))],
            key=f"lis_curves_{file_path}",
        )
        if not selected:
            return

        import numpy as np
        depth_name = ch_names[0]
        depth = curves[depth_name].flatten()
        df = pd.DataFrame({"DEPTH": depth})
        for name in selected:
            try:
                arr = curves[name]
                if arr.ndim > 1:
                    arr = arr[:, 0]
                df[name] = arr
            except Exception:
                pass

        _las_log_plot(df, "DEPTH", selected, file_path + "_lis")

    except Exception as e:
        st.error(f"LIS read failed: {e}")
        _view_text_fallback(file_path)


# =============================================================================
# SEGY
# =============================================================================

def _view_segy(file_path: str):
    import pandas as pd
    try:
        import segyio
    except ImportError:
        st.error("segyio not installed — `pip install segyio`")
        return

    try:
        with segyio.open(file_path, ignore_geometry=True) as f:
            # Binary header
            bin_header = {
                "Traces":        f.tracecount,
                "Samples/trace": f.bin[segyio.BinField.Samples],
                "Sample interval (us)": f.bin[segyio.BinField.Interval],
                "Format":        f.bin[segyio.BinField.Format],
            }
            with _vsection("🗂️ Binary header"):
                st.dataframe(
                    pd.DataFrame([{"Field": k, "Value": str(v)}
                                  for k, v in bin_header.items()]),
                    hide_index=True, use_container_width=True,
                )

            # Trace header sample
            n_preview = min(50, f.tracecount)
            header_keys = [
                segyio.TraceField.CDP,
                segyio.TraceField.FieldRecord,
                segyio.TraceField.TRACE_SEQUENCE_FILE,
                segyio.TraceField.offset,
                segyio.TraceField.CDP_X,
                segyio.TraceField.CDP_Y,
            ]
            h_rows = []
            for i in range(n_preview):
                h = f.header[i]
                h_rows.append({
                    "Trace":  i + 1,
                    "CDP":    h[segyio.TraceField.CDP],
                    "FieldRec": h[segyio.TraceField.FieldRecord],
                    "CDP_X":  h[segyio.TraceField.CDP_X],
                    "CDP_Y":  h[segyio.TraceField.CDP_Y],
                    "Offset": h[segyio.TraceField.offset],
                })
            with _vsection(f"📋 Trace headers (first {n_preview})"):
                st.dataframe(pd.DataFrame(h_rows),
                             hide_index=True, use_container_width=True)

            # Wiggle / density plot
            n_plot = min(100, f.tracecount)
            traces_to_plot = st.slider(
                "Traces to display", 10, min(500, f.tracecount), n_plot,
                key=f"segy_n_{file_path}",
            )

            import numpy as np
            data = np.zeros((f.samples.size, traces_to_plot))
            for i in range(traces_to_plot):
                data[:, i] = f.trace[i]

            _segy_plot(data, f.samples, traces_to_plot, file_path)

    except Exception as e:
        st.error(f"SEGY read failed: {e}")


def _segy_plot(data, samples, n_traces, file_path):
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np

        fig, axes = plt.subplots(1, 2, figsize=(14, 8),
                                 gridspec_kw={"width_ratios": [1, 1]})
        fig.patch.set_alpha(0.0)

        # Density (variable area)
        ax1 = axes[0]
        extent = [0, n_traces, samples[-1], samples[0]]
        vmax   = np.percentile(np.abs(data), 95) or 1
        ax1.imshow(data, aspect="auto", cmap="RdBu_r",
                   vmin=-vmax, vmax=vmax, extent=extent)
        ax1.set_title("Density (variable area)", color="white")
        ax1.set_xlabel("Trace", color="white")
        ax1.set_ylabel("Sample", color="white")
        ax1.tick_params(colors="white")
        ax1.set_facecolor("#1A2B4A")

        # Wiggle (first 30 traces max)
        ax2 = axes[1]
        n_wig = min(30, n_traces)
        norm  = vmax * 2 if vmax > 0 else 1
        for i in range(n_wig):
            trace = data[:, i] / norm + i
            ax2.plot(trace, samples, "k-", linewidth=0.4, alpha=0.7)
            ax2.fill_betweenx(samples, i, trace,
                              where=(trace > i), color="#C8922A", alpha=0.4)
        ax2.set_xlim(-1, n_wig)
        ax2.invert_yaxis()
        ax2.set_title(f"Wiggle (first {n_wig} traces)", color="white")
        ax2.set_xlabel("Trace", color="white")
        ax2.set_ylabel("Sample", color="white")
        ax2.tick_params(colors="white")
        ax2.set_facecolor("#1A2B4A")

        fig.tight_layout()
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)

    except Exception as e:
        st.warning(f"SEGY plot failed: {e}")


# =============================================================================
# P190
# =============================================================================

def _view_p190(file_path: str):
    import pandas as pd
    try:
        from dataview.file_catalog.p190_catalog import parse_p190
        result = parse_p190(file_path)
    except ImportError:
        result = _parse_p190_basic(file_path)
    except Exception as e:
        st.error(f"P190 parse failed: {e}")
        return

    header = result.get("header", {})
    shots  = result.get("shots", [])

    if header:
        with _vsection("📋 File header"):
            st.dataframe(
                pd.DataFrame([{"Field": k, "Value": str(v)}
                              for k, v in header.items() if v]),
                hide_index=True, use_container_width=True,
            )

    if shots:
        df = pd.DataFrame(shots)
        st.metric("Shot points", len(df))
        st.dataframe(df.head(200), hide_index=True, use_container_width=True)

        # Map if lat/lon available
        lat_col = next((c for c in df.columns
                        if c.upper() in ("LAT","LATITUDE","SHOT_LAT")), None)
        lon_col = next((c for c in df.columns
                        if c.upper() in ("LON","LONG","LONGITUDE","SHOT_LON")), None)
        if lat_col and lon_col:
            _folium_points(df, lat_col, lon_col, label="Shot point")

        st.download_button(
            "⬇ Download shot points CSV",
            data=df.to_csv(index=False),
            file_name=Path(file_path).stem + "_shots.csv",
            mime="text/csv",
            key=f"p190_dl_{file_path}",
        )
    else:
        st.info("No shot point data found.")
        _view_text_fallback(file_path)


def _parse_p190_basic(file_path: str) -> dict:
    """Minimal P190 parser — reads H records (header) and S records (shots)."""
    import re
    header = {}
    shots  = []
    try:
        with open(file_path, "r", errors="replace") as fh:
            for line in fh:
                if not line.strip():
                    continue
                rec = line[0].upper()
                if rec == "H":
                    parts = line[1:].split()
                    if len(parts) >= 2:
                        header[parts[0]] = " ".join(parts[1:])
                elif rec == "S":
                    parts = line.split()
                    if len(parts) >= 6:
                        try:
                            shots.append({
                                "LINE":   parts[1],
                                "SP":     parts[2],
                                "EASTING":  float(parts[3]),
                                "NORTHING": float(parts[4]),
                                "DEPTH":    float(parts[5]) if len(parts) > 5 else None,
                            })
                        except ValueError:
                            pass
    except Exception:
        pass
    return {"header": header, "shots": shots}


# =============================================================================
# Shapefile
# =============================================================================

def _view_shapefile(file_path: str):
    import pandas as pd
    try:
        import geopandas as gpd
        gdf = gpd.read_file(file_path)
    except ImportError:
        st.error("geopandas not installed.")
        return
    except Exception as e:
        st.error(f"Shapefile read failed: {e}")
        return

    m1, m2, m3 = st.columns(3)
    m1.metric("Features",  len(gdf))
    m2.metric("Columns",   len(gdf.columns))
    m3.metric("CRS",       str(gdf.crs) if gdf.crs else "None")

    # Attribute table
    display_df = gdf.drop(columns=["geometry"], errors="ignore")
    with _vsection("📋 Attribute table"):
        st.dataframe(display_df.head(500),
                     hide_index=True, use_container_width=True)

    # Map
    try:
        _shp_map(gdf, file_path)
    except Exception as e:
        st.warning(f"Map unavailable: {e}")

    # Download attribute CSV
    st.download_button(
        "⬇ Download attributes CSV",
        data=display_df.to_csv(index=False),
        file_name=Path(file_path).stem + "_attrs.csv",
        mime="text/csv",
        key=f"shp_dl_{file_path}",
    )


def _shp_map(gdf, file_path: str):
    import folium
    from streamlit_folium import st_folium
    import geopandas as gpd

    # Reproject to WGS84 for display
    try:
        gdf_wgs = gdf.to_crs(epsg=4326)
    except Exception:
        gdf_wgs = gdf

    bounds = gdf_wgs.total_bounds  # minx, miny, maxx, maxy
    center = [(bounds[1] + bounds[3]) / 2,
              (bounds[0] + bounds[2]) / 2]

    m = folium.Map(location=center, zoom_start=8,
                   tiles="CartoDB positron")

    # Add features
    geom_type = gdf_wgs.geometry.geom_type.iloc[0] if len(gdf_wgs) else "Unknown"

    label_col = next(
        (c for c in gdf_wgs.columns
         if c.upper() in ("NAME","WELL_NAME","UWI","ID","LABEL","API")),
        None,
    )

    for _, row in gdf_wgs.iterrows():
        tooltip = str(row[label_col]) if label_col else ""
        try:
            geom = row.geometry
            if geom is None or geom.is_empty:
                continue
            if geom.geom_type == "Point":
                folium.CircleMarker(
                    location=[geom.y, geom.x],
                    radius=4,
                    color="#C8922A",
                    fill=True,
                    fill_opacity=0.7,
                    tooltip=tooltip,
                ).add_to(m)
            elif geom.geom_type in ("LineString", "MultiLineString"):
                folium.GeoJson(
                    geom.__geo_interface__,
                    style_function=lambda x: {
                        "color": "#1A2B4A", "weight": 2},
                    tooltip=tooltip,
                ).add_to(m)
            else:
                folium.GeoJson(
                    geom.__geo_interface__,
                    style_function=lambda x: {
                        "fillColor": "#1A7BBF",
                        "color": "#1A2B4A",
                        "weight": 1,
                        "fillOpacity": 0.3},
                    tooltip=tooltip,
                ).add_to(m)
        except Exception:
            pass

    m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
    st_folium(m, use_container_width=True, height=500,
              key=f"shp_map_{file_path}")


# =============================================================================
# Excel / CSV
# =============================================================================

def _view_excel(file_path: str):
    import pandas as pd
    try:
        xl = pd.ExcelFile(file_path)
    except Exception as e:
        st.error(f"Excel read failed: {e}")
        return

    sheet = st.selectbox("Sheet", xl.sheet_names,
                         key=f"xl_sheet_{file_path}")
    try:
        df = xl.parse(sheet)
        st.metric("Rows", f"{len(df):,}")
        st.dataframe(df.head(500), hide_index=True, use_container_width=True)
        st.download_button(
            "⬇ Download sheet CSV",
            data=df.to_csv(index=False),
            file_name=f"{Path(file_path).stem}_{sheet}.csv",
            mime="text/csv",
            key=f"xl_dl_{file_path}",
        )
    except Exception as e:
        st.error(f"Sheet read failed: {e}")


def _view_csv(file_path: str):
    import pandas as pd
    ext = Path(file_path).suffix.lower()
    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(file_path, sep=sep, nrows=5000,
                         encoding="utf-8", on_bad_lines="skip")
        st.metric("Rows (preview)", f"{len(df):,}")
        st.dataframe(df, hide_index=True, use_container_width=True)
        st.download_button(
            "⬇ Download CSV",
            data=df.to_csv(index=False),
            file_name=Path(file_path).name,
            mime="text/csv",
            key=f"csv_dl_{file_path}",
        )
    except Exception as e:
        st.error(f"CSV read failed: {e}")
        _view_text_fallback(file_path)


# =============================================================================
# Word / DOCX
# =============================================================================

def _view_word(file_path: str):
    import pandas as pd
    ext = Path(file_path).suffix.lower()

    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
        except ImportError:
            st.error("python-docx not installed.")
            _view_text_fallback(file_path)
            return
        except Exception as e:
            st.error(f"DOCX read failed: {e}")
            return

        # Paragraphs
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        with _vsection("📝 Document text"):
            st.text(text[:5000])
            if len(text) > 5000:
                st.caption(f"... ({len(text):,} chars total, showing first 5,000)")

        # Tables
        if doc.tables:
            with _vsection(f"📊 Tables ({len(doc.tables)})"):
                for i, tbl in enumerate(doc.tables):
                    rows = [[cell.text for cell in row.cells]
                            for row in tbl.rows]
                    if rows:
                        headers = rows[0]
                        data    = rows[1:]
                        try:
                            df = pd.DataFrame(data, columns=headers)
                            st.markdown(f"**Table {i+1}**")
                            st.dataframe(df, hide_index=True,
                                         use_container_width=True)
                        except Exception:
                            pass
    else:
        # .doc — try antiword / textract fallback
        _view_text_fallback(file_path)


# =============================================================================
# Image
# =============================================================================

def _view_image(file_path: str):
    ext = Path(file_path).suffix.lower()
    try:
        if ext in (".tif", ".tiff"):
            try:
                from PIL import Image
                import io
                img = Image.open(file_path)
                # Convert to RGB for display
                if img.mode not in ("RGB", "RGBA", "L"):
                    img = img.convert("RGB")
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                st.image(buf.getvalue(), use_container_width=True,
                         caption=Path(file_path).name)
                st.caption(
                    f"Size: {img.width} × {img.height} px · "
                    f"Mode: {img.mode} · "
                    f"Bands: {len(img.getbands())}"
                )
            except ImportError:
                st.error("Pillow not installed — `pip install Pillow`")
        else:
            st.image(file_path, use_container_width=True,
                     caption=Path(file_path).name)
    except Exception as e:
        st.error(f"Image render failed: {e}")


# =============================================================================
# Text fallback
# =============================================================================

def _view_text_fallback(file_path: str, max_chars: int = 5000):
    st.markdown("**Raw text (first 5,000 chars)**")
    try:
        encodings = ["utf-8", "latin-1", "cp1252"]
        text = None
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc, errors="replace") as fh:
                    text = fh.read(max_chars)
                break
            except Exception:
                continue
        if text:
            st.code(text, language=None)
        else:
            st.warning("Could not read file as text.")
    except Exception as e:
        st.error(f"Text read failed: {e}")


# =============================================================================
# Folium point helper (shared)
# =============================================================================

def _folium_points(df, lat_col: str, lon_col: str, label: str = ""):
    try:
        import folium
        from streamlit_folium import st_folium

        valid = df[[lat_col, lon_col]].dropna()
        if valid.empty:
            return

        center = [valid[lat_col].mean(), valid[lon_col].mean()]
        m = folium.Map(location=center, zoom_start=8,
                       tiles="CartoDB positron")

        for _, row in df.iterrows():
            try:
                lat = float(row[lat_col])
                lon = float(row[lon_col])
                folium.CircleMarker(
                    location=[lat, lon],
                    radius=4,
                    color="#C8922A",
                    fill=True,
                    fill_opacity=0.8,
                ).add_to(m)
            except Exception:
                pass

        st_folium(m, use_container_width=True, height=400,
                  key=f"fmap_{lat_col}_{lon_col}")
    except ImportError:
        st.info("Install streamlit-folium for map display.")
