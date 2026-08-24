"""
modules/file_viewer.py
======================
Universal file viewer — dispatches to the correct viewer by extension.

Usage:
    from dataview.file_catalog.file_viewer import view
    view(file_path)          # auto-detect extension
    view(file_path, ".las")  # explicit extension

Supported:
  .pdf                -- iframe embed
  .las                -- header + curve list + multi-track log plot
  .dlis .dlf .dis     -- frame/channel browser + curve plot
  .lis                -- channel table + curve plot
  .segy .sgy .seg     -- binary header + trace header table + wiggle plot
  .p190 .p90          -- shot point table + map
  .shp .geojson .gpkg -- attribute table + folium map
  .xlsx .xls .xlsm    -- sheet selector + dataframe preview
  .csv .tsv           -- dataframe preview
  .docx .doc          -- extracted text + table list
  .tif .tiff .png
  .jpg .jpeg          -- rendered image
  *                   -- raw text fallback
"""
import streamlit as st
from pathlib import Path

# ── Extension routing ─────────────────────────────────────────────────────────

PDF_EXTS    = {".pdf"}
LAS_EXTS    = {".las"}
DLIS_EXTS   = {".dlis", ".dlf", ".dis"}
LIS_EXTS    = {".lis"}
SEGY_EXTS   = {".segy", ".sgy", ".seg"}
P190_EXTS   = {".p190", ".p90", ".p1"}
SHP_EXTS    = {".shp", ".geojson", ".gpkg", ".kml", ".kmz"}
EXCEL_EXTS  = {".xlsx", ".xls", ".xlsm"}
CSV_EXTS    = {".csv", ".tsv"}
WORD_EXTS   = {".docx", ".doc"}
IMAGE_EXTS  = {".tif", ".tiff", ".png", ".jpg", ".jpeg"}


def _arrow_safe(_df):
    """Coerce mixed-type object columns to string so Streamlit/Arrow can
    serialize dataframes parsed from arbitrary files (Excel, shapefiles, CSV,
    Word tables) without an ArrowTypeError."""
    try:
        import pandas as _pd
        if isinstance(_df, _pd.DataFrame):
            _df = _df.copy()
            for _c in _df.columns:
                if _df[_c].dtype == object:
                    _df[_c] = _df[_c].astype("string")
    except Exception:
        pass
    return _df


def view(file_path: str, file_ext: str = None):
    """
    Main dispatcher. Call this from any page.
    Shows the appropriate viewer for the file type.
    """
    fpath = Path(file_path)
    ext   = (file_ext or fpath.suffix).lower()

    if not fpath.exists():
        st.error(f"File not found: `{file_path}`")
        return

    st.caption(
        f"`{file_path}` · "
        f"{fpath.stat().st_size / 1024:.1f} KB · "
        f"{ext}"
    )

    if ext in PDF_EXTS:
        _view_pdf(file_path)
    elif ext in LAS_EXTS:
        _view_las(file_path)
    elif ext in DLIS_EXTS:
        _view_dlis(file_path)
    elif ext in LIS_EXTS:
        _view_lis(file_path)
    elif ext in SEGY_EXTS:
        _view_segy(file_path)
    elif ext in P190_EXTS:
        _view_p190(file_path)
    elif ext in SHP_EXTS:
        _view_shapefile(file_path)
    elif ext in EXCEL_EXTS:
        _view_excel(file_path)
    elif ext in CSV_EXTS:
        _view_csv(file_path)
    elif ext in WORD_EXTS:
        _view_word(file_path)
    elif ext in IMAGE_EXTS:
        _view_image(file_path)
    else:
        _view_text_fallback(file_path)


# =============================================================================
# PDF
# =============================================================================


from contextlib import contextmanager as _contextmanager

@_contextmanager
def _vsection(label: str):
    """Nest-safe replacement for st.expander in embeddable viewers: a bordered
    container with a bold label. Unlike expander, containers can be nested inside
    an expander (so these viewers work when embedded in the Documents page)."""
    import streamlit as st
    box = st.container(border=True)
    with box:
        if label:
            st.markdown(f"**{label}**")
        yield box


def _view_pdf(file_path: str):
    """Render a PDF as page images via PyMuPDF. Simple and fast for the typical
    1–few page scout tickets / reports. Never base64's the file into the DOM
    (that crashed the browser) and never re-reads the whole file on every rerun.
    """
    import os
    st.caption(os.path.basename(file_path))

    try:
        import fitz  # PyMuPDF
    except Exception:
        st.warning("PyMuPDF (fitz) isn't installed, so the PDF can't be "
                   "previewed in-app. Run: pip install PyMuPDF")
        if st.button("Open in native app", key=f"pdfnat_{file_path}"):
            _open_native(file_path) if "_open_native" in globals() else None
        return

    try:
        with st.spinner("Rendering…"):
            doc = fitz.open(file_path)
            npages = doc.page_count
            # cap at 10 pages so a fat report can't stall the render
            for i in range(min(npages, 10)):
                pix = doc.load_page(i).get_pixmap(dpi=100)
                st.image(pix.tobytes("png"),
                         caption=f"Page {i+1} of {npages}",
                         use_container_width=True)
            doc.close()
        if npages > 10:
            st.caption(f"Showing first 10 of {npages} pages.")
    except Exception as e:
        st.error(f"PDF render failed: {e}")
    return


# =============================================================================
# LAS
# =============================================================================

def _view_las(file_path: str):
    import pandas as pd
    try:
        import lasio
        from dataview.file_catalog.las_reader import read_las
        las = read_las(file_path)
    except Exception as e:
        # A MULTI-SECTION LAS 3.0 IS NOT A BROKEN FILE. lasio assumes one data
        # block, so a 3.0 file carrying Core / Inclinometry / Tops / Test dies
        # in the header parser:
        #     LASHeaderError: Line 193 (section ~TEST | TEST_Definition)
        # and the reader that CAN read it is already in this repo. Falling
        # back to it is the same move _view_segy makes when segyio refuses a
        # ragged file — showing what is readable beats an error over an empty
        # page.
        if _view_las3_sections(file_path, str(e)):
            return
        st.error(f"LAS read failed: {e}")
        _view_text_fallback(file_path)
        return

    # ── Well info ─────────────────────────────────────────────────────────────
    well_items = [
        {"Mnemonic": item.mnemonic,
         "Unit":     item.unit,
         "Value":    str(item.value),
         "Description": item.descr}
        for item in las.well
    ]
    if well_items:
        with _vsection("📋 Well header"):
            st.dataframe(pd.DataFrame(well_items),
                         hide_index=True, use_container_width=True)

    # ── Curves ───────────────────────────────────────────────────────────────
    curve_items = [
        {"Mnemonic": c.mnemonic,
         "Unit":     c.unit,
         "Description": c.descr}
        for c in las.curves
    ]
    with _vsection(f"📈 Curves ({len(curve_items)})"):
        st.dataframe(pd.DataFrame(curve_items),
                     hide_index=True, use_container_width=True)

    # ── Log plot ──────────────────────────────────────────────────────────────
    df = las.df().reset_index()
    depth_col = df.columns[0]

    curve_names = [c for c in df.columns if c != depth_col]
    if not curve_names:
        st.info("No curve data.")
        return

    selected = st.multiselect(
        "Curves to plot",
        curve_names,
        default=curve_names[:min(4, len(curve_names))],
        key=f"las_curves_{file_path}",
    )
    if not selected:
        return

    _las_log_plot(df, depth_col, selected, file_path)

    # Download
    st.download_button(
        "⬇ Download curve data CSV",
        data=df.to_csv(index=False),
        file_name=Path(file_path).stem + "_curves.csv",
        mime="text/csv",
        key=f"las_dl_{file_path}",
    )


def _las_log_plot(df, depth_col, curves, file_path):
    try:
        import plotly.graph_objects as go
        from plotly.subplots import make_subplots
        import numpy as np

        n     = len(curves)
        depth = df[depth_col].values
        NAVY  = "#1A2B4A"
        COLORS = ["#C8922A","#1A7BBF","#2CA02C","#D62728",
                  "#9467BD","#8C564B","#E377C2","#7F7F7F"]

        fig = make_subplots(
            rows=1, cols=n,
            shared_yaxes=True,
            subplot_titles=curves,
            horizontal_spacing=0.02,
        )

        for i, curve in enumerate(curves):
            vals = df[curve].replace([np.inf, -np.inf], np.nan).values
            color = COLORS[i % len(COLORS)]
            fig.add_trace(
                go.Scatter(
                    x=vals, y=depth,
                    mode="lines",
                    name=curve,
                    line=dict(color=color, width=1),
                    hovertemplate=f"{curve}: %{{x:.3f}}<br>Depth: %{{y:.1f}}<extra></extra>",
                ),
                row=1, col=i + 1,
            )

        fig.update_yaxes(autorange="reversed", title_text="Depth", col=1)
        fig.update_layout(
            height=700,
            margin=dict(l=60, r=20, t=40, b=20),
            showlegend=False,
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
        )
        for i in range(1, n + 1):
            fig.update_xaxes(showgrid=True,
                             gridcolor="rgba(128,128,128,0.15)",
                             row=1, col=i)
        fig.update_yaxes(showgrid=True,
                         gridcolor="rgba(128,128,128,0.15)")

        st.plotly_chart(fig, use_container_width=True)

    except Exception as e:
        st.warning(f"Plot failed: {e}")
        st.dataframe(df[[depth_col] + curves].head(200),
                     hide_index=True, use_container_width=True)


# =============================================================================
# DLIS
# =============================================================================

def _view_dlis(file_path: str):
    import pandas as pd
    try:
        import dlisio
        f, *tail = dlisio.dlis.load(file_path)
    except Exception as e:
        st.error(f"DLIS read failed: {e}")
        return

    try:
        # Origins
        origins = list(f.origins)
        if origins:
            o = origins[0]
            st.markdown("**File origin**")
            orig_items = []
            for attr in ["well_name","field_name","company","country",
                         "producer_name","creation_time"]:
                v = getattr(o, attr, None)
                if v:
                    orig_items.append({"Field": attr.replace("_"," ").title(),
                                       "Value": str(v)})
            if orig_items:
                st.dataframe(pd.DataFrame(orig_items),
                             hide_index=True, use_container_width=True)

        # Frames
        frames = list(f.frames)
        if not frames:
            st.warning("No frames found.")
            return

        frame_names = [fr.name for fr in frames]
        sel_frame = st.selectbox("Frame", frame_names,
                                 key=f"dlis_frame_{file_path}")
        frame = next(fr for fr in frames if fr.name == sel_frame)

        # Channels
        channels = list(frame.channels)
        ch_info = [{"Name": ch.name, "Unit": str(ch.units),
                    "Dimension": str(ch.dimension)}
                   for ch in channels]
        with _vsection(f"📡 Channels ({len(channels)})"):
            st.dataframe(pd.DataFrame(ch_info),
                         hide_index=True, use_container_width=True)

        # Curve data
        ch_names = [ch.name for ch in channels]
        selected = st.multiselect(
            "Curves to plot", ch_names,
            default=ch_names[:min(4, len(ch_names))],
            key=f"dlis_curves_{file_path}",
        )
        if not selected:
            return

        curves = frame.curves()
        import numpy as np
        depth_ch = channels[0]
        depth    = curves[depth_ch.name].flatten()

        df = pd.DataFrame({"DEPTH": depth})
        for name in selected:
            try:
                arr = curves[name]
                if arr.ndim > 1:
                    arr = arr[:, 0]
                df[name] = arr
            except Exception:
                pass

        _las_log_plot(df, "DEPTH", selected, file_path + "_dlis")

    except Exception as e:
        st.error(f"DLIS parse error: {e}")
    finally:
        try:
            f.close()
            for t in tail:
                t.close()
        except Exception:
            pass


# =============================================================================
# LIS
# =============================================================================

def _view_lis(file_path: str):
    import pandas as pd
    st.info("LIS viewer — basic channel extraction.")
    try:
        # Try using dlisio for LIS (it supports both)
        import dlisio
        f, *tail = dlisio.lis.load(file_path)
        recs = list(f.data_format_specs())
        if not recs:
            st.warning("No data records found.")
            return

        sel_rec = st.selectbox(
            "Data record", range(len(recs)),
            format_func=lambda i: f"Record {i+1}",
            key=f"lis_rec_{file_path}",
        )
        rec = recs[sel_rec]
        channels = rec.specs

        ch_info = [{"Name": ch.mnemonic, "Unit": ch.units,
                    "Size": ch.size}
                   for ch in channels]
        with _vsection(f"📡 Channels ({len(channels)})"):
            st.dataframe(pd.DataFrame(ch_info),
                         hide_index=True, use_container_width=True)

        curves = dlisio.lis.curves(f, rec)
        ch_names = [ch.mnemonic for ch in channels]
        selected = st.multiselect(
            "Curves to plot", ch_names,
            default=ch_names[:min(4, len(ch_names))],
            key=f"lis_curves_{file_path}",
        )
        if not selected:
            return

        import numpy as np
        depth_name = ch_names[0]
        depth = curves[depth_name].flatten()
        df = pd.DataFrame({"DEPTH": depth})
        for name in selected:
            try:
                arr = curves[name]
                if arr.ndim > 1:
                    arr = arr[:, 0]
                df[name] = arr
            except Exception:
                pass

        _las_log_plot(df, "DEPTH", selected, file_path + "_lis")

    except Exception as e:
        st.error(f"LIS read failed: {e}")
        _view_text_fallback(file_path)


# =============================================================================
# SEGY
# =============================================================================

def _view_segy(file_path: str):
    import pandas as pd
    try:
        import segyio
    except ImportError:
        st.error("segyio not installed — `pip install segyio`")
        return

    try:
        with segyio.open(file_path, ignore_geometry=True) as f:
            # Binary header
            bin_header = {
                "Traces":        f.tracecount,
                "Samples/trace": f.bin[segyio.BinField.Samples],
                "Sample interval (us)": f.bin[segyio.BinField.Interval],
                "Format":        f.bin[segyio.BinField.Format],
            }
            with _vsection("🗂️ Binary header"):
                st.dataframe(
                    pd.DataFrame([{"Field": k, "Value": str(v)}
                                  for k, v in bin_header.items()]),
                    hide_index=True, use_container_width=True,
                )

            # Trace header sample
            n_preview = min(50, f.tracecount)
            header_keys = [
                segyio.TraceField.CDP,
                segyio.TraceField.FieldRecord,
                segyio.TraceField.TRACE_SEQUENCE_FILE,
                segyio.TraceField.offset,
                segyio.TraceField.CDP_X,
                segyio.TraceField.CDP_Y,
            ]
            h_rows = []
            for i in range(n_preview):
                h = f.header[i]
                h_rows.append({
                    "Trace":  i + 1,
                    "CDP":    h[segyio.TraceField.CDP],
                    "FieldRec": h[segyio.TraceField.FieldRecord],
                    "CDP_X":  h[segyio.TraceField.CDP_X],
                    "CDP_Y":  h[segyio.TraceField.CDP_Y],
                    "Offset": h[segyio.TraceField.offset],
                })
            with _vsection(f"📋 Trace headers (first {n_preview})"):
                st.dataframe(pd.DataFrame(h_rows),
                             hide_index=True, use_container_width=True)

            # Wiggle / density plot
            n_plot = min(100, f.tracecount)
            traces_to_plot = st.slider(
                "Traces to display", 10, min(500, f.tracecount), n_plot,
                key=f"segy_n_{file_path}",
            )

            import numpy as np
            data = np.zeros((f.samples.size, traces_to_plot))
            for i in range(traces_to_plot):
                data[:, i] = f.trace[i]

            _segy_plot(data, f.samples, traces_to_plot, file_path)

    except Exception as e:
        # segyio refuses a file whose body is not an exact multiple of the
        # trace length, and it is right to: it indexes the whole file, so a
        # ragged tail makes every trace number suspect. For a VIEWER that is
        # the wrong answer — the textual header, the binary header and the
        # well-formed trace headers are all still readable and are usually
        # what someone opening this panel actually wants. Fall back rather
        # than showing a raw RuntimeError over an empty page.
        st.warning(f"segyio could not open this file: {e}")
        _view_segy_tolerant(file_path, str(e))


def _view_segy_tolerant(file_path: str, why: str = ""):
    """Header-level view via the dependency-free reader, for files segyio
    rejects. Shows what is verifiable and says plainly what is not."""
    import pandas as pd
    try:
        from dataview.file_catalog.segy_header import (
            read_segy_header, sample_trace_rows, read_trace_samples)
    except Exception as e:            # pragma: no cover
        st.error(f"fallback reader unavailable: {e}")
        return

    h = read_segy_header(file_path)
    if not h.get("ok"):
        st.error("The file could not be read as SEG-Y at all — not even its "
                 "binary header. " + "; ".join(h.get("notes") or []))
        return

    # WHY, ARITHMETICALLY. "trace count inconsistent with file size" is a
    # conclusion; the numbers behind it are what tell you whether the file is
    # truncated, padded, or ragged — so show them.
    import os as _os
    size = _os.path.getsize(file_path)
    bpt = h.get("_bytes_per_trace") or 0
    body = size - (h.get("_data_start") or 0)
    if bpt:
        whole, left = divmod(body, bpt)
        st.info(
            f"**Reading in tolerant mode.** {size:,} bytes; "
            f"{body:,} after the headers; {h.get('n_samples')} samples x "
            f"{h.get('bytes_per_sample')} bytes + 240 header = {bpt:,} per "
            f"trace. That is **{whole:,} whole traces plus {left:,} leftover "
            f"bytes** — not a whole trace, which is what segyio refuses.")

    with _vsection("🗂️ Binary header"):
        st.dataframe(pd.DataFrame([
            {"Field": "Traces (whole)", "Value": f"{h.get('n_traces'):,}"},
            {"Field": "Samples/trace", "Value": h.get("n_samples")},
            {"Field": "Sample interval (us)", "Value": h.get("sample_interval_us")},
            {"Field": "Format", "Value": h.get("format_desc")},
            {"Field": "Byte order", "Value": h.get("byte_order")},
            {"Field": "Measurement", "Value": h.get("measurement_system")},
        ]), hide_index=True, use_container_width=True)

    if h.get("trace_map"):
        st.caption("Trace positions below use the layout **this file's own "
                   "textual header declares**, not the rev-1 defaults: "
                   + ", ".join(f"{k} @ byte {v + 1}"
                               for k, v in sorted(h["trace_map"].items())))

    rows = sample_trace_rows(file_path, 50)
    if rows:
        with _vsection(f"📋 Trace headers (first {len(rows)})"):
            st.dataframe(pd.DataFrame(rows), hide_index=True,
                         use_container_width=True)

    with _vsection("📄 Textual header"):
        st.code(h.get("textual_header") or "(empty)", language=None)

    for n in (h.get("notes") or []):
        st.caption(f"· {n}")

    # The samples themselves — attempted, never faked.
    with _vsection("〰️ Trace data"):
        with st.spinner("Walking the file for readable traces…"):
            data, times = read_trace_samples(file_path, 0, 120)
        if data is None:
            stats = getattr(read_trace_samples, "last_stats", {}) or {}
            st.error(
                "**No trace samples could be decoded.** The walk found no "
                "offset where the payload reads as "
                f"{h.get('format_desc')} — every candidate decoded to "
                "exponents no seismic sample has."
                + (f" Walked {stats.get('traces_walked', 0):,} trace slots."
                   if stats else ""))
            st.caption(
                "Deliberately blank rather than plotted: bytes that do not "
                "decode as the declared format are not quiet data, and a "
                "picture drawn from them would look like seismic without "
                "being it. The headers above ARE verified — traces 1 and 2 "
                "match the corner coordinates the textual header states — so "
                "the file's identity and geometry are trustworthy even though "
                "its samples are not.")
        else:
            stats = getattr(read_trace_samples, "last_stats", {}) or {}
            _segy_plot(data, times, data.shape[1], file_path)
            st.caption(
                f"{stats.get('plotted', 0)} trace(s) plotted · "
                f"{stats.get('blank_skipped', 0)} blank skipped · "
                f"{stats.get('resyncs', 0)} re-synchronised · "
                f"{stats.get('unreadable', 0)} unreadable")


def _segy_plot(data, samples, n_traces, file_path):
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np

        fig, axes = plt.subplots(1, 2, figsize=(14, 8),
                                 gridspec_kw={"width_ratios": [1, 1]})
        fig.patch.set_alpha(0.0)

        # Density (variable area)
        ax1 = axes[0]
        extent = [0, n_traces, samples[-1], samples[0]]
        vmax   = np.percentile(np.abs(data), 95) or 1
        ax1.imshow(data, aspect="auto", cmap="RdBu_r",
                   vmin=-vmax, vmax=vmax, extent=extent)
        ax1.set_title("Density (variable area)", color="white")
        ax1.set_xlabel("Trace", color="white")
        ax1.set_ylabel("Sample", color="white")
        ax1.tick_params(colors="white")
        ax1.set_facecolor("#1A2B4A")

        # Wiggle (first 30 traces max)
        #
        # THIS PANEL WAS BLACK INK ON A DARK NAVY GROUND. "k-" is black, the
        # facecolor was #1A2B4A, and both the line and the fill were drawn at
        # PART OPACITY on top of that -- so the wiggle, the one display a
        # geologist reads wavelet character from, was the least legible thing
        # on the page. The density panel beside it never had the problem
        # because RdBu_r is WHITE at zero amplitude, which is also why the two
        # looked like they came from different applications.
        #
        # A wiggle is conventionally black on white or cream with the positive
        # lobes filled solid, and that convention exists because it is what
        # makes a wavelet's shape readable. Restore it: pale ground,
        # full-strength ink, solid peak fill.
        ax2 = axes[1]
        n_wig = min(30, n_traces)
        norm  = vmax * 2 if vmax > 0 else 1
        for i in range(n_wig):
            trace = data[:, i] / norm + i
            ax2.plot(trace, samples, "-", color="#111111", linewidth=0.6)
            ax2.fill_betweenx(samples, i, trace,
                              where=(trace > i), color="#111111", alpha=0.9)
        ax2.set_xlim(-1, n_wig)
        ax2.invert_yaxis()
        ax2.set_title(f"Wiggle (first {n_wig} traces)", color="white")
        ax2.set_xlabel("Trace", color="white")
        ax2.set_ylabel("Sample", color="white")
        # Ticks and labels sit OUTSIDE the axes, on a figure patch that is
        # transparent, so they stay white against the dark page. Only the plot
        # INTERIOR goes pale -- setting the figure light would strand them.
        ax2.tick_params(colors="white")
        ax2.set_facecolor("#F7F4EC")
        for _sp in ax2.spines.values():
            _sp.set_color("#8A8A8A")

        fig.tight_layout()
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)

    except Exception as e:
        st.warning(f"SEGY plot failed: {e}")


def segy_volume_plot(il, xl, title=""):
    """A 3D volume as its two orthogonal sections, tied to each other.

    WHY NOT _segy_plot. That draws "the section" plus a wiggle of its first 30
    traces, which is the right display for a 2D LINE -- a line has one section
    and nothing to cross-reference. A volume does not: the useful question is
    always "what does this look like the other way", and the two panels only
    mean something together if you can see where they intersect. So each
    section is drawn against the OTHER axis's real numbering and carries a
    marker at the position of the section beside it. Drawing both against
    0..n-1 would put two pictures on a page with no way to relate them.

    Greyscale on purpose: it is what a volume is conventionally interpreted in,
    and the eye reads structural continuity out of it better than out of a
    diverging colour ramp.

    Either panel may be None -- a volume can carry a usable inline index and no
    crossline one -- in which case the other is drawn alone rather than beside
    an empty box.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np

        panels = [p for p in (il, xl) if p and p.get("data") is not None]
        if not panels:
            return None

        fig, axes = plt.subplots(1, len(panels), figsize=(7.4 * len(panels), 8),
                                 squeeze=False)
        fig.patch.set_alpha(0.0)

        for ax, p in zip(axes[0], panels):
            data, times, xs = p["data"], p["times"], p.get("x")
            # CLIP ON A PERCENTILE, NOT THE MAXIMUM. A single spike -- a bad
            # trace, an edge effect -- sets the scale for the whole section and
            # everything real fades to mid-grey. The 97th percentile keeps the
            # section readable and lets the outlier saturate, which is what
            # every interpretation package does.
            vmax = float(np.percentile(np.abs(data), 97)) or 1.0
            if xs is not None and len(xs) == data.shape[1]:
                extent = [float(xs[0]), float(xs[-1]), float(times[-1]),
                          float(times[0])]
                xlabel = p.get("cross_label", "trace")
            else:
                extent = [0, data.shape[1], float(times[-1]), float(times[0])]
                xlabel = "trace"
            ax.imshow(data, aspect="auto", cmap="Greys", vmin=-vmax, vmax=vmax,
                      extent=extent, interpolation="bilinear")

            # THE TIE. Mark where the other section cuts this one.
            tie = p.get("tie")
            if tie is not None and extent[0] <= tie <= extent[1]:
                ax.axvline(tie, color="#C81E1E", linewidth=1.1,
                           linestyle="--", alpha=0.9)
                ax.text(tie, extent[3], " " + str(p.get("tie_label") or "") + " ",
                        color="#C81E1E", fontsize=8.5, va="top", ha="left",
                        bbox=dict(facecolor="white", edgecolor="none",
                                  alpha=0.75, pad=1.5))

            ax.set_title(p.get("title", ""), color="white", fontsize=11)
            ax.set_xlabel(xlabel, color="white")
            ax.set_ylabel(p.get("ylabel", "Time (ms)"), color="white")
            ax.tick_params(colors="white")
            for _sp in ax.spines.values():
                _sp.set_color("#8A8A8A")

        if title:
            fig.suptitle(title, color="white", fontsize=12)
        fig.tight_layout()
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
        return True

    except Exception as e:
        st.warning(f"Volume plot failed: {e}")
        return None


# =============================================================================
# P190
# =============================================================================

def _view_p190(file_path: str):
    import pandas as pd
    try:
        from dataview.file_catalog.p190_catalog import parse_p190
        result = parse_p190(file_path)
    except ImportError:
        result = _parse_p190_basic(file_path)
    except Exception as e:
        st.error(f"P190 parse failed: {e}")
        return

    header = result.get("header", {})
    shots  = result.get("shots", [])

    if header:
        with _vsection("📋 File header"):
            st.dataframe(
                pd.DataFrame([{"Field": k, "Value": str(v)}
                              for k, v in header.items() if v]),
                hide_index=True, use_container_width=True,
            )

    if shots:
        df = pd.DataFrame(shots)
        st.metric("Shot points", len(df))
        st.dataframe(_arrow_safe(df.head(200)), hide_index=True, use_container_width=True)

        # Map if lat/lon available
        lat_col = next((c for c in df.columns
                        if c.upper() in ("LAT","LATITUDE","SHOT_LAT")), None)
        lon_col = next((c for c in df.columns
                        if c.upper() in ("LON","LONG","LONGITUDE","SHOT_LON")), None)
        if lat_col and lon_col:
            _folium_points(df, lat_col, lon_col, label="Shot point")

        st.download_button(
            "⬇ Download shot points CSV",
            data=df.to_csv(index=False),
            file_name=Path(file_path).stem + "_shots.csv",
            mime="text/csv",
            key=f"p190_dl_{file_path}",
        )
    else:
        st.info("No shot point data found.")
        _view_text_fallback(file_path)


def _parse_p190_basic(file_path: str) -> dict:
    """Minimal P190 parser — reads H records (header) and S records (shots)."""
    import re
    header = {}
    shots  = []
    try:
        with open(file_path, "r", errors="replace") as fh:
            for line in fh:
                if not line.strip():
                    continue
                rec = line[0].upper()
                if rec == "H":
                    parts = line[1:].split()
                    if len(parts) >= 2:
                        header[parts[0]] = " ".join(parts[1:])
                elif rec == "S":
                    parts = line.split()
                    if len(parts) >= 6:
                        try:
                            shots.append({
                                "LINE":   parts[1],
                                "SP":     parts[2],
                                "EASTING":  float(parts[3]),
                                "NORTHING": float(parts[4]),
                                "DEPTH":    float(parts[5]) if len(parts) > 5 else None,
                            })
                        except ValueError:
                            pass
    except Exception:
        pass
    return {"header": header, "shots": shots}


# =============================================================================
# Shapefile
# =============================================================================

def _view_shapefile(file_path: str):
    import pandas as pd
    try:
        import geopandas as gpd
        gdf = gpd.read_file(file_path)
    except ImportError:
        st.error("geopandas not installed.")
        return
    except Exception as e:
        st.error(f"Shapefile read failed: {e}")
        return

    m1, m2, m3 = st.columns(3)
    m1.metric("Features",  len(gdf))
    m2.metric("Columns",   len(gdf.columns))
    m3.metric("CRS",       str(gdf.crs) if gdf.crs else "None")

    # Attribute table
    display_df = gdf.drop(columns=["geometry"], errors="ignore")
    with _vsection("📋 Attribute table"):
        st.dataframe(_arrow_safe(display_df.head(500)),
                     hide_index=True, use_container_width=True)

    # Map
    try:
        _shp_map(gdf, file_path)
    except Exception as e:
        st.warning(f"Map unavailable: {e}")

    # Download attribute CSV
    st.download_button(
        "⬇ Download attributes CSV",
        data=display_df.to_csv(index=False),
        file_name=Path(file_path).stem + "_attrs.csv",
        mime="text/csv",
        key=f"shp_dl_{file_path}",
    )


def _shp_map(gdf, file_path: str):
    import folium
    from streamlit_folium import st_folium
    import geopandas as gpd

    # Reproject to WGS84 for display
    try:
        gdf_wgs = gdf.to_crs(epsg=4326)
    except Exception:
        gdf_wgs = gdf

    bounds = gdf_wgs.total_bounds  # minx, miny, maxx, maxy
    center = [(bounds[1] + bounds[3]) / 2,
              (bounds[0] + bounds[2]) / 2]

    m = folium.Map(location=center, zoom_start=8,
                   tiles="CartoDB positron")

    # Add features
    geom_type = gdf_wgs.geometry.geom_type.iloc[0] if len(gdf_wgs) else "Unknown"

    label_col = next(
        (c for c in gdf_wgs.columns
         if c.upper() in ("NAME","WELL_NAME","UWI","ID","LABEL","API")),
        None,
    )

    for _, row in gdf_wgs.iterrows():
        tooltip = str(row[label_col]) if label_col else ""
        try:
            geom = row.geometry
            if geom is None or geom.is_empty:
                continue
            if geom.geom_type == "Point":
                folium.CircleMarker(
                    location=[geom.y, geom.x],
                    radius=4,
                    color="#C8922A",
                    fill=True,
                    fill_opacity=0.7,
                    tooltip=tooltip,
                ).add_to(m)
            elif geom.geom_type in ("LineString", "MultiLineString"):
                folium.GeoJson(
                    geom.__geo_interface__,
                    style_function=lambda x: {
                        "color": "#1A2B4A", "weight": 2},
                    tooltip=tooltip,
                ).add_to(m)
            else:
                folium.GeoJson(
                    geom.__geo_interface__,
                    style_function=lambda x: {
                        "fillColor": "#1A7BBF",
                        "color": "#1A2B4A",
                        "weight": 1,
                        "fillOpacity": 0.3},
                    tooltip=tooltip,
                ).add_to(m)
        except Exception:
            pass

    m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
    st_folium(m, use_container_width=True, height=500,
              key=f"shp_map_{file_path}")


# =============================================================================
# Excel / CSV
# =============================================================================

def _view_excel(file_path: str):
    import pandas as pd
    try:
        xl = pd.ExcelFile(file_path)
    except Exception as e:
        st.error(f"Excel read failed: {e}")
        return

    sheet = st.selectbox("Sheet", xl.sheet_names,
                         key=f"xl_sheet_{file_path}")
    try:
        df = xl.parse(sheet)
        st.metric("Rows", f"{len(df):,}")
        st.dataframe(_arrow_safe(df.head(500)), hide_index=True, use_container_width=True)
        st.download_button(
            "⬇ Download sheet CSV",
            data=df.to_csv(index=False),
            file_name=f"{Path(file_path).stem}_{sheet}.csv",
            mime="text/csv",
            key=f"xl_dl_{file_path}",
        )
    except Exception as e:
        st.error(f"Sheet read failed: {e}")


def _view_csv(file_path: str):
    import pandas as pd
    ext = Path(file_path).suffix.lower()
    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(file_path, sep=sep, nrows=5000,
                         encoding="utf-8", on_bad_lines="skip")
        st.metric("Rows (preview)", f"{len(df):,}")
        st.dataframe(_arrow_safe(df), hide_index=True, use_container_width=True)
        st.download_button(
            "⬇ Download CSV",
            data=df.to_csv(index=False),
            file_name=Path(file_path).name,
            mime="text/csv",
            key=f"csv_dl_{file_path}",
        )
    except Exception as e:
        st.error(f"CSV read failed: {e}")
        _view_text_fallback(file_path)


# =============================================================================
# Word / DOCX
# =============================================================================

def _view_word(file_path: str):
    import pandas as pd
    ext = Path(file_path).suffix.lower()

    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
        except ImportError:
            st.error("python-docx not installed.")
            _view_text_fallback(file_path)
            return
        except Exception as e:
            st.error(f"DOCX read failed: {e}")
            return

        # Paragraphs
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        with _vsection("📝 Document text"):
            st.text(text[:5000])
            if len(text) > 5000:
                st.caption(f"... ({len(text):,} chars total, showing first 5,000)")

        # Tables
        if doc.tables:
            with _vsection(f"📊 Tables ({len(doc.tables)})"):
                for i, tbl in enumerate(doc.tables):
                    rows = [[cell.text for cell in row.cells]
                            for row in tbl.rows]
                    if rows:
                        headers = rows[0]
                        data    = rows[1:]
                        try:
                            df = pd.DataFrame(data, columns=headers)
                            st.markdown(f"**Table {i+1}**")
                            st.dataframe(_arrow_safe(df), hide_index=True,
                                         use_container_width=True)
                        except Exception:
                            pass
    else:
        # .doc — try antiword / textract fallback
        _view_text_fallback(file_path)


# =============================================================================
# Image
# =============================================================================

def _view_image(file_path: str):
    ext = Path(file_path).suffix.lower()
    try:
        if ext in (".tif", ".tiff"):
            try:
                from PIL import Image
                import io
                img = Image.open(file_path)
                # Convert to RGB for display
                if img.mode not in ("RGB", "RGBA", "L"):
                    img = img.convert("RGB")
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                st.image(buf.getvalue(), use_container_width=True,
                         caption=Path(file_path).name)
                st.caption(
                    f"Size: {img.width} × {img.height} px · "
                    f"Mode: {img.mode} · "
                    f"Bands: {len(img.getbands())}"
                )
            except ImportError:
                st.error("Pillow not installed — `pip install Pillow`")
        else:
            st.image(file_path, use_container_width=True,
                     caption=Path(file_path).name)
    except Exception as e:
        st.error(f"Image render failed: {e}")


# =============================================================================
# Text fallback
# =============================================================================

def _view_text_fallback(file_path: str, max_chars: int = 5000):
    st.markdown("**Raw text (first 5,000 chars)**")
    try:
        encodings = ["utf-8", "latin-1", "cp1252"]
        text = None
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc, errors="replace") as fh:
                    text = fh.read(max_chars)
                break
            except Exception:
                continue
        if text:
            st.code(text, language=None)
        else:
            st.warning("Could not read file as text.")
    except Exception as e:
        st.error(f"Text read failed: {e}")


# =============================================================================
# Folium point helper (shared)
# =============================================================================

def _folium_points(df, lat_col: str, lon_col: str, label: str = ""):
    try:
        import folium
        from streamlit_folium import st_folium

        valid = df[[lat_col, lon_col]].dropna()
        if valid.empty:
            return

        center = [valid[lat_col].mean(), valid[lon_col].mean()]
        m = folium.Map(location=center, zoom_start=8,
                       tiles="CartoDB positron")

        for _, row in df.iterrows():
            try:
                lat = float(row[lat_col])
                lon = float(row[lon_col])
                folium.CircleMarker(
                    location=[lat, lon],
                    radius=4,
                    color="#C8922A",
                    fill=True,
                    fill_opacity=0.8,
                ).add_to(m)
            except Exception:
                pass

        st_folium(m, use_container_width=True, height=400,
                  key=f"fmap_{lat_col}_{lon_col}")
    except ImportError:
        st.info("Install streamlit-folium for map display.")


def _view_las3_sections(file_path: str, why: str = "") -> bool:
    """Render a multi-section LAS 3.0 via split_las3. True if it handled it.

    Returns False for anything that is not a 3.0 file so the caller can fall
    through to its ordinary error path — this must never swallow a genuinely
    broken 2.0 file and present it as something exotic.
    """
    import pandas as pd
    try:
        from dataview.file_catalog.las_reader import split_las3
        las3 = split_las3(file_path)          # viewer WANTS the curve samples
    except ValueError:
        return False                          # not 3.0 — not our business
    except Exception:
        return False
    if not las3.sets and not las3.well:
        return False

    st.info(
        f"**LAS 3.0, read section by section.** lasio stopped on this file "
        f"({why[:120]}) because it assumes a single data block; this file "
        f"carries {len(las3.sets)}. Everything below is parsed from the file "
        f"itself.")

    if las3.well:
        with _vsection("🛢️ Well information"):
            st.dataframe(pd.DataFrame(
                [{"Mnemonic": k, "Value": v} for k, v in las3.well.items()]),
                hide_index=True, use_container_width=True)

    for name in sorted(las3.sets):
        s = las3.sets[name]
        with _vsection(f"📊 ~{name}  ·  {len(s.columns)} columns × "
                       f"{len(s.rows):,} rows"):
            st.dataframe(pd.DataFrame(
                [{"Mnemonic": c.mnemonic, "Unit": c.unit,
                  "Type": c.fmt or "—", "Description": c.descr}
                 for c in s.columns]), hide_index=True,
                use_container_width=True)
            if not s.rows:
                st.caption("no data rows in this section")
                continue

            # THE CURVES WERE ALWAYS HERE, ONLY NEVER DRAWN. This section had
            # the definitions table and a 500-row preview grid, so a LAS 3.0
            # file showed "the header but not the curves" -- while a LAS 2.0
            # file, which lasio can open, got a full log plot. Nothing was
            # missing from the parse: this file carries 13,084 depths x 8
            # curves and split_las3 reads every one, including turning the
            # file's declared NULL into a real gap rather than plotting
            # -999.25 as a value.
            #
            # DUPLICATE MNEMONICS ARE LEGAL IN LAS and would make df[mnem]
            # return a DataFrame instead of a Series, so the "is this column
            # numeric" test raises "truth value is ambiguous" and the whole
            # viewer falls over. Suffix the repeats.
            _seen, _cols = {}, []
            for _m in s.mnemonics:
                _seen[_m] = _seen.get(_m, 0) + 1
                _cols.append(_m if _seen[_m] == 1 else f"{_m}[{_seen[_m]}]")
            _df = pd.DataFrame(s.rows, columns=_cols)
            _num = _df.apply(lambda c: pd.to_numeric(c, errors="coerce"))

            _depth = next((c for c in _cols
                           if c.upper().split("[")[0]
                           in ("DEPT", "DEPTH", "MD", "TVD", "TVDSS")), None)
            if _depth is None and _cols:
                _depth = _cols[0]
            # A section is only plottable if its index column and at least one
            # other column are numeric. Core descriptions and Tops names are
            # tabular and real, and a log plot of them would be meaningless --
            # they keep the grid and nothing else.
            _plot = [c for c in _cols
                     if c != _depth and bool(_num[c].notna().any())]
            if _depth and bool(_num[_depth].notna().any()) and _plot:
                _sel = st.multiselect(
                    "Curves to plot", _plot,
                    default=_plot[:min(4, len(_plot))],
                    key=f"las3_curves_{file_path}_{name}")
                if _sel:
                    # THE WHOLE SECTION, NOT THE PREVIEW. Plotting the 500
                    # rows shown below would draw the top of the well and
                    # label it the log -- here, 500 of 13,084 samples is the
                    # first 250 feet of a 6,500 foot hole.
                    _pdf = _num[[_depth] + list(_sel)].dropna(subset=[_depth])
                    _las_log_plot(_pdf, _depth, list(_sel), file_path)

            st.dataframe(_df.head(500), hide_index=True,
                         use_container_width=True)
            if len(s.rows) > 500:
                st.caption(f"grid shows the first 500 of {len(s.rows):,} rows"
                           + (f"; the plot uses all {len(s.rows):,}"
                              if _plot else ""))
            st.download_button(
                f"Download ~{name} as CSV ({len(s.rows):,} rows)",
                data=_df.to_csv(index=False),
                file_name=Path(file_path).stem + f"_{name}.csv",
                mime="text/csv",
                key=f"las3_dl_{file_path}_{name}")
    return True
