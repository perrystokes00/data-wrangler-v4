"""
worker_core.py — the streamlit-free per-file processor (THE KEYSTONE)
=====================================================================
One code path for processing a single file into the file_catalog.cat_* mirrors,
callable from BOTH:
  * the parallel worker pool (background processes draining the work queue), and
  * the interactive UI's "process this one file" button.

No Streamlit, no page_workbench import. Fully importable and picklable so it
runs inside a spawned worker process. Each worker builds its OWN SQLAlchemy
engine from (server, database) — engines can't cross the process boundary — and
calls process_file() per claimed file.

This is the foundation of the corporate-scale pipeline
(see CORPORATE_SCALE_ARCHITECTURE.md). Build order step 1.

CONTRACT
--------
    process_file(engine, file_rec, log=None) -> FileResult

    file_rec: dict with FILE_PATH, FILE_EXT, MATCHED_UWI (opt), INVENTORY_ID (opt),
              FILE_NAME (opt).
    Returns FileResult(status, rows_written, detail, error).
    MUST be idempotent — capture() replaces rows scoped to INVENTORY_ID, so a
    file re-claimed after a worker crash is reprocessed safely (replace, not
    duplicate).

DESIGN NOTES
------------
* capture(engine, ...) takes a SQLAlchemy engine and opens engine.begin()
  itself, and calls _ensure_fast_executemany(engine). So workers need a real
  engine, not a raw pyodbc connection.
* The parser modules (pdf_survey_catalog, pdf_db_loader, lasio, dlisio,
  lis_catalog, shapefile/office/json catalogers) are already streamlit-free.
  The ONLY Streamlit coupling in the legacy path was st.error() in exception
  handlers — replaced here by the `log` callback.
* Engine rebuild mirrors pipeline_run._engine() exactly (trusted_connection,
  driver braces) — the proven pattern that avoids the IM002 child-process bug.
"""
from __future__ import annotations

import os
import re
from dataclasses import dataclass, field

# The seismic header MERGE is IMPORTED, never copied. It was a verbatim
# duplicate of extract_core._SQL_SEIS_MERGE until 23 Aug, so the multicore
# path — which is the DEFAULT — silently kept the old behaviour every time
# the canonical one was fixed. This table had four writers; the escapechar
# bug came back through exactly that kind of fourth writer.
from dataview.file_catalog.extract_core import (
    _SQL_SEIS_MERGE, ensure_seis_columns)


# ── result contract ──────────────────────────────────────────────────────────
@dataclass
class FileResult:
    status: str                       # "done" | "error" | "skip"
    rows_written: int = 0
    detail: dict = field(default_factory=dict)
    error: str | None = None
    rt: str = ""                      # report type / format label


# ── per-worker engine (built once inside each worker process) ────────────────
def make_engine(server: str, database: str,
                driver: str = "ODBC Driver 17 for SQL Server"):
    """Build a SQLAlchemy engine for ONE worker process.

    Uses odbc_connect with a verbatim DSN — NOT the keyword URL form
    (mssql+pyodbc://@host\\instance/db?...). The keyword form mangles the
    backslash in a host\\instance server name and resolves the database
    inconsistently across connections, which surfaces as "Invalid column name
    PROC_STATUS" in some connections (they land in the wrong DB / default
    catalog) even though the column exists in the intended DB. odbc_connect
    passes SERVER=host\\instance;DATABASE=db exactly as written, so every
    connection from this engine is pinned to the same database.

    (pipeline_run._engine uses the keyword form and carries this same
    vulnerability in the child-process path — which is why _engine_spec exists
    to rebuild it there. We sidestep it entirely with odbc_connect.)

    NullPool: a worker uses one connection at a time; avoids any pooled
    connection landing in a different context.
    """
    from sqlalchemy import create_engine
    from sqlalchemy.pool import NullPool
    import urllib.parse as _up
    odbc = (f"DRIVER={{{driver}}};SERVER={server};DATABASE={database};"
            f"Trusted_Connection=yes;")
    return create_engine(
        "mssql+pyodbc:///?odbc_connect=" + _up.quote_plus(odbc),
        fast_executemany=True, poolclass=NullPool)


# ── the keystone ─────────────────────────────────────────────────────────────
def process_file(engine, file_rec: dict, log=None) -> FileResult:
    """Process ONE file end to end into the cat_* mirrors. Streamlit-free.

    Dispatches by extension to the existing streamlit-free parsers + capture().
    Idempotent per file (capture replaces by INVENTORY_ID).
    """
    say = log or (lambda *_a, **_k: None)

    fpath = file_rec.get("FILE_PATH") or file_rec.get("path")
    fext  = (file_rec.get("FILE_EXT") or file_rec.get("ext")
             or os.path.splitext(fpath or "")[1] or "").lower()
    uwi   = (file_rec.get("MATCHED_UWI") or file_rec.get("uwi") or "")
    uwi   = str(uwi).strip() or None
    inv   = file_rec.get("INVENTORY_ID") or file_rec.get("inventory_id")
    fname = file_rec.get("FILE_NAME") or os.path.basename(fpath or "")

    if not fpath or not os.path.exists(fpath):
        return FileResult("error", error=f"file not found: {fpath}")

    res = None
    try:
        if fext == ".las":
            res = _do_las(engine, fpath, uwi, inv, say)
        elif fext in (".dlis", ".dlf", ".dis"):
            res = _do_dlis(engine, fpath, uwi, inv, fname, say)
        elif fext == ".lis":
            res = _do_lis(engine, fpath, uwi, inv, fname, say)
        elif fext == ".pdf":
            res = _do_pdf(engine, fpath, uwi, inv, say)
        elif fext in (".shp",):
            res = _do_shapefile(engine, fpath, uwi, inv, say)
        elif fext in (".xml",):
            res = _do_witsml(engine, fpath, uwi, inv, say)
        elif fext in (".xlsx", ".xls", ".docx", ".doc", ".odt", ".ods", ".csv"):
            res = _do_office(engine, fpath, uwi, inv, say)
        elif fext == ".json":
            res = _do_json(engine, fpath, uwi, inv, say)
        elif fext in (".segy", ".sgy", ".seg", ".p190"):
            res = _do_segy(engine, fpath, inv, say)
        else:
            return FileResult("skip", error=f"no handler for {fext}")
    except Exception as e:                       # noqa: BLE001
        return FileResult("error", error=f"{type(e).__name__}: {e}")

    # Stamp this file's resolved identity onto its GLOBAL_FILE_CATALOG row so the
    # catalog is directly queryable: WHERE UWI14=:x returns EVERY document for a
    # well (LAS + scout + completion + dir-survey), WHERE SURVEY_NAME=:y every doc
    # for a survey. Identity is read from what THIS file just wrote to staging
    # (cat_well / FILE_SEIS_HEADER), scoped by INVENTORY_ID. Best-effort: a
    # tagging failure never fails the file.
    if inv is not None and res is not None and getattr(res, "status", None) == "done":
        try:
            _tag_catalog_from_stage(engine, inv, say)
        except Exception as e:                   # noqa: BLE001
            say(f"catalog identity tag: {str(e)[:200]}")
    return res


# ── catalog identity tagging ─────────────────────────────────────────────────
# Give GLOBAL_FILE_CATALOG each document's resolved well/survey key so the
# catalog is queryable on its own: one catalog row per file, so the tag is
# single-valued per document. Multi-well files (e.g. a well-point shapefile)
# resolve to >1 uwi and are left to FILE_WELL_HEADER. UWI14 / SURVEY_NAME columns
# are added by add_catalog_identity_cols.py. Everything here is best-effort.

_WELL_STAGE = None  # cached [(schema, table, uwi_col), ...] of well-bearing cat_*


def _uwi14(uwi):
    """Digits-only canonical 14-char UWI, or None if not exactly 14 digits."""
    if uwi is None:
        return None
    d = re.sub(r"\D", "", str(uwi))
    return d if len(d) == 14 else None


def _well_stage_tables(engine):
    """Discover, once, every cat_* staging table that carries BOTH a UWI-ish
    column and INVENTORY_ID — cat_well (LAS/PDF) plus the office mirrors
    cat_well_completion / cat_well_formation_top / cat_prod_entity, etc. The
    tagger unions the UWI across all of them, so office-derived wells (which
    never touch cat_well) are covered without coupling to that loader, and any
    future mirror is picked up automatically."""
    global _WELL_STAGE
    if _WELL_STAGE is not None:
        return _WELL_STAGE
    from sqlalchemy import text as _t
    found = []
    try:
        with engine.connect() as con:
            rows = con.execute(_t("""
                SELECT c.TABLE_SCHEMA, c.TABLE_NAME, c.COLUMN_NAME
                FROM INFORMATION_SCHEMA.COLUMNS c
                WHERE c.TABLE_NAME LIKE 'cat[_]%'
                  AND c.COLUMN_NAME IN ('UWI','MATCHED_UWI','UWI14')
                  AND EXISTS (SELECT 1 FROM INFORMATION_SCHEMA.COLUMNS i
                              WHERE i.TABLE_SCHEMA = c.TABLE_SCHEMA
                                AND i.TABLE_NAME  = c.TABLE_NAME
                                AND i.COLUMN_NAME = 'INVENTORY_ID')""")).fetchall()
            best = {}  # prefer UWI over MATCHED_UWI/UWI14 when several exist
            for sch, tbl, col in rows:
                key = (sch, tbl)
                if key not in best or col.upper() == "UWI":
                    best[key] = col
            found = [(s, t, c) for (s, t), c in best.items()]
    except Exception:
        pass
    _WELL_STAGE = found
    return found


def _tag_catalog(engine, inv, uwi14=None, survey_name=None, say=None):
    """Set-based UPDATE of the GLOBAL_FILE_CATALOG row (keyed by INVENTORY_ID)
    with the resolved identity. Best-effort — never raises to the caller."""
    from sqlalchemy import text as _t
    if inv is None or (uwi14 is None and survey_name is None):
        return 0
    sets, params = [], {"inv": inv}
    if uwi14 is not None:
        sets.append("UWI14 = :u")
        params["u"] = uwi14
    if survey_name is not None:
        sets.append("SURVEY_NAME = :s")
        params["s"] = str(survey_name)[:255]
    sql = ("UPDATE file_catalog.GLOBAL_FILE_CATALOG SET " + ", ".join(sets) +
           " WHERE INVENTORY_ID = :inv")
    try:
        with engine.begin() as con:
            return con.execute(_t(sql), params).rowcount or 0
    except Exception as e:                        # noqa: BLE001
        if say:
            say(f"catalog tag: {str(e)[:200]}")
        return 0


def _tag_catalog_from_stage(engine, inv, say=None):
    """Resolve this file's identity from staging and stamp the catalog. Collect
    the DISTINCT UWI this file wrote across every well-bearing cat_* table
    (cat_well + office mirrors), scoped by INVENTORY_ID; tag UWI14 only when it
    resolves to exactly ONE well (multi-well files → left to FILE_WELL_HEADER).
    Survey tag comes from the seismic header."""
    from sqlalchemy import text as _t
    uwi14 = survey = None
    uwis = set()
    with engine.connect() as con:
        for sch, tbl, ucol in _well_stage_tables(engine):
            try:
                rows = con.execute(_t(f"""
                    SELECT DISTINCT [{ucol}] FROM {sch}.{tbl}
                    WHERE INVENTORY_ID = :inv AND [{ucol}] IS NOT NULL
                      AND LTRIM(RTRIM(CONVERT(varchar(64), [{ucol}]))) <> ''"""),
                    {"inv": inv}).fetchall()
                for r in rows:
                    n = _uwi14(r[0])
                    if n:
                        uwis.add(n)
            except Exception as e:                # noqa: BLE001
                if say:
                    say(f"{tbl} identity: {str(e)[:100]}")
        if len(uwis) == 1:
            uwi14 = next(iter(uwis))
        try:
            srow = con.execute(_t("""
                SELECT TOP 1 SURVEY_NAME FROM file_catalog.FILE_SEIS_HEADER
                WHERE INVENTORY_ID = :inv AND SURVEY_NAME IS NOT NULL"""),
                {"inv": inv}).fetchone()
            if srow:
                survey = srow[0]
        except Exception as e:                    # noqa: BLE001
            if say:
                say(f"seis identity: {str(e)[:120]}")
    if uwi14 or survey:
        _tag_catalog(engine, inv, uwi14=uwi14, survey_name=survey, say=say)


# ── format handlers — thin wrappers over existing streamlit-free logic ───────
# All handlers are ported from page_workbench._load_rows_to_catalog: PDF, LAS,
# DLIS, LIS, shapefile, office, json (OSDU), witsml. Each returns a FileResult
# and uses capture() with the REAL schema (LOG_ID/MNEMONIC/TOP_DEPTH for logs,
# cat_well columns for headers, etc.), with st.error() replaced by the `say`
# callback so the path is fully streamlit-free and worker-safe.

def _capture():
    """Import capture() the same way the rest of the app does (modules.* or flat)."""
    try:
        from dataview.file_catalog.catalog_capture import capture, reset_replace_state
    except ImportError:
        from dataview.file_catalog.catalog_capture import capture, reset_replace_state
    return capture, reset_replace_state


def _wellname_from_filename(fname, fpath=None):
    """Extract the WELL NAME *substring* from a binary-log filename, stripping
    vendor prefixes, run/tool/interval metadata, and format qualifiers — so the
    result is a clean name triage can match to dv_well / gold.

    Handles the conventions seen in real data:
      ANADARKO_BURK_145            → 'ANADARKO BURK 145'   (name as-is)
      a0501DEVON_EN_BURK_188t01    → 'DEVON EN BURK 188'   (strip a0501.. t01)
      a0501t01                     → None                  (wrapper only, no name)
      A12-A-08_Run4_8375in_RM_...  → 'A12-A-08'            (cut at _Run/_8375in)
      Chevron_A12a-CPP-A6_Combined → 'Chevron A12a-CPP-A6' (strip _Combined)
      A151_Composite-TOTAL         → 'A151'                (strip _Composite..)
      G030088972__A-5-1__REFS1234  → 'A-5-1'               (middle of __ split)

    Returns a name or None. Heuristic by design — triage decides what matches.
    """
    base = os.path.splitext(os.path.basename(fname or fpath or ""))[0]
    if not base:
        return None

    # Case: double-underscore delimited <fileid>__<well>__<ref> → take middle.
    if "__" in base:
        parts = [p for p in base.split("__") if p]
        if len(parts) >= 2:
            # middle token is usually the well; skip pure-numeric file ids
            for p in parts:
                if re.search(r"[A-Za-z]", p) and not re.fullmatch(r"[A-Z]?\d{6,}", p) \
                        and not p.upper().startswith("REFS"):
                    base = p
                    break

    # Strip a leading vendor wrapper like 'a0501' ONLY when it's glued directly
    # to more name (letters follow with no separator), e.g. 'a0501DEVON...'.
    # A standalone token like 'A151' or 'A151_...' is the name itself — keep it.
    base = re.sub(r"^[a-zA-Z]\d{3,}(?=[A-Za-z])", "", base)
    # Strip a trailing wrapper like 't01' (letter+digits) glued to the end.
    base = re.sub(r"t\d{1,3}$", "", base)

    # Truncate at the first run/tool/interval/format marker (with its separator).
    base = re.split(
        r"[_\-\s]+(?:run\d*|\d{2,}in|rm|combined|composite|merged|final|raw|"
        r"edited|welllog|logs?|data|total|\d+-\d+m)(?:[_\-\s]|$)",
        base, maxsplit=1, flags=re.IGNORECASE)[0]

    # Normalize separators to spaces.
    name = re.sub(r"[_]+", " ", base)
    name = re.sub(r"\s+", " ", name).strip(" -_")

    # Must contain letters AND be more than a stray fragment to be a real name.
    if not name or not re.search(r"[A-Za-z]", name) or len(name) < 3:
        return None
    return name


def _uwi_from_filename(fname, fpath=None):
    """Pull a 14-digit UWI from a filename/path, or None. DLIS/LIS/LAS files are
    commonly named like 42999000020000_welllog.dlis, ..._42-330-00035-00-00, or
    42_475_10072_0000.las (underscore-delimited API). 10-digit APIs are
    zero-padded to 14. Lets binary logs that lack an embedded UWI still be keyed
    by their real well, so triage/promote match normally."""
    base = os.path.basename(fname or fpath or "")
    stem = os.path.splitext(base)[0]
    for _rx in (r"(\d{2}-\d{3}-\d{5}-\d{2}-\d{2})",    # dash-delimited API
                r"(\d{2}_\d{3}_\d{5}_\d{2}_\d{2})",    # underscore API (5-part)
                r"(\d{2}_\d{3}_\d{5}_\d{4})",          # underscore API (4-part)
                r"(?<!\d)(\d{14})(?!\d)",              # 14 contiguous digits
                r"(?<!\d)(\d{10})(?!\d)"):             # 10 contiguous digits
        m = re.search(_rx, base)
        if m:
            d = re.sub(r"\D", "", m.group(1))
            if len(d) >= 10:
                u = d[:14] if len(d) >= 14 else d.ljust(14, "0")
                if u != "0" * 14:
                    return u
    # Last resort: if the stem is ENTIRELY separator-delimited digit groups
    # (e.g. 42_475_10072_0000 or 35-049-10121), strip separators and see if it
    # yields a 10–14 digit API. Guarded to all-digits-and-separators so we don't
    # mangle arbitrary names.
    if re.fullmatch(r"[\d_\-]+", stem):
        d = re.sub(r"\D", "", stem)
        if 10 <= len(d) <= 14:
            u = d[:14] if len(d) >= 14 else d.ljust(14, "0")
            if u != "0" * 14:
                return u
    return None


def _valid_uwi(uwi):
    """UWI looks valid if its first 10 chars are digits forming a plausible
    API10 (2-digit state code 01-62). Rejects blanks, FN_ fallbacks, garbage."""
    if not uwi:
        return False
    d = "".join(ch for ch in str(uwi) if ch.isdigit())
    if len(d) < 10:
        return False
    try:
        return 1 <= int(d[:2]) <= 62      # API state codes: 01-50 states, 55-62 offshore
    except ValueError:
        return False


def _do_las(engine, fpath, uwi, inv, say) -> FileResult:
    """LAS → cat_well + cat_well_log + cat_well_log_curve (per-curve INVENTORY).

    Ported from page_workbench's `elif fext in LAS_EXTS` block — the canonical,
    promote-clean schema. One lasio.read; header to cat_well, log header to
    cat_well_log, one row PER CURVE to cat_well_log_curve, all sharing one
    LOG_ID so curves resolve their FK to dv_well_log on promote. No sample
    arrays retained (inventory grain).
    """
    import lasio, uuid as _uuid
    from dataview.file_catalog.las_reader import read_las
    from datetime import datetime as _dt
    capture, reset_replace_state = _capture()
    reset_replace_state()
    _now = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    res = FileResult("done", rt="WELL_LOG")

    las = read_las(fpath, ignore_data=True)   # header + curve defs only (skip samples)

    def _wv(*keys):
        for k in keys:
            try:
                v = str(las.well[k].value).strip()
                if v and v.lower() not in ("", "unknown", "none", "--"):
                    return v
            except Exception:
                pass
        return None

    def _safe_num(v):
        try:
            return float(str(v).strip())
        except (TypeError, ValueError):
            return None

    def _safe_coord(v):
        f = _safe_num(v)
        if f is None:
            return None
        return f if -180.0 <= f <= 180.0 else None

    # The LAS header carries the authoritative UWI in the ~WELL section
    # (UWI. 17-031-10035-0000 / API. ...). Read and normalize it — strip
    # separators, pad a 10-digit API to 14 — and PREFER it over whatever the
    # caller derived from the filename. Only if the header has no usable UWI do
    # we keep the passed-in value (filename-derived, possibly an FN_ fallback).
    def _norm_uwi(raw):
        if not raw:
            return None
        d = re.sub(r"\D", "", str(raw))
        if len(d) < 10:
            return None
        u = d[:14] if len(d) >= 14 else d.ljust(14, "0")
        return u if u != "0" * 14 else None

    _hdr_uwi = _norm_uwi(_wv("UWI", "API", "APINUM", "APINO", "API_NO", "WELLID"))
    if _hdr_uwi:
        uwi = _hdr_uwi      # header UWI wins — it's the real well identity

    # THE FILENAME, WHEN THE HEADER AND THE CALLER BOTH HAVE NOTHING.
    # _do_dlis already does this; LAS did not, so it depended on an
    # EARLIER STAGE having written MATCHED_UWI into the catalog. Run
    # capture before the matcher and the file skips with 'invalid UWI'
    # — 20 of the 185 Teapot LAS, exactly the 20 whose ~WELL section
    # carries an empty UWI field, which is a shape real LAS files have.
    # Deriving it here makes the stage order stop mattering.
    if not _valid_uwi(uwi):
        _fn_uwi = _norm_uwi(_uwi_from_filename(os.path.basename(fpath), fpath))
        if _fn_uwi:
            uwi = _fn_uwi
            say(f"uwi from filename: {uwi} ({os.path.basename(fpath)})")

    if not _valid_uwi(uwi):
        say(f"skip (invalid UWI {uwi!r}): {fpath}")
        res.status = "skip"
        res.detail["note"] = f"invalid UWI: {uwi!r}"
        return res

    # depth frame from the LAS index (no arrays retained)
    def _fnum(v):
        try:
            return float(str(v).strip())
        except Exception:
            return None
    d_start = _fnum(_wv("STRT", "START"))
    d_stop  = _fnum(_wv("STOP"))
    try:
        d_uom = (las.curves[0].unit or "").strip() or None
    except Exception:
        d_uom = None

    _logid = _wv("LOG_ID", "LOGID") or (f"{uwi}-LAS" if uwi else None)

    # one transaction; header shares _con (no nested begin)
    with engine.begin() as _con:

        # 1) well header → cat_well (promote builds dv_well)
        if uwi:
            try:
                n = capture(engine, "cat_well", [{
                    "WELL_NAME":         _wv("WELL") or uwi,
                    "OPERATOR_NAME":     _wv("COMP", "PROV"),
                    "FIELD_NAME":        _wv("FLD", "FIELD"),
                    "PROVINCE_STATE":    _wv("STAT", "STATE"),
                    "COUNTY":            _wv("CNTY", "COUNTY"),
                    "COUNTRY":           _wv("CTRY", "CTRY.", "COUNTRY"),
                    "SURFACE_LATITUDE":  _safe_coord(_wv("LATI", "LAT")),
                    "SURFACE_LONGITUDE": _safe_coord(_wv("LONG", "LON")),
                    "FINAL_TD":          _wv("STOP", "TD"),
                    "ACTIVE_IND":        "Y",
                    "ROW_QUALITY":       "FINAL",
                    "PPDM_GUID":         str(_uuid.uuid4()),
                    "ROW_CREATED_BY":    "DataWrangler",
                    "ROW_CREATED_DATE":  _now,
                }], uwi=uwi, inventory_id=inv, source_path=fpath,
                   source="LAS_HEADER", conn=_con)
                if n:
                    res.detail["cat_well"] = n
                # also populate FILE_WELL_HEADER (documents map / triage read this)
                _write_well_header(engine, inv, uwi, {
                    "well_name":  _wv("WELL") or uwi,
                    "operator":   _wv("COMP", "PROV"),
                    "field":      _wv("FLD", "FIELD"),
                    "state":      _wv("STAT", "STATE"),
                    "county":     _wv("CNTY", "COUNTY"),
                    "latitude":   _safe_coord(_wv("LATI", "LAT")),
                    "longitude":  _safe_coord(_wv("LONG", "LON")),
                    "total_depth": _wv("STOP", "TD"),
                }, "LAS", say, con=_con)
            except Exception as e:
                say(f"header capture: {e}")

        # 2) per-curve inventory → cat_well_log_curve
        curve_rows = []
        for c in las.curves:
            mnem = (getattr(c, "mnemonic", "") or "").strip()
            if not mnem:
                continue
            curve_rows.append({
                "UWI":               uwi or None,
                "LOG_ID":            _logid,
                "CURVE_ID":          mnem[:40],
                "MNEMONIC":          mnem,
                "CURVE_DESCRIPTION": (getattr(c, "descr", "") or "").strip() or None,
                "CURVE_UNIT":        (getattr(c, "unit", "") or "").strip() or None,
                "TOP_DEPTH":         _safe_num(d_start),
                "BASE_DEPTH":        _safe_num(d_stop),
                "DEPTH_OUOM":        d_uom,
                "NULL_VALUE":        _safe_num(_wv("NULL")),
                "ACTIVE_IND":        "Y",
                "ROW_CREATED_BY":    "DataWrangler",
                "ROW_CREATED_DATE":  _now,
            })
        if curve_rows:
            n = capture(engine, "cat_well_log_curve", curve_rows,
                        uwi=uwi, inventory_id=inv, source_path=fpath, source="LAS", conn=_con)
            res.rows_written += (n or 0)
            if n:
                res.detail["cat_well_log_curve"] = n

        # 3) log header → cat_well_log (shares _logid with the curves)
        if uwi:
            try:
                _srvc = _wv("SRVC", "SERVICE", "COMP")
                n = capture(engine, "cat_well_log", [{
                    "LOG_ID":                _logid,
                    "LOG_TYPE":              _wv("TYPE", "LOGTYPE"),
                    "RUN_NUM":               _wv("RUN", "RUN_NUMBER"),
                    "LOG_DATE":              _wv("DATE", "LOGDATE", "DATE_LOG"),
                    "SERVICE_COMPANY_BA_ID": None,
                    "TOP_DEPTH":             _safe_num(d_start),
                    "BASE_DEPTH":            _safe_num(d_stop),
                    "DEPTH_OUOM":            d_uom,
                    "NULL_VALUE":            _safe_num(_wv("NULL")),
                    "FILE_PATH":             fpath,
                    "FILE_FORMAT":           "LAS",
                    "ACTIVE_IND":            "Y",
                    "REMARK":                (f"service_company={_srvc}"
                                              if _srvc else None),
                    "ROW_CREATED_BY":        "DataWrangler",
                    "ROW_CREATED_DATE":      _now,
                }], uwi=uwi, inventory_id=inv, source_path=fpath, source="LAS", conn=_con)
                if n:
                    res.detail["cat_well_log"] = n
            except Exception as e:
                say(f"well_log capture: {e}")

    return res


# ── stubs for the remaining formats (next build steps port them in) ──────────
def _do_pdf(engine, fpath, uwi, inv, say, dialect="mssql") -> FileResult:
    """PDF capture, consolidated onto the Directory Loader's extractor.

    Delegates to catalog_doc_capture.capture_document, which runs
    pdf_document_loader.extract_file — self-resolving the UWI from the document
    itself (18 label spellings), so a scout ticket whose UWI triage could not
    read still captures — then maps its rows into the cat_* mirrors. Streamlit-
    free, so it is safe inside a pool worker. (The former pdf_survey_catalog /
    pdf_db_loader body is retired; _extract_pdf_rows / _load_directional below
    are now unused but left in place to keep this change minimal.)
    """
    from dataview.file_catalog.catalog_doc_capture import capture_document
    r = capture_document(engine, dialect, fpath, ".pdf", uwi, inv, log=say)
    n = int(r.get("loaded", 0) or 0)
    errs = r.get("errors") or []
    if n:
        status = "done"
    elif errs:
        status = "error"
    else:
        status = "skip"          # uwi_unresolved / no_rows -> nothing captured, not an error
    return FileResult(status, rows_written=n, detail=r.get("detail", {}),
                      error=(str(errs[0]) if errs else (r.get("note") or None)),
                      rt=r.get("rt", "PDF"))


def _extract_pdf_rows(fpath, rt, say):
    """Produce the detail rows for a PDF report type — the _do_extract step,
    ported streamlit-free. Returns a list (may be empty for header-only types)."""
    from dataview.file_catalog.pdf_survey_catalog import (
        extract_stations, extract_eowr, extract_rft_data, extract_well_test,
        extract_casing_cement, extract_ddr, extract_scout_ticket,
        RT_DIRECTIONAL, RT_EOWR, RT_RFT, RT_WELL_TEST, RT_CASING, RT_DDR, RT_SCOUT)
    try:
        from dataview.file_catalog.pdf_survey_catalog import RT_FORMATION
    except ImportError:
        RT_FORMATION = "FORMATION_TOPS"
    try:
        if rt == RT_DIRECTIONAL:
            return extract_stations(fpath).get("stations", [])
        if rt in (RT_EOWR, RT_FORMATION):
            return extract_eowr(fpath).get("strat", [])
        if rt == RT_RFT:
            return extract_rft_data(fpath).get("rows", [])
        if rt == RT_WELL_TEST:
            return extract_well_test(fpath).get("flow_rows", [])
        if rt == RT_CASING:
            r = extract_casing_cement(fpath)
            return (r.get("casing", []) or []) + (r.get("cement", []) or [])
        if rt == RT_DDR:
            return extract_ddr(fpath).get("ops", [])
        if rt == RT_SCOUT:
            # scout self-parses inside load_scout; no pre-extracted rows needed
            return []
    except Exception as e:
        say(f"pdf extract ({rt}): {e}")
    return []


# Canonical station keys emitted by pdf_survey_catalog._parse_station_row -> the
# cat_well_dir_srvy_sta column names the survey loader expects. Without this remap the
# stations load with NULL md/incl/azim (values parsed under MD/INC/AZI, loader reads
# md/incl/azim).
_SURVEY_KEY_MAP = {
    "MD": "md", "INC": "incl", "AZI": "azim", "TVD": "tvd",
    "NS": "ns_offset", "EW": "ew_offset", "DLS": "dls", "VSEC": "vsec",
}


def _remap_station_keys(rows):
    """Add lowercase column-name keys alongside the canonical keys so the loader
    finds the values whichever naming it uses."""
    out = []
    for st in (rows or []):
        if not isinstance(st, dict):
            out.append(st); continue
        m = dict(st)
        for canon, col in _SURVEY_KEY_MAP.items():
            if canon in st and col not in m:
                m[col] = st[canon]
        out.append(m)
    return out


def _load_directional(engine, dialect, well_info, rows, say):
    """Directional survey loader — CORRECTED path first (survey_loader, which
    writes cat_well_dir_srvy_hdr/_sta), falling back to legacy load_to_ppdm if
    the survey_loader refactor isn't deployed. Returns the loader's result dict.
    """
    rows = _remap_station_keys(rows)   # canonical MD/INC/AZI -> md/incl/azim columns
    # corrected path (prior-session fix): format-agnostic station loader
    for _imp in ("survey_loader", "dataview.file_catalog.survey_loader"):
        try:
            mod = __import__(_imp, fromlist=["load_directional_survey"])
            return mod.load_directional_survey(
                engine=engine, dialect=dialect, well_info=well_info,
                stations=rows)
        except Exception:
            continue
    # also try the pdf_db_loader delegate if present
    try:
        from dataview.file_catalog.pdf_db_loader import load_directional_survey as _ld
        return _ld(engine=engine, dialect=dialect, well_info=well_info,
                   rows=rows)
    except Exception:
        pass
    # legacy fallback (writes dbo.* — only works on older schemas)
    try:
        from dataview.file_catalog.pdf_survey_catalog import load_to_ppdm
        return load_to_ppdm(well_info=well_info, stations=rows,
                            engine=engine, dialect=dialect)
    except Exception as e:
        say(f"directional load: {e}")
        return {"loaded": 0, "detail": {}, "errors": [str(e)]}

def _do_dlis(engine, fpath, uwi, inv, fname, say) -> FileResult:
    """DLIS → cat_well_log + cat_well_log_curve as curve INVENTORY (one row per
    channel, NOT per sample). Same schema/LOG_ID pattern as the LAS handler, so
    it promotes to dv_well_log / dv_well_log_curve identically.

    DLIS files carry huge sample volume (100s-1000s of channels × thousands of
    samples × multiple frames). We store the inventory — mnemonic, unit, depth
    range, sample count — like SEGY stores survey geometry, not every value.
    numpy does min/max/count per channel; there is NO per-sample Python loop
    (that was the earlier hang). dlisio reads (C ext, GIL released) cleanly —
    verified on all 4 real test files.

    No UWI? DLIS identity is the filename → deterministic synthetic FN_<hash>
    so the parse path proves out; these HOLD at promote until a real dv_well
    exists (expected — real name→well resolution is a follow-up).
    """
    import uuid as _uuid, hashlib as _hashlib
    import numpy as _np
    from dlisio import dlis as _dlis
    from datetime import datetime as _dt
    capture, reset_replace_state = _capture()
    reset_replace_state()
    _now = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    res = FileResult("done", rt="WELL_LOG")

    # Load the file once; we read ORIGIN identity AND curves from it.
    f, *tail = _dlis.load(fpath)
    lfs = [f] + list(tail)

    # Identity, in order of trust: passed-in UWI → filename UWI → ORIGIN well_id
    # (if UWI-shaped) → synthetic FN_<hash>. The WELL NAME for triage comes from
    # the ORIGIN well_name FIRST (what the logging company recorded), falling
    # back to a name parsed from the filename — ORIGIN is trusted over the
    # filename because filenames are often renamed/invented on disk.
    # Identity from the FILENAME: UWI if present, else a parsed well name, else
    # synthetic FN_<hash>. We deliberately do NOT read the DLIS ORIGIN here:
    # dlisio parses origin objects in a separate pass (~1.3s/file on top of the
    # frame read, a 4.5x slowdown), and for our data the filenames carry the
    # identity. If a deployment has unreliable filenames where the ORIGIN
    # well_name is the only true identity, re-enable an ORIGIN read here and
    # accept the cost — it's a per-site tradeoff, off by default for speed.
    fn_uwi = _uwi_from_filename(fname, fpath) if not uwi else None
    well_name = _wellname_from_filename(fname, fpath)
    real_uwi = uwi or fn_uwi
    key = real_uwi or ("FN_" + _hashlib.sha1(
        (os.path.splitext(fname or "")[0] or "UNKNOWN").upper()
        .encode("utf-8")).hexdigest()[:14].upper())
    _logid = f"{key}-DLIS"

    # Well-identity row → cat_well with the filename WELL_NAME so triage can
    # match name→gold. Skipped for nameless wrapper-only files.
    if well_name or real_uwi:
        try:
            capture(engine, "cat_well", [{
                "WELL_NAME":        well_name or key,
                "ACTIVE_IND":       "Y",
                "ROW_CREATED_BY":   "DataWrangler",
                "ROW_CREATED_DATE": _now,
            }], uwi=real_uwi, inventory_id=inv,
               source_path=fpath, source="DLIS")
        except Exception as e:
            say(f"dlis cat_well: {e}")

    curve_rows = []
    overall_top = overall_base = None
    depth_uom = "ft"
    try:
        for lf in lfs:
            unit_by = {}
            # Building unit_by iterates EVERY channel in the logical file and
            # touches .fingerprint/.units on each — which can force dlisio to
            # materialize the full channel object pool (expensive on files with
            # hundreds of channels). Units are a nice-to-have on curve inventory;
            # set DW_DLIS_UNITS=0 to skip the scan and stay fast.
            if os.environ.get("DW_DLIS_UNITS", "1") != "0":
                try:
                    for c in lf.channels:
                        unit_by[str(getattr(c, "fingerprint", "") or "")] = \
                            getattr(c, "units", "") or ""
                except Exception:
                    pass
            for fr in lf.frames:
                # FAST inventory path — no sample materialization:
                #   • curve names  ← fr.dtype.names (structured dtype, metadata)
                #   • depth range  ← fr.index_min / fr.index_max (precomputed)
                #   • index name   ← fr.index
                # This avoids fr.curves(), which decoded every sample of every
                # channel (the ~700ms+/file cost). top/base need not be precise,
                # and the frame index range is exactly the right cheap source.
                try:
                    _dt = fr.dtype() if callable(getattr(fr, "dtype", None)) else fr.dtype
                    names = list(_dt.names or [])
                except Exception:
                    names = []
                if not names:
                    continue
                idx_name = getattr(fr, "index", None)
                if idx_name not in names:
                    idx_name = names[0] if names else None

                def _fnum(v):
                    try:
                        x = float(v)
                        return x if x == x else None  # NaN guard
                    except (TypeError, ValueError):
                        return None
                f_top = _fnum(getattr(fr, "index_min", None))
                f_base = _fnum(getattr(fr, "index_max", None))
                if f_top is not None and f_base is not None and f_top > f_base:
                    f_top, f_base = f_base, f_top   # normalize if reversed

                depth_uom = _unit_for(unit_by, idx_name) or depth_uom
                if f_top is not None:
                    overall_top = f_top if overall_top is None else min(overall_top, f_top)
                if f_base is not None:
                    overall_base = f_base if overall_base is None else max(overall_base, f_base)

                for nm in names:
                    if nm == idx_name:
                        continue
                    curve_rows.append({
                        "UWI":               key,
                        "LOG_ID":            _logid,
                        "CURVE_ID":          str(nm)[:40],
                        "MNEMONIC":          str(nm),
                        "CURVE_UNIT":        _unit_for(unit_by, nm) or None,
                        "TOP_DEPTH":         f_top,
                        "BASE_DEPTH":        f_base,
                        "DEPTH_OUOM":        depth_uom,
                        "SAMPLE_COUNT":      None,
                        "ACTIVE_IND":        "Y",
                        "ROW_CREATED_BY":    "DataWrangler",
                        "ROW_CREATED_DATE":  _now,
                    })
    finally:
        for lf in lfs:
            try:
                lf.close()
            except Exception:
                pass

    if curve_rows:
        n = capture(engine, "cat_well_log_curve", curve_rows,
                    uwi=key, inventory_id=inv, source_path=fpath, source="DLIS")
        res.rows_written += (n or 0)
        if n:
            res.detail["cat_well_log_curve"] = n
    # log header
    try:
        n = capture(engine, "cat_well_log", [{
            "LOG_ID":           _logid,
            "TOP_DEPTH":        overall_top,
            "BASE_DEPTH":       overall_base,
            "DEPTH_OUOM":       depth_uom,
            "FILE_PATH":        fpath,
            "FILE_FORMAT":      "DLIS",
            "ACTIVE_IND":       "Y",
            "ROW_CREATED_BY":   "DataWrangler",
            "ROW_CREATED_DATE": _now,
        }], uwi=key, inventory_id=inv, source_path=fpath, source="DLIS")
        if n:
            res.detail["cat_well_log"] = n
    except Exception as e:
        say(f"dlis log header: {e}")
    if not curve_rows:
        res.status = "skip"
    return res


def _do_lis(engine, fpath, uwi, inv, fname, say) -> FileResult:
    """LIS → curve INVENTORY via the project's lis_catalog (same module the
    extract stage uses). Writes the same cat_well_log + cat_well_log_curve
    schema as LAS/DLIS. Uses the curve METADATA lis_catalog already returns
    (curve_names, n_curves, depth range) — no binary sample layout guessing. If
    lis_catalog exposes no curve names, returns skip gracefully.
    """
    import hashlib as _hashlib
    from datetime import datetime as _dt
    capture, reset_replace_state = _capture()
    reset_replace_state()
    _now = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    res = FileResult("done", rt="WELL_LOG")

    try:
        from dataview.file_catalog.lis_catalog import classify_lis
    except ImportError:
        try:
            from dataview.file_catalog.lis_catalog import classify_lis
        except ImportError:
            return FileResult("skip", rt="LIS",
                              detail={"note": "lis_catalog unavailable"})

    cl = classify_lis(fpath)
    cnames = cl.get("curve_names") or []
    if not cnames:
        return FileResult("skip", rt="LIS",
                          detail={"note": "no curve names from lis_catalog"})

    fn_uwi = _uwi_from_filename(fname, fpath) if not uwi else None
    # Prefer the well name lis_catalog read from the file over a filename guess
    # (same principle as DLIS ORIGIN-first). Filename is the fallback.
    cl_name = ((cl.get("well_name") or "").strip() or None) if isinstance(cl, dict) else None
    cl_oper = ((cl.get("operator") or "").strip() or None) if isinstance(cl, dict) else None
    cl_uwi = _uwi_from_filename(str(cl.get("uwi") or "")) if isinstance(cl, dict) else None
    well_name = cl_name or _wellname_from_filename(fname, fpath)
    real_uwi = uwi or fn_uwi or cl_uwi
    key = real_uwi or ("FN_" + _hashlib.sha1(
        (os.path.splitext(fname or "")[0] or "UNKNOWN").upper()
        .encode("utf-8")).hexdigest()[:14].upper())
    _logid = f"{key}-LIS"
    if well_name or real_uwi:
        try:
            capture(engine, "cat_well", [{
                "WELL_NAME":        well_name or key,
                "OPERATOR_NAME":    cl_oper,
                "ACTIVE_IND":       "Y",
                "ROW_CREATED_BY":   "DataWrangler",
                "ROW_CREATED_DATE": _now,
            }], uwi=real_uwi, inventory_id=inv,
               source_path=fpath, source="LIS")
        except Exception as e:
            say(f"lis cat_well: {e}")
    d_top = cl.get("depth_start")
    d_base = cl.get("depth_stop")
    d_uom = cl.get("depth_ouom") or "ft"

    curve_rows = [{
        "UWI":               key,
        "LOG_ID":            _logid,
        "CURVE_ID":          str(m)[:40],
        "MNEMONIC":          str(m),
        "TOP_DEPTH":         d_top,
        "BASE_DEPTH":        d_base,
        "DEPTH_OUOM":        d_uom,
        "ACTIVE_IND":        "Y",
        "ROW_CREATED_BY":    "DataWrangler",
        "ROW_CREATED_DATE":  _now,
    } for m in cnames if m]

    if curve_rows:
        n = capture(engine, "cat_well_log_curve", curve_rows,
                    uwi=key, inventory_id=inv, source_path=fpath, source="LIS")
        res.rows_written += (n or 0)
        if n:
            res.detail["cat_well_log_curve"] = n
    try:
        n = capture(engine, "cat_well_log", [{
            "LOG_ID":           _logid,
            "TOP_DEPTH":        d_top,
            "BASE_DEPTH":       d_base,
            "DEPTH_OUOM":       d_uom,
            "FILE_PATH":        fpath,
            "FILE_FORMAT":      "LIS",
            "ACTIVE_IND":       "Y",
            "ROW_CREATED_BY":   "DataWrangler",
            "ROW_CREATED_DATE": _now,
        }], uwi=key, inventory_id=inv, source_path=fpath, source="LIS")
        if n:
            res.detail["cat_well_log"] = n
    except Exception as e:
        say(f"lis log header: {e}")
    if not curve_rows:
        res.status = "skip"
    return res


def _unit_for(unit_by, mnem):
    """Best-effort unit lookup: dlisio fingerprints embed the mnemonic."""
    if not unit_by or not mnem:
        return ""
    for fp, u in unit_by.items():
        if mnem in fp:
            return u
    return ""

def _do_shapefile(engine, fpath, uwi, inv, say) -> FileResult:
    """Shapefile → classify, route by feature type. WELL points → cat_well.
    FIELD/LEASE/BOUNDARY/PIPELINE → per-feature capture into their cat_* table
    (one row per polygon/line with geometry). SEISMIC skips here (handled via the
    header → dv_seis_set promote path)."""
    res = FileResult("done", rt="SHAPEFILE")
    try:
        from dataview.mapping.shapefile_catalog import (
            classify_shapefile, capture_wells_to_catalog,
            capture_features_to_catalog, FT_WELL)
    except ImportError:
        from dataview.mapping.shapefile_catalog import (
            classify_shapefile, capture_wells_to_catalog,
            capture_features_to_catalog, FT_WELL)
    well_info = {"uwi": uwi, "source_path": fpath, "inventory_id": inv}
    cl = classify_shapefile(fpath)
    ft = cl.get("feature_type")
    # classifier feature_type -> capture category for the generic per-feature path
    _FEAT_CAT = {"FIELD": "FIELD", "LEASE": "LAND_TRACT",
                 "BOUNDARY": "BOUNDARY", "PIPELINE": "PIPELINE"}
    if ft == FT_WELL:
        r = capture_wells_to_catalog(
            file_path=fpath, column_map=cl.get("column_map") or {},
            engine=engine, well_info=well_info, dialect="mssql",
            source="SHAPEFILE")
        n = r.get("loaded", 0)
        res.rows_written = n
        if n:
            res.detail["cat_well"] = n
        real = [e for e in r.get("errors", [])
                if not str(e).startswith("header capture:")]
        if real:
            res.status = "error"
            res.error = str(real[0])[:400]
        elif not n:
            res.status = "skip"
    elif ft in _FEAT_CAT:
        r = capture_features_to_catalog(
            file_path=fpath, feature_category=_FEAT_CAT[ft],
            engine=engine, well_info=well_info, dialect="mssql",
            source="SHAPEFILE")
        n = r.get("loaded", 0)
        res.rows_written = n
        res.detail.update(r.get("detail", {}))
        real = [e for e in r.get("errors", [])
                if not str(e).startswith("header capture:")]
        if real:
            res.status = "error"
            res.error = str(real[0])[:400]
        elif not n:
            res.status = "skip"
    else:
        res.status = "skip"
        res.detail["note"] = f"shapefile_skip:{ft}"
    return res


def _docx_uwi(fpath):
    """Pull a UWI from a .docx by scanning BOTH paragraphs and TABLE CELLS.

    Completion/well reports routinely put the UWI in a table (e.g.
    | API / UWI | 42-330-00035-00-00 |), and a paragraph-only text scan misses
    it — which is why these errored "Could not find UWI" despite the UWI being
    right there. We read paragraph text AND every table cell, then apply the
    standard UWI patterns. Returns bare-digit UWI or None.
    """
    ext = os.path.splitext(fpath)[1].lower()
    if ext != ".docx":
        return None
    try:
        import docx  # python-docx
    except ImportError:
        return None
    try:
        d = docx.Document(fpath)
    except Exception:
        return None
    chunks = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    text = "\n".join(c for c in chunks if c)

    patterns = [
        r'(?:UWI|API)[^\d]{0,12}(\d{2}-\d{3}-\d{5}-\d{2}-\d{2})',
        r'(\d{2}-\d{3}-\d{5}-\d{2}-\d{2})',
        r'(?:UWI|API)[^\d]{0,12}(\d{14})',
        r'\b(\d{14})\b',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return re.sub(r"[\-\s/]", "", m.group(1)).strip() or None
    return None


def _do_office(engine, fpath, uwi, inv, say) -> FileResult:
    """Office (xlsx/docx/...) → dv_office_loader.dispatch self-parses and routes
    to the formation-tops / completion / production sub-loader. The loader now
    extracts the docx UWI from table cells itself (incl. label-value grids), so
    we just dispatch."""
    import os as _os
    if _os.path.splitext(fpath)[1].lower() == ".docx":
        # Final Well Reports (well/core/survey) -> the loader's DOCX extractor via
        # capture_document. If it captures nothing (e.g. a formation-tops or
        # completion .docx it does not read), fall through to dv_office_loader below.
        from dataview.file_catalog.catalog_doc_capture import (
            capture_document, has_docx_extractor)
        if has_docx_extractor():
            _r = capture_document(engine, "mssql", fpath, ".docx", uwi, inv, log=say)
            _n = int(_r.get("loaded", 0) or 0)
            if _n:
                _fr = FileResult("done", rows_written=_n,
                                 detail=_r.get("detail", {}), rt=_r.get("rt", "DOCX"))
                return _fr
            # 0 rows -> not an FWR the loader reads; try dv_office_loader instead

    res = FileResult("done", rt="OFFICE")
    try:
        from dataview.file_catalog.dv_office_loader import dispatch as _office
    except ImportError:
        from dataview.file_catalog.dv_office_loader import dispatch as _office
    r = _office(engine, fpath, source="OFFICE")
    n = int(r.get("loaded", 0) or 0)
    res.rows_written = n
    errs = [str(e) for e in (r.get("errors") or [])]
    if n:
        res.detail["office"] = n
    elif any("No loader found" in e for e in errs):
        res.status = "skip"
        res.detail["note"] = "not_impl:OFFICE"
    elif errs:
        res.status = "error"
        res.error = errs[0][:400]
    else:
        res.status = "skip"
    return res


def _do_json(engine, fpath, uwi, inv, say) -> FileResult:
    """OSDU / JSON-Well-Log → json_well_log_catalog self-parses, routing by the
    `kind` field. Non-well kinds return a no_target note (clean skip). Ported
    from the JSON_LOG branch."""
    res = FileResult("done", rt="OSDU")
    try:
        from dataview.file_catalog.json_well_log_catalog import load_json_well_log as _jwl
    except ImportError:
        from dataview.file_catalog.json_well_log_catalog import load_json_well_log as _jwl
    well_info = {"uwi": uwi, "source_path": fpath, "inventory_id": inv}
    r = _jwl(engine, fpath, uwi=uwi, inventory_id=inv,
             source_path=fpath, well_info=well_info)
    n = int(r.get("loaded", 0) or 0)
    res.rows_written = n
    res.rt = r.get("rt") or "OSDU"
    res.detail.update(r.get("detail") or {})
    errs = [str(e) for e in (r.get("errors") or [])]
    if n:
        pass
    elif errs and not any("no_target" in e.lower() for e in errs):
        res.status = "error"
        res.error = errs[0][:400]
    else:
        res.status = "skip"
        if r.get("note"):
            res.detail["note"] = r["note"]
    return res


def _do_segy(engine, fpath, inv, say) -> FileResult:
    """SEGY/P190 → seismic HEADER ONLY (no trace data) into
    file_catalog.FILE_SEIS_HEADER, which promote lifts into dataview.dv_seis_set.

    Brings seismic into the unified worker path (previously handled by the
    separate extract stage). extract_core._extract_fields does the heavy header
    parse + CDP reprojection + survey-outline build (streamlit-free; segyio is a
    C ext that releases the GIL → parallelizes well). We MERGE the resulting
    fields into FILE_SEIS_HEADER with the same SQL page_workbench uses. No trace
    samples are ever read — headers only, by design.
    """
    import uuid as _uuid, os as _os
    from sqlalchemy import text as _t
    import sys as _sys
    _sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
    from dataview.file_catalog import extract_core

    res = FileResult("done", rt="SEISMIC")
    fext = _os.path.splitext(fpath)[1].lower()
    fields = extract_core._extract_fields(fpath, fext)
    if fields.get("file_category") != "SEIS":
        return FileResult("skip", rt="SEISMIC",
                          detail={"note": "not classified as seismic"})

    def _trunc(v, n):
        if v is None:
            return None
        s = str(v).strip()
        return s[:n] if s else None

    def _num(v):
        try:
            return float(str(v).strip())
        except (TypeError, ValueError):
            return None

    def _coord(v):
        f = _num(v)
        return f if (f is not None and -180.0 <= f <= 180.0) else None

    def _int(v):
        try:
            return int(float(str(v).strip()))
        except (TypeError, ValueError):
            return None

    inv_id = str(inv) if inv is not None else fpath
    params = {
        "hid":      _uuid.uuid5(_uuid.NAMESPACE_URL, inv_id + "_s").hex.upper(),
        "inv_id":   inv,
        "sn":       _trunc(fields.get("survey_name"), 255),
        "snsrc":    _trunc(fields.get("survey_name_source"), 30),
        "ln":       _trunc(fields.get("line_name"), 255),
        "stype":    _trunc(fields.get("seis_set_type"), 40),
        "sd":       _trunc(fields.get("survey_date"), 20),
        "contr":    _trunc(fields.get("contractor"), 255),
        "bmin_lat": _coord(fields.get("bbox_min_lat")),
        "bmax_lat": _coord(fields.get("bbox_max_lat")),
        "bmin_lon": _coord(fields.get("bbox_min_lon")),
        "bmax_lon": _coord(fields.get("bbox_max_lon")),
        "epsg":     _int(fields.get("epsg_code")),
        "si":       _num(fields.get("sample_interval")),
        "tc":       _int(fields.get("trace_count")),
        "sf":       _trunc(fields.get("shot_first"), 20),
        "sl":       _trunc(fields.get("shot_last"), 20),
        "il_min":   fields.get("il_min"),
        "il_max":   fields.get("il_max"),
        "xl_min":   fields.get("xl_min"),
        "xl_max":   fields.get("xl_max"),
        "outline":  fields.get("survey_outline"),
    }
    try:
        with engine.begin() as con:
            ensure_seis_columns(con)
            con.execute(_t(_SQL_SEIS_MERGE), params)
        res.rows_written = 1
        res.detail["FILE_SEIS_HEADER"] = 1
    except Exception as e:
        return FileResult("error", rt="SEISMIC",
                          error=f"seis header write: {str(e)[:300]}")
    return res


# Well header MERGE (verbatim from page_workbench's _SQL_WELL_MERGE, streamlit-
# free). The pool writes cat_well for promote, but the documents map and triage
# read file_catalog.FILE_WELL_HEADER — so we ALSO upsert that here, keyed on the
# same deterministic WELL_HEADER_ID = uuid5(NAMESPACE_URL, inv_id) page_workbench
# uses, so a re-process updates in place rather than duplicating.
_SQL_WELL_MERGE = """
    MERGE file_catalog.FILE_WELL_HEADER AS tgt
    USING (SELECT :hid AS WELL_HEADER_ID) src
    ON tgt.WELL_HEADER_ID = src.WELL_HEADER_ID
    WHEN MATCHED THEN UPDATE SET
        UWI=:uwi, UWI14=:uwi14, WELL_NAME=:wn, OPERATOR=:op,
        WELL_FIELD=:fld, STATE=:st, COUNTY=:co,
        LATITUDE=:lat, LONGITUDE=:lon,
        TOTAL_DEPTH=:td, SPUD_DATE=:spud,
        RIG_RELEASE=:rig, REPORT_TYPE=:rt,
        SURVEY_TYPE=:stype, CONTRACTOR=:contr,
        CONFIDENCE=:conf, EXTRACTED_DATE=GETUTCDATE()
    WHEN NOT MATCHED THEN INSERT (
        WELL_HEADER_ID,INVENTORY_ID,
        UWI,UWI14,WELL_NAME,OPERATOR,WELL_FIELD,
        STATE,COUNTY,LATITUDE,LONGITUDE,
        TOTAL_DEPTH,SPUD_DATE,RIG_RELEASE,
        REPORT_TYPE,SURVEY_TYPE,CONTRACTOR,CONFIDENCE,
        EXTRACTED_DATE,EXTRACTED_BY
    ) VALUES (
        :hid,:inv_id,
        :uwi,:uwi14,:wn,:op,:fld,
        :st,:co,:lat,:lon,
        :td,:spud,:rig,
        :rt,:stype,:contr,:conf,
        GETUTCDATE(),'DataWrangler'
    );
"""


def _write_well_header(engine, inv, uwi, cl, report_type, say, con=None):
    """Upsert one FILE_WELL_HEADER row from the pool's classify dict `cl`, so the
    documents map / triage (which read FILE_WELL_HEADER, not cat_well) see wells
    the pool loads. Mirrors page_workbench._well_params. Best-effort: a header
    failure must not fail the file (the cat_well write is the source of truth for
    promote)."""
    import uuid as _uuid
    from sqlalchemy import text as _t

    def _tr(v, n):
        if v is None:
            return None
        s = str(v).strip()
        return s[:n] if s else None

    inv_id = str(inv) if inv is not None else "_nofid_"
    params = {
        "hid":    _uuid.uuid5(_uuid.NAMESPACE_URL, inv_id).hex.upper(),
        "inv_id": inv,
        "uwi":    _tr(uwi, 40),
        "uwi14":  _uwi14(uwi),
        "wn":     _tr(cl.get("well_name"), 255),
        "op":     _tr(cl.get("operator"), 255),
        "fld":    _tr(cl.get("field"), 100),
        "st":     _tr(cl.get("state"), 50),
        "co":     _tr(cl.get("county"), 100),
        "lat":    _tr(cl.get("latitude"), 30),
        "lon":    _tr(cl.get("longitude"), 30),
        "td":     _tr(cl.get("total_depth"), 20),
        "spud":   _tr(cl.get("spud_date"), 20),
        "rig":    _tr(cl.get("rig_release"), 20),
        "rt":     _tr(report_type, 50),
        "stype":  _tr(cl.get("survey_type"), 50),
        "contr":  _tr(cl.get("contractor"), 255),
        "conf":   None,
    }
    try:
        import contextlib as _ctx
        _cm = _ctx.nullcontext(con) if con is not None else engine.begin()
        with _cm as _c:
            _c.execute(_t(_SQL_WELL_MERGE), params)
        return 1
    except Exception as e:
        say(f"FILE_WELL_HEADER write: {str(e)[:200]}")
        return 0


def _do_witsml(engine, fpath, uwi, inv, say) -> FileResult:
    """WITSML (.xml) → witsml_catalog self-parses (well header / trajectory /
    log curves). Ported from the WITSML branch. Gates non-WITSML .xml itself."""
    res = FileResult("done", rt="WITSML")
    try:
        from dataview.file_catalog.witsml_catalog import load_witsml as _witsml
    except ImportError:
        from dataview.file_catalog.witsml_catalog import load_witsml as _witsml
    well_info = {"uwi": uwi, "source_path": fpath, "inventory_id": inv}
    r = _witsml(engine, fpath, uwi=uwi, inventory_id=inv,
                source_path=fpath, well_info=well_info)
    n = int(r.get("loaded", 0) or 0)
    res.rows_written = n
    res.rt = r.get("rt") or "WITSML"
    res.detail.update(r.get("detail") or {})
    errs = [str(e) for e in (r.get("errors") or [])]
    if n:
        pass
    elif errs:
        res.status = "error"
        res.error = errs[0][:400]
    else:
        res.status = "skip"
        if r.get("note"):
            res.detail["note"] = r["note"]
    return res
