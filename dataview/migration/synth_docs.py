"""
dataview/migration/synth_docs.py
===============================
Synthetic well documents for File Catalog testing, modelled on real reports.

WHY THE DETAIL MATTERS
----------------------
An earlier version of this produced a header block and one table per document.
That exercises "can we find a UWI in a PDF" and nothing else. Real well
documents are dense — a casing and cementing record carries a string programme,
a cement job summary AND a CBL evaluation; an end-of-well report carries a
stratigraphic column, an NPT ledger with cost categories, and a completion
summary. Extractors that cope with a thin document routinely fail on a real one,
because the real failure modes are multi-table pages, repeated column headers,
numbers that look like depths but aren't, and units baked into values.

So these follow the shape of actual reports: several tables per document,
realistic magnitudes, and the incidental noise — confidentiality footers,
service company names, interpreter credentials — that surrounds the data you
actually want.

GROUND TRUTH
------------
Every document is generated from a known well, and MANIFEST.csv records the UWI
each file should resolve to plus the case it represents. That makes catalog
testing a score — matched correctly / matched wrongly / not matched — rather
than an impression. A confidently wrong UWI is worse than none: it attaches a
document to another operator's well and nothing complains.

THE HARD CASES ARE DELIBERATE
-----------------------------
A share of the output is difficult on purpose, each labelled in the manifest:
UWI only in the filename; only in the text; dashed API rather than 14-digit; a
UWI belonging to no well; none at all; two wells in one document; and an
image-only scan with no text layer. Worth noting the real reports these were
modelled on had EMPTY "UWI / API:" fields — a document naming an operator, a
field and a county but never its own well is common, so it belongs in the set.

USAGE (from the repo root)
--------------------------
    py -m dataview.migration.synth_docs --wells-csv C:\\synth\\dv_well.csv ^
        --out C:\\synth_docs --per-well 3
"""
from __future__ import annotations

import argparse
import csv
import math
import os
import random
from datetime import date, timedelta

HARD_CASE_SHARE = 0.22

# Stratigraphy by region, so a Texas well doesn't come back with Kansas tops.
STRAT_PERMIAN = [
    ("Rustler", 1240, ""), ("Salado", 1420, "Salt section"),
    ("Castile", 3310, ""), ("Bell Canyon", 3730, ""),
    ("Cherry Canyon", 4070, ""), ("Brushy Canyon", 4550, ""),
    ("Bone Spring 1st", 5070, "Shows"), ("Bone Spring 2nd", 5410, "Shows"),
    ("Bone Spring 3rd", 5980, ""), ("Wolfcamp A", 6840, "TARGET"),
    ("Wolfcamp B", 7320, ""), ("Wolfcamp C", 7940, ""),
    ("Dean", 8460, ""), ("Spraberry Upper", 8890, ""),
]
STRAT_MIDCON = [
    ("Stone Corral", 1450, "Anhydrite"), ("Hutchinson Salt", 1600, "Salt"),
    ("Chase", 2400, ""), ("Council Grove", 2700, ""), ("Admire", 2900, ""),
    ("Lansing", 3200, ""), ("Kansas City", 3350, ""), ("Marmaton", 3600, ""),
    ("Cherokee", 3750, "TARGET"), ("Mississippian", 3950, "Shows"),
    ("Arbuckle", 4300, ""), ("Precambrian", 4600, "Basement"),
]
FIELDS_BY_STATE = {
    "42": ["Spraberry Trend", "Delaware Basin", "Midland Basin", "Wolfcamp Shale"],
    "35": ["SCOOP/STACK", "Anadarko Basin", "Cana-Woodford"],
    "30": ["Delaware Basin", "Northwest Shelf"],
    "15": ["Hugoton", "El Dorado", "Chase-Silica", "Trapp"],
}
_ABBR = {"42": "TX", "35": "OK", "30": "NM", "15": "KS", "05": "CO", "49": "WY"}
SERVICE_COS = ["Halliburton", "Schlumberger", "Baker Hughes", "Weatherford"]
RIGS = ["Patterson UTI #219", "Helmerich & Payne #451", "Nabors X-12",
        "Precision Drilling #308", "Cactus Rig #145"]
INTERPRETERS = ["J. Rodriguez, M.Sc. Petrophysics", "A. Whitfield, P.Geo.",
                "M. Okonkwo, Senior Petrophysicist", "L. Trevino, M.Sc."]

CASING_PROGRAMME = [
    ("Conductor",    '20"',      94,   "K-55",  0.016, "Float shoe"),
    ("Surface",      '13-3/8"',  54.5, "K-55",  0.17,  "Float shoe + collar"),
    ("Intermediate", '9-5/8"',   47,   "L-80",  0.55,  "Float shoe + collar"),
    ("Production",   '5-1/2"',   20,   "P-110", 1.00,  "Float shoe + collar"),
]
NPT_EVENTS = [
    ("Stuck pipe — worked free", "Formation-related"),
    ("BHA twist-off — fishing op", "Mechanical"),
    ("Lost circulation — cement squeeze", "Formation-related"),
    ("MWD failure — replacement", "Equipment"),
    ("Waiting on weather", "Weather"),
    ("Mud motor failure", "Equipment"),
]
CURVES = ["GR", "CALI", "SP", "RILD", "RT", "RHOB", "NPHI", "DPHI", "PEF",
          "DT", "DTSM"]


def _dashed(uwi):
    return f"{uwi[:2]}-{uwi[2:5]}-{uwi[5:10]}" if len(uwi) >= 10 else uwi


def _long_uwi(uwi):
    """42-317-12345-00-00 — the form the Word examples used."""
    return f"{_dashed(uwi)}-00-00"


def _st(w):
    return str(w.get("province_state", ""))[:2]


def _strat(w):
    return STRAT_PERMIAN if _st(w) in ("42", "30") else STRAT_MIDCON


def _field(w, rng):
    return rng.choice(FIELDS_BY_STATE.get(_st(w), ["Unnamed Field"]))


def _abbr(w):
    return _ABBR.get(_st(w), "US")


# --------------------------------------------------------------------------- #
# PDF scaffolding
# --------------------------------------------------------------------------- #
def _pdf(path, title, subtitle, ident_rows, blocks, footer):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)
    ss = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=ss["Title"], fontSize=15, spaceAfter=2)
    sub = ParagraphStyle("sub", parent=ss["Normal"], fontSize=9,
                         textColor=colors.HexColor("#555555"), spaceAfter=10)
    hd = ParagraphStyle("hd", parent=ss["Heading2"], fontSize=10.5, spaceBefore=9)
    body = ParagraphStyle("body", parent=ss["Normal"], fontSize=8.5, leading=11)
    foot = ParagraphStyle("foot", parent=ss["Normal"], fontSize=7,
                          textColor=colors.HexColor("#777777"), spaceBefore=13)

    story = [Paragraph(title, h1), Paragraph(subtitle, sub)]
    if ident_rows:
        pairs, row = [], []
        for k, v in ident_rows:
            row += [k, str(v)]
            if len(row) == 4:
                pairs.append(row)
                row = []
        if row:
            pairs.append(row + [""] * (4 - len(row)))
        t = Table(pairs, colWidths=[1.25*inch, 2.1*inch, 1.25*inch, 2.1*inch])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
        ]))
        story += [t, Spacer(1, 6)]

    for heading, content in blocks:
        if heading:
            story.append(Paragraph(heading, hd))
        if isinstance(content, tuple):
            head, rows, widths = content
            t = Table([head] + rows, repeatRows=1,
                      colWidths=[x*inch for x in widths] if widths else None)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6e6e6")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.2),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(t)
        else:
            story.append(Paragraph(content, body))
        story.append(Spacer(1, 4))
    story.append(Paragraph(footer, foot))
    SimpleDocTemplate(path, pagesize=letter, title=title, leftMargin=44,
                      rightMargin=44, topMargin=44, bottomMargin=36).build(story)


def _ident(w, shown, rng, extra=()):
    return [("Operator", w.get("operator_name", "")),
            ("Well Name", w.get("well_name", "")),
            ("UWI / API", shown),
            ("Field", _field(w, rng)),
            ("State", _abbr(w)), ("County", w.get("county", ""))] + list(extra)


# --------------------------------------------------------------------------- #
# PDF document types
# --------------------------------------------------------------------------- #
def scout_ticket(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    casing = [[n, od, wt, gr, f"{round(td*fr):,}", f"{rng.randint(100,900):,}"]
              for n, od, wt, gr, fr, _ in CASING_PROGRAMME]
    n_st = rng.randint(14, 42)
    stages, top = [], tvd + rng.uniform(50, 300)
    for i in range(1, n_st + 1):
        if i in (1, 2, 5, 10, 15, 20, n_st):
            stages.append([i, f"{top:,.0f}", f"{top+350:,.0f}", "350",
                           f"{rng.randint(3000,3600):,}",
                           f"{rng.choice([450000,500000,550000]):,}",
                           f"{rng.randint(8100,8500):,}", rng.randint(58, 64)])
        top += (td - tvd) / n_st
    ip, q, d0 = [], rng.uniform(900, 2100), date(2024, 6, 20)
    for days in (5, 35, 65, 95):
        gas = q * rng.uniform(1.4, 1.9)
        ip.append([(d0 + timedelta(days=days)).isoformat(), days, f"{q:,.0f}",
                   f"{gas:,.0f}", f"{q*rng.uniform(0.15,0.55):,.0f}",
                   f"{gas/q*1000:,.0f}", f"{rng.choice([32,40,48,56])}/64",
                   f"{rng.randint(3200,4300):,}"])
        q *= rng.uniform(0.78, 0.88)
    _pdf(path, "WELL SCOUT TICKET", "IHS Markit · Basin Intelligence",
         _ident(w, shown, rng, [
             ("Well Type", f"{w.get('well_type','Oil')} — Horizontal"),
             ("Status", w.get("well_status", "Producing")),
             ("Spud Date", w.get("spud_date", "")),
             ("Completion Date", w.get("completion_date", "")),
             ("Rig", rng.choice(RIGS)),
             ("KB Elevation", f"{w.get('kb_elevation','')} ft"),
             ("Total Depth", f"{td:,.0f} ft MD"), ("TVD", f"{tvd:,.0f} ft"),
             ("Lateral Length", f"{td-tvd:,.0f} ft"),
             ("Azimuth", f"{rng.randint(0,359)}°")]),
         [("Completion Information",
           (["Casing String", "Size (in)", "Weight (ppf)", "Grade",
             "Set Depth (ft MD)", "Cement (sacks)"], casing,
            [1.1, 0.8, 1.0, 0.7, 1.3, 1.2])),
          ("Perforation &amp; Stimulation",
           (["Stage #", "Top (ft MD)", "Base (ft MD)", "Interval (ft)",
             "Fluid (bbl)", "Proppant (lbs)", "Max Press (psi)", "Rate (bpm)"],
            stages, [0.6, 0.9, 0.9, 0.8, 0.85, 1.05, 1.0, 0.7])),
          ("Initial Production (IP)",
           (["Test Date", "Days On", "Oil (bbl/d)", "Gas (Mcf/d)",
             "Water (bbl/d)", "GOR (scf/bbl)", "Choke (in)", "WHP (psi)"], ip,
            [0.95, 0.7, 0.9, 0.9, 0.95, 1.0, 0.75, 0.8]))],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — Scout Ticket 2024")


def end_of_well(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    try:
        spud = date.fromisoformat(str(w.get("spud_date"))[:10])
    except Exception:
        spud = date(2024, 1, 8)
    days = rng.randint(95, 165)
    tops, prev = [], 0.0
    for name, nominal, note in _strat(w):
        if nominal > tvd * 0.98:
            break
        d = nominal * rng.uniform(0.9, 1.1)
        if d <= prev:
            continue
        prev = d
        tops.append([name, f"{d*td/max(tvd,1):,.0f}", f"{d:,.0f}",
                     f"{rng.randint(120,900):,}", note])
    npt, th, tc = [], 0, 0
    for day, (ev, cat) in zip(sorted(rng.sample(range(8, days), 4)),
                              rng.sample(NPT_EVENTS, 4)):
        h = rng.choice([12, 18, 24, 36, 48])
        c = round(h * rng.uniform(6, 9))
        th += h
        tc += c
        npt.append([day, ev, f"{rng.uniform(0.2,0.9)*td:,.0f}", h, c, cat])
    npt.append(["", "TOTAL NPT", "", th, tc, ""])
    n_st = rng.randint(24, 48)
    _pdf(path, "END OF WELL REPORT", "Final Well Summary",
         _ident(w, shown, rng, [
             ("Spud Date", spud.isoformat()),
             ("Rig Release", (spud + timedelta(days=days)).isoformat()),
             ("Total Depth", f"{td:,.0f} ft MD"), ("TVD", f"{tvd:,.0f} ft"),
             ("Lateral", f"{td-tvd:,.0f} ft"),
             ("Azimuth", f"{rng.randint(0,359)}°"), ("Elapsed Days", days),
             ("AFE Days", days - rng.randint(-12, 14)),
             ("Actual Cost", f"${rng.uniform(6,16):.1f} MM"),
             ("AFE Cost", f"${rng.uniform(6,15):.1f} MM")]),
         [("Stratigraphic Summary",
           (["Formation", "Top (ft MD)", "Top (ft TVD)", "Thickness (ft)",
             "Note"], tops, [1.5, 1.0, 1.0, 1.1, 1.6])),
          ("Drilling Events &amp; NPT",
           (["Day", "Event", "MD (ft)", "Duration (hrs)", "Cost ($K)",
             "Category"], npt, [0.5, 2.4, 0.9, 1.0, 0.85, 1.3])),
          ("Completion Summary",
           f"Stages completed: {n_st} of {n_st} planned. Total fluid: "
           f"{rng.randint(80,160)*1000:,} bbl. Total proppant: "
           f"{rng.randint(12,24)*1000000:,} lbs. Average stage spacing: "
           f"{rng.randint(240,340)} ft. Cluster spacing: "
           f"{rng.choice([25,30,35])} ft. Clusters per stage: "
           f"{rng.choice([4,5,6])}.")],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — End of Well Report 2024")


def casing_cement(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    prog = [[n, od, wt, gr, f"{round(td*fr):,}", f"{min(round(td*fr), tvd):,}",
             rng.randint(3, 220), fe] for n, od, wt, gr, fr, fe in CASING_PROGRAMME]
    toc = round(td * rng.uniform(0.3, 0.5))
    job = [["Cement Type", "Class H + Silica Flour", "Lead Slurry",
            f"{rng.uniform(12.8,14.2):.1f} ppg"],
           ["Tail Slurry", f"{rng.uniform(15.8,16.8):.1f} ppg",
            "Top of Cement (TOC)", f"{toc:,} ft MD"],
           ["Sacks Lead", f"{rng.randint(900,1600):,}", "Sacks Tail",
            f"{rng.randint(400,900):,}"],
           ["Mix Water (bbl)", rng.randint(300, 520), "Displacement (bbl)",
            rng.randint(600, 980)],
           ["BHCT (°F)", rng.randint(180, 240), "BHST (°F)",
            rng.randint(200, 260)],
           ["WOC Time (hrs)", 24, "Thickening Time",
            f"{rng.uniform(4.5,8.0):.1f} hrs"],
           ["24-hr Compressive Strength", f"{rng.randint(2400,3800):,} psi",
            "48-hr Strength", f"{rng.randint(4000,5600):,} psi"]]
    cbl, d = [], round(td * 0.17)
    while d < td and len(cbl) < 6:
        nxt = min(td, d + rng.uniform(1200, 3400))
        amp = rng.randint(6, 26)
        bond = ("Excellent (&gt;80%)" if amp < 10 else "Good (70–80%)"
                if amp < 16 else "Moderate (50–70%)")
        cbl.append([f"{d:,.0f}–{nxt:,.0f}", amp, bond,
                    "Possible microannulus" if amp > 18
                    else ("Lateral section" if d > tvd else "")])
        d = nxt
    _pdf(path, "CASING &amp; CEMENTING RECORD",
         f"{rng.choice(SERVICE_COS)} Cementing Services", _ident(w, shown, rng),
         [("Casing Programme",
           (["String", "OD (in)", "Weight (ppf)", "Grade", "Shoe MD (ft)",
             "Shoe TVD (ft)", "Centralizers", "Float Equipment"], prog,
            [1.0, 0.7, 0.9, 0.6, 0.95, 0.95, 0.95, 1.3])),
          ("Cement Job Summary — Production String",
           (["Parameter", "Value", "Parameter", "Value"], job,
            [1.7, 1.5, 1.7, 1.5])),
          ("Cement Evaluation (CBL/VDL)",
           (["Interval (ft MD)", "CBL Amplitude (mV)", "Cement Bond", "Note"],
            cbl, [1.6, 1.4, 1.6, 1.8]))],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — Casing &amp; Cementing Record 2024")


def well_test(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    zone = _strat(w)[-4][0]
    pres = rng.randint(3800, 5200)
    flow, q = [], rng.uniform(700, 1200)
    for i, (kind, hrs) in enumerate([("Flow", 8), ("Shut-in", 4), ("Flow", 8),
                                     ("Shut-in", 4), ("Flow", 8),
                                     ("Buildup", 24)], start=1):
        if kind == "Flow":
            gas = q * rng.uniform(1.4, 1.8)
            flow.append([i, kind, hrs, f"{rng.choice([24,32,40,48])}/64",
                         f"{q:,.0f}", f"{gas:,.0f}",
                         f"{q*rng.uniform(0.08,0.14):,.0f}",
                         f"{rng.randint(3100,3900):,}",
                         f"{rng.randint(3700,4300):,}"])
            q *= rng.uniform(1.25, 1.55)
        else:
            flow.append([i, kind, hrs, "—", "—", "—", "—", "—",
                         f"{pres-rng.randint(0,120):,}"])
    api_g, gor = rng.uniform(38, 48), rng.randint(900, 2200)
    res = [["Static Reservoir Pressure", f"{pres:,}", "psi",
            "Horner plot extrapolation"],
           ["Formation Permeability (k)", f"{rng.uniform(0.05,2.4):.2f}", "mD",
            "Buildup analysis"],
           ["Skin Factor (S)", f"{rng.uniform(-5,2):.1f}", "—",
            "Buildup analysis (stimulated)"],
           ["Productivity Index (PI)", f"{rng.uniform(0.1,0.9):.2f}",
            "bbl/d/psi", "Flow period 3"],
           ["Drainage Radius", f"{rng.randint(1200,3600):,}", "ft",
            "Transient analysis"],
           ["Reservoir Temperature", rng.randint(160, 240), "°F",
            "BHT measured"],
           ["Fluid Gravity", f"{api_g:.1f}", "°API", "Surface sample"],
           ["GOR", f"{gor:,}", "scf/bbl", "Flow period 3 average"]]
    samples = [["Separator oil", "Separator", 85, 120, f"{api_g:.1f}",
                f"{gor:,}", "0.2"],
               ["Separator gas", "Separator", 85, 120, "—", "—", "—"],
               ["Recombined", "Lab", rng.randint(190, 235), f"{pres:,}",
                f"{api_g-rng.uniform(0.5,2):.1f}",
                f"{gor+rng.randint(50,200):,}", "—"]]
    _pdf(path, "WELL TEST REPORT", "Production Test — Multi-Rate Flow Test",
         _ident(w, shown, rng, [
             ("Test Date", w.get("completion_date", "")),
             ("Test Type", "Multi-Rate Flow Test"), ("Zone", zone),
             ("Perforations", f"{tvd:,.0f}–{td:,.0f} ft MD"),
             ("Tubing", '2-7/8" EUE'), ("Separator", "Portable 3-phase")]),
         [("Flow Test Summary",
           (["Period", "Type", "Hrs", "Choke (in)", "Oil (bbl/d)",
             "Gas (Mcf/d)", "Water (bbl/d)", "FWHP (psi)", "FBHP (psi)"], flow,
            [0.55, 0.75, 0.5, 0.8, 0.9, 0.95, 1.0, 0.9, 0.9])),
          ("Reservoir Analysis",
           (["Parameter", "Value", "Units", "Method"], res,
            [1.9, 1.0, 0.9, 2.4])),
          ("Fluid Samples",
           (["Sample Type", "Collection Point", "Temp (°F)", "Press (psi)",
             "API Gravity", "GOR (scf/bbl)", "BS&amp;W (%)"], samples,
            [1.2, 1.3, 0.85, 0.95, 1.0, 1.15, 0.85]))],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — Well Test Report 2024")


def petrophysics(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    top = max(500.0, tvd - rng.uniform(600, 1400))
    zones, detail = [], []
    for name, _n, _x in _strat(w)[-5:]:
        gross = rng.uniform(220, 460)
        ng = rng.uniform(0.24, 0.78)
        zones.append([name, f"{top:,.0f}", f"{top+gross:,.0f}", f"{gross:,.0f}",
                      f"{gross*ng:,.0f}", f"{ng:.2f}",
                      f"{rng.uniform(4.0,9.5):.1f}", f"{rng.uniform(20,70):.0f}",
                      f"{rng.uniform(20,58):.0f}",
                      "High" if ng > 0.65 else "Moderate" if ng > 0.4 else "Low"])
        top += gross
    d = float(zones[0][1].replace(",", ""))
    for _ in range(8):
        detail.append([f"{d:,.0f}", f"{rng.uniform(48,92):.0f}",
                       f"{rng.uniform(2.38,2.58):.2f}",
                       f"{rng.uniform(9,16):.1f}", f"{rng.uniform(8,32):.1f}",
                       f"{rng.uniform(5,11):.1f}", f"{rng.uniform(16,42):.0f}",
                       f"{rng.uniform(20,42):.0f}",
                       rng.choice(["Silt/Lime", "Calcareous shale", "Shale",
                                   "Lime-rich silt", "Lime silt"])])
        d += 50
    cutoffs = [["Vcl (clay volume)", "&lt; 50%", "Core-log calibration"],
               ["PHIE (effective porosity)", "&gt; 4%", "Core plug analysis"],
               ["SW (water saturation)", "&lt; 65%",
                f"Archie, m={rng.uniform(1.7,2.0):.1f} n=2.0 "
                f"Rw={rng.uniform(0.02,0.08):.2f}"],
               ["RT (resistivity)", "&gt; 8 ohm·m", "Pickett plot"]]
    _pdf(path, "PETROPHYSICAL INTERPRETATION REPORT",
         f"{rng.choice(SERVICE_COS)} Petrotechnical Services",
         _ident(w, shown, rng, [
             ("Log Suite", " · ".join(rng.sample(CURVES, 6))),
             ("Interval", f"{zones[0][1]}–{zones[-1][2]} ft MD"),
             ("Interpreter", rng.choice(INTERPRETERS)),
             ("Date", w.get("completion_date", ""))]),
         [("Zone Summary",
           (["Zone", "Top (ft)", "Base (ft)", "Gross (ft)", "Net Pay (ft)",
             "N/G", "PHIE (%)", "SW (%)", "Vcl (%)", "HC Pore Vol"], zones,
            [1.35, 0.72, 0.72, 0.68, 0.78, 0.5, 0.68, 0.6, 0.6, 0.85])),
          ("Detailed Interval Analysis",
           (["Depth (ft)", "GR (API)", "RHOB (g/cc)", "NPHI (%)",
             "RT (ohm·m)", "PHIE (%)", "SW (%)", "Vcl (%)", "Lithology"],
            detail, [0.8, 0.75, 0.9, 0.75, 0.85, 0.75, 0.65, 0.65, 1.35])),
          ("Cutoff Criteria Applied",
           (["Parameter", "Cutoff Value", "Basis"], cutoffs, [1.9, 1.2, 3.1]))],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — Petrophysical Report 2024")


def dir_survey(path, w, rng, shown):
    td = float(w.get("final_td") or 12000)
    rows, md, incl, azi, tvd, ns, ew = [], 0.0, 0.0, rng.uniform(0, 360), 0.0, 0.0, 0.0
    while md < td and len(rows) < 34:
        step = rng.uniform(90, 160)
        md += step
        incl = min(90.0, incl + rng.uniform(0, 4.2))
        azi = (azi + rng.uniform(-3, 3)) % 360
        tvd += step * math.cos(math.radians(incl))
        ns += step * math.sin(math.radians(incl)) * math.cos(math.radians(azi))
        ew += step * math.sin(math.radians(incl)) * math.sin(math.radians(azi))
        rows.append([f"{md:,.0f}", f"{incl:.2f}", f"{azi:.2f}", f"{tvd:,.0f}",
                     f"{ns:,.0f}", f"{ew:,.0f}", f"{math.hypot(ns, ew):,.0f}",
                     f"{rng.uniform(0,3.5):.2f}"])
    _pdf(path, "DIRECTIONAL SURVEY REPORT",
         f"{rng.choice(SERVICE_COS)} Drilling Services",
         _ident(w, shown, rng, [
             ("Survey Type", rng.choice(["GYRO", "MWD", "MAGNETIC"])),
             ("Calculation", "Minimum curvature"), ("Depth Reference", "KB"),
             ("Survey Date", w.get("spud_date", ""))]),
         [("Survey Stations",
           (["MD (ft)", "Incl (°)", "Azim (°)", "TVD (ft)", "N/S (ft)",
             "E/W (ft)", "Closure (ft)", "DLS (°/100ft)"], rows,
            [0.85, 0.8, 0.85, 0.85, 0.85, 0.85, 0.95, 1.1]))],
         f"CONFIDENTIAL — {w.get('operator_name','')} — "
         f"{w.get('well_name','')} — Directional Survey 2024")


def image_only_pdf(path, w, rng):
    """A scan with no text layer — pdftotext returns nothing. The catalog should
    record 'no text extracted', not 'no well found'. Different failures needing
    different handling, and easy to conflate."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (1275, 1650), "white")
    d = ImageDraw.Draw(img)
    d.text((90, 80), "WELL SCOUT TICKET   (scanned)", fill="black")
    d.text((90, 130), f"API {_dashed(w['uwi'])}", fill="black")
    d.text((90, 175), w.get("well_name", ""), fill="black")
    d.text((90, 220), w.get("operator_name", ""), fill="black")
    for i in range(22):
        d.line((90, 280 + i*40, 1180, 280 + i*40), fill=(205, 205, 205))
    c = canvas.Canvas(path, pagesize=letter)
    c.drawImage(ImageReader(img), 0, 0, width=612, height=792)
    c.save()


# --------------------------------------------------------------------------- #
# LAS files
# --------------------------------------------------------------------------- #
# Bed types and the log response each produces. Values are centres; the curve
# generator walks toward them rather than sampling them independently.
BEDS = {
    #            GR    RHOB  NPHI   RT    DT    PEF   SP
    "shale":    (105,  2.52, 0.34,   4,  105,  3.2, -12),
    "sand":     ( 42,  2.31, 0.22,  28,   82,  2.0, -68),
    "limestone":( 22,  2.68, 0.06, 180,   50,  5.1,  -6),
    "dolomite": ( 28,  2.80, 0.03, 240,   45,  3.1,  -8),
    "anhydrite":( 14,  2.95, 0.01, 900,   50,  5.0,   0),
    "coal":     ( 38,  1.45, 0.48,  90,  130,  0.2, -30),
}
BED_SEQ = ["shale", "sand", "shale", "limestone", "shale", "dolomite",
           "sand", "shale", "limestone", "shale"]


def _walk(prev, target, step_ft, rng, sd, tightness=0.06):
    """One sample of a depth-correlated curve.

    A log is not white noise: adjacent samples half a foot apart are strongly
    correlated, because the tool has a vertical resolution of a foot or more and
    the rock itself changes gradually within a bed. Sampling each depth
    independently — which the example LAS did — produces something that passes a
    format check and fails any statistic computed over it, and looks obviously
    wrong to anyone who reads logs.

    So each sample pulls toward the bed's characteristic value and adds a small
    excursion, giving autocorrelation within beds and a step at each boundary.
    """
    pull = (target - prev) * tightness * max(1.0, step_ft / 0.5)
    return prev + pull + rng.gauss(0, sd)


# THE FOUR EXCEPTIONS. LAS 1.2 orders the ~W section descr:value for every
# mnemonic EXCEPT these, which keep the 2.0 order value:descr. That is not
# a lasio quirk - it is the 1.2 spec, and lasio encodes it in
# defaults.ORDER_DEFINITIONS[1.2]["Well"]. Swapping these four as well
# hands a numeric column the string "STOP DEPTH".
_LAS12_VALUE_FIRST = {"STRT", "STOP", "STEP", "NULL"}


def _hline(mnem, unit, value, descr, version="2.0"):
    """One ~WELL / ~PARAMETER line, in the field order the declared VERSION
    actually specifies.

        LAS 2.0    MNEM.UNIT   VALUE : DESCRIPTION
        LAS 1.2    MNEM.UNIT   DESCRIPTION : VALUE

    THE 1.2 CASE USED TO BE A LIE. Every line was written 2.0-ordered and only
    the VERS number changed, so the "las_1.2" variant was not a LAS 1.2 file at
    all. lasio read it exactly as the standard says to and handed back every
    field as its own label:

        UWI    value='UNIQUE WELL IDENTIFIER'
        LOG_ID value='LOG ID'

    Four files in the corpus, and all four stuck: their curves staged under
    log_id 'LOG ID' and their log header never staged, so promote held them on
    a missing parent forever. The generator's own comment says this case exists
    because "a parser that only ever meets unwrapped LAS 2.0 with a populated
    UWI is not a tested parser" — true, and it was testing nothing, because the
    fixture was malformed rather than merely old.

    The order above was established by writing both bodies under both versions
    and asking lasio which came back right, not from memory. The matrix is
    symmetric: each order reads correctly under its own version and swapped
    under the other.

    AND THEN IT WAS WRONG AGAIN, because the sample was wrong. That check
    used UWI and LOG_ID, which really are descr:value, so a blanket swap
    passed it. STRT/STOP/STEP/NULL are value:descr even in 1.2 - see
    _LAS12_VALUE_FIRST above. Swapping them too wrote

        STOP.FT   STOP DEPTH : 5165.0

    so lasio returned value='STOP DEPTH', TOTAL_DEPTH got a string, the
    FILE_WELL_HEADER MERGE failed nvarchar -> numeric, and because a failed
    write leaves the file pending, the extract stage re-claimed the same
    seven files ~570 times before the run was killed.
    """
    swap = (str(version).startswith("1.2")
            and str(mnem).strip().upper() not in _LAS12_VALUE_FIRST)
    left, right = (descr, value) if swap else (value, descr)
    return f" {mnem:<4}.{unit:<17}{str(left):<30}: {right}"


# --------------------------------------------------------------------------- #
# LAS 3.0 enrichment: a wider curve suite, and the sections that make a file 3.0
# --------------------------------------------------------------------------- #
#
# WHAT WAS THIN. The corpus had exactly one 3.0 file, carrying a single ~Log
# section -- so the very thing 3.0 exists for, and the thing that makes lasio
# fail on it, was never exercised by a generated file. split_las3 and
# _view_las3_sections were both written against the spec's published samples
# and one hand-made fixture.
#
# The extra sections are not decoration. las3_capture already maps
# ~Inclinometry into cat_well_dir_srvy_hdr/sta, and its own docstring says Core,
# Tops and Test "have mirrors waiting and are the obvious next additions, but
# each needs its own column mapping read off the real files rather than
# assumed". These files are that input.

# The extended suite. Every one of these is DERIVED from the bed model that
# already drives GR/RHOB/NPHI/RT/DT/PEF/SP/CALI, so the curves agree with each
# other and with the lithology. A suite of independently random tracks would
# pass a format check and look wrong to anyone who reads logs -- the same
# reasoning _walk was written for.
_RICH_CURVES = [
    ("BS",   "IN",   "Bit Size"),
    ("DRHO", "G/CC", "Density Correction"),
    ("DTS",  "US/F", "Shear Sonic Delta-T"),
    ("RESD", "OHMM", "Deep Resistivity"),
    ("RESM", "OHMM", "Medium Resistivity"),
    ("RESS", "OHMM", "Shallow Resistivity"),
    ("SGR",  "GAPI", "Total Gamma Ray"),
    ("CGR",  "GAPI", "Computed Gamma Ray (Th+K)"),
    ("TENS", "LBF",  "Cable Tension"),
]

# Vp/Vs by lithology. Shear is slower than compressional by a ratio the rock
# type sets -- a constant multiplier would make DTS a redrawn DT.
_VPVS = {"shale": 1.92, "sand": 1.62, "limestone": 1.87,
         "dolomite": 1.80, "anhydrite": 1.85, "coal": 2.10}
# Beds that take mud filtrate. Invasion is what separates the three resistivity
# curves; in a tight or shaly bed they read almost the same, and drawing a big
# separation there would be a picture of something that does not happen.
_PERMEABLE = ("sand", "limestone", "dolomite")


def _rich_values(cur, lith, depth, td, rng, bit=8.5):
    """The extended curves for one depth, in _RICH_CURVES order."""
    washout = cur["CALI"] - bit
    drho = max(-0.04, min(0.30, washout * 0.085 + rng.gauss(0, 0.006)))
    dts = cur["DT"] * _VPVS.get(lith, 1.85) + rng.gauss(0, 1.2)
    rt = cur["RT"]
    if lith in _PERMEABLE:
        f_m = rng.uniform(0.62, 0.80)
        f_s = f_m * rng.uniform(0.48, 0.72)
    else:
        f_m = rng.uniform(0.92, 1.00)
        f_s = f_m * rng.uniform(0.93, 0.995)
    resm = max(0.2, rt * f_m)
    ress = max(0.2, rt * f_s)
    sgr = cur["GR"]
    # CGR is SGR less the uranium contribution, so it is always the lower of
    # the two -- clamped rather than merely scaled, because noise must not put
    # the computed curve above the total.
    cgr = min(sgr, max(0.0, sgr * 0.74 + rng.gauss(0, 1.5)))
    tens = max(300.0, 2600.0 - (td - depth) * 0.018 + rng.gauss(0, 12))
    return [bit, drho, dts, rt, resm, ress, sgr, cgr, tens]


def _q(s):
    """A 3.0 cell, quoted when it carries the delimiter.

    A core description is exactly where a comma lands inside a value
    ("shale, silty"), and an unquoted one shifts every column after it -- the
    kind of wrong that lands in a table looking plausible. split_las3 honours
    the quotes; this is what produces them.
    """
    t = str(s)
    return '"' + t + '"' if ("," in t or '"' in t) else t


def _defsec(name, cols):
    """A ~*_Definition block. cols is [(mnem, unit, descr, fmt)]."""
    out = [f"~{name}_Definition"]
    for mnem, unit, descr, fmt in cols:
        out.append(f" {mnem:<10}.{unit:<8} : {descr} {{{fmt}}}")
    return out


def _datasec(name, cols, rows):
    """A data block that NAMES ITS DEFINITION.

    That association is the rule split_las3 keys on -- "a data section is one
    that names a definition" -- and it is how the reader knows which columns
    belong to these rows without guessing from the name.
    """
    out = ["", f"~{name} | {name.split('[')[0]}_Definition"]
    for r in rows:
        out.append(",".join(_q(v) for v in r))
    return out


_CORE_LITH = ["sandstone", "shaly sandstone", "limestone", "dolomite",
              "shale", "silty shale"]
_CORE_NOTE = ["fine grained, well sorted, oil stained",
              "medium grained, calcite cement, fair porosity",
              "argillaceous, thin bedded, poor recovery",
              "vuggy, partly dolomitised, good shows",
              "laminated, carbonaceous, no shows",
              "burrowed, glauconitic, trace fluorescence"]
_TEST_TYPE = ["DST", "DST", "RFT", "PRODUCTION"]
_RECOVERY = ["gas to surface, no water",
              "oil and gas cut mud, rising pressure",
              "salt water, no shows",
              "gas cut mud, slow build-up",
              "oil, 32 API, trace water"]


def las3_extra_sections(w, rng, start, stop, td, units="FT", null_val=-999.25):
    """The Core / Tops / Inclinometry / Test blocks of a rich LAS 3.0 file.

    Returns a list of lines. Every section is a Definition/Data PAIR joined by
    an association, which is what makes it a data set rather than a header
    block as far as split_las3 is concerned.

    The mnemonics are not free choices. las3_capture looks for MD/DEPT/DEPTH,
    TVD, AZIM/AZI/AZ and DEVI/INCL/INC/DEV, so ~Inclinometry uses names from
    those lists -- a survey written with mnemonics outside them parses into a
    data set that the capture stage then silently declines to map, which is
    exactly the failure this corpus is meant to catch rather than cause.
    """
    L = []
    ft = 1.0 if units == "FT" else 0.3048

    # ---- Core: a few cut intervals inside the logged section ---------------
    core_cols = [("CORE_TOP", units, "Core interval top", "F"),
                 ("CORE_BOT", units, "Core interval base", "F"),
                 ("RECOV", "%", "Recovery", "F"),
                 ("POR", "%", "Core porosity", "F"),
                 ("PERM", "MD", "Core permeability", "F"),
                 ("SO", "%", "Oil saturation", "F"),
                 ("SW", "%", "Water saturation", "F"),
                 ("LITH", "", "Core lithology", "S"),
                 ("DESC", "", "Core description", "S")]
    core_rows = []
    for _i in range(rng.randint(2, 4)):
        top = round(rng.uniform(start + 200 * ft, stop - 120 * ft), 1)
        bot = round(top + rng.uniform(18, 60) * ft, 1)
        por = round(rng.uniform(3.5, 24.0), 1)
        core_rows.append([
            top, bot, round(rng.uniform(72, 100), 1), por,
            # Permeability tracks porosity the way a Kozeny-style trend does;
            # an independent random k next to phi is the giveaway.
            round(max(0.01, 10 ** (0.22 * por - 3.1) * rng.uniform(0.5, 2.0)), 3),
            round(rng.uniform(0, 42), 1), round(rng.uniform(18, 78), 1),
            rng.choice(_CORE_LITH), rng.choice(_CORE_NOTE)])
    core_rows.sort(key=lambda r: r[0])
    L += [""] + _defsec("Core", core_cols) + _datasec("Core[1]", core_cols,
                                                      core_rows)

    # ---- Tops: the state's stratigraphic column, down to TD ----------------
    tops_cols = [("TOPMD", units, "Formation top, measured depth", "F"),
                 ("TOPTVD", units, "Formation top, true vertical depth", "F"),
                 ("FORMATION", "", "Formation name", "S"),
                 ("SOURCE", "", "Pick source", "S"),
                 ("REMARK", "", "Remark", "S")]
    tops_rows = []
    for nm, dep, note in _strat(w):
        md = dep * ft
        if md < start or md > min(stop, td):
            continue
        tops_rows.append([round(md, 1), round(md * rng.uniform(0.93, 1.0), 1),
                          nm, rng.choice(["LOG", "LOG", "MUDLOG", "SEISMIC"]),
                          note or ""])
    if tops_rows:
        L += [""] + _defsec("Tops", tops_cols) + _datasec("Tops", tops_cols,
                                                          tops_rows)

    # ---- Inclinometry: a build-and-hold survey ------------------------------
    # MINIMUM CURVATURE, not a straight cosine. The difference is small per
    # station and accumulates over a hole, and this file is fixture data for a
    # directional-survey loader -- feeding it a survey whose TVD does not close
    # would teach the loader the wrong answer.
    incl_cols = [("MD", units, "Measured depth", "F"),
                 ("INCL", "DEG", "Inclination", "F"),
                 ("AZIM", "DEG", "Azimuth", "F"),
                 ("TVD", units, "True vertical depth", "F"),
                 ("NS", units, "North-South offset", "F"),
                 ("EW", units, "East-West offset", "F"),
                 ("DLS", "DEG", "Dogleg severity per 100", "F")]
    import math as _m
    kop = td * rng.uniform(0.35, 0.55)
    build = rng.uniform(1.4, 3.0)                 # degrees per 100 ft
    hold = rng.uniform(18, 62)                    # final inclination
    azi0 = rng.uniform(0, 360)
    step_md = 100.0 * ft
    md, inc, azi, tvd, ns, ew = 0.0, 0.0, azi0, 0.0, 0.0, 0.0
    incl_rows = [[0.0, 0.0, round(azi0, 2), 0.0, 0.0, 0.0, 0.0]]
    while md < td:
        nmd = min(md + step_md, td)
        dmd = nmd - md
        ninc = inc if nmd <= kop else min(hold, inc + build * (dmd / (100 * ft)))
        nazi = azi + rng.gauss(0, 0.6)
        i1, i2 = _m.radians(inc), _m.radians(ninc)
        a1, a2 = _m.radians(azi), _m.radians(nazi)
        cosdl = (_m.cos(i2 - i1)
                 - _m.sin(i1) * _m.sin(i2) * (1 - _m.cos(a2 - a1)))
        cosdl = max(-1.0, min(1.0, cosdl))
        dl = _m.acos(cosdl)
        rf = 1.0 if dl < 1e-9 else (2.0 / dl) * _m.tan(dl / 2.0)
        tvd += (dmd / 2.0) * (_m.cos(i1) + _m.cos(i2)) * rf
        ns += (dmd / 2.0) * (_m.sin(i1) * _m.cos(a1)
                             + _m.sin(i2) * _m.cos(a2)) * rf
        ew += (dmd / 2.0) * (_m.sin(i1) * _m.sin(a1)
                             + _m.sin(i2) * _m.sin(a2)) * rf
        md, inc, azi = nmd, ninc, nazi
        incl_rows.append([round(md, 1), round(inc, 2), round(azi % 360, 2),
                          round(tvd, 1), round(ns, 1), round(ew, 1),
                          round(_m.degrees(dl) / (dmd / (100 * ft)), 2)])
    L += [""] + _defsec("Inclinometry", incl_cols) + _datasec(
        "Inclinometry", incl_cols, incl_rows)

    # ---- Test: drill stem / formation tests ---------------------------------
    test_cols = [("TESTNO", "", "Test number", "F"),
                 ("TESTTOP", units, "Test interval top", "F"),
                 ("TESTBOT", units, "Test interval base", "F"),
                 ("TESTTYPE", "", "Test type", "S"),
                 ("ISIP", "PSI", "Initial shut-in pressure", "F"),
                 ("FSIP", "PSI", "Final shut-in pressure", "F"),
                 ("DURATION", "MIN", "Test duration", "F"),
                 ("RECOVERY", "", "Recovery description", "S")]
    test_rows = []
    for i in range(rng.randint(1, 3)):
        top = round(rng.uniform(start + 300 * ft, stop - 200 * ft), 1)
        bot = round(top + rng.uniform(30, 140) * ft, 1)
        isip = round(rng.uniform(900, 4200), 1)
        test_rows.append([
            float(i + 1), top, bot, rng.choice(_TEST_TYPE), isip,
            # A final shut-in below the initial is the normal case; a test that
            # built back above its ISIP would be the anomaly, not the fixture.
            round(isip * rng.uniform(0.72, 0.99), 1),
            float(rng.choice([30, 45, 60, 90, 120])),
            rng.choice(_RECOVERY)])
    test_rows.sort(key=lambda r: r[1])
    L += [""] + _defsec("Test", test_cols) + _datasec("Test", test_cols,
                                                      test_rows)
    return L


def las_file(path, w, rng, shown="full", wrap=False, version="2.0",
             units="FT", null_val=-999.25, rich=False):
    """A LAS file with lithology-driven, depth-correlated curves.

    `version` governs BOTH the VERS line and the ~WELL/~PARAMETER field order —
    see _hline. Writing one without the other produces a file no conforming
    reader can parse.
    """
    td = float(w.get("final_td") or 9000)
    start = round(td * rng.uniform(0.12, 0.25), 1)
    stop = round(td * rng.uniform(0.90, 0.99), 1)
    step = 0.5 if units == "FT" else 0.1524          # 0.5 ft, or its metric twin
    if units == "M":
        start, stop = round(start * 0.3048, 2), round(stop * 0.3048, 2)

    uwi_txt = {"full": w["uwi"], "dashed": _long_uwi(w["uwi"]), "none": ""}[shown]
    curves = [("DEPT", units, "DEPTH"), ("GR", "GAPI", "Gamma Ray"),
              ("CALI", "IN", "Caliper"), ("SP", "MV", "Spontaneous Potential"),
              ("RHOB", "G/CC", "Bulk Density"), ("NPHI", "V/V", "Neutron Porosity"),
              ("RT", "OHMM", "True Resistivity"), ("DT", "US/F", "Sonic Delta-T"),
              ("PEF", "B/E", "Photoelectric Factor")]
    _rich = bool(rich) and str(version).startswith("3")
    if _rich:
        curves = curves + list(_RICH_CURVES)

    # Bed boundaries down the logged interval — thick enough to be readable.
    n = int((stop - start) / step) + 1
    bounds, d = [], start
    seq = BED_SEQ * 6
    i = 0
    while d < stop:
        thick = rng.uniform(40, 260) * (1 if units == "FT" else 0.3048)
        bounds.append((d, min(d + thick, stop), seq[i % len(seq)]))
        d += thick
        i += 1

    _v3 = str(version).startswith("3")
    hdr = [
        "~VERSION INFORMATION",
        f" VERS.                 {version}   : CWLS Log ASCII Standard - VERSION {version}",
        f" WRAP.                 {'YES' if wrap else 'NO'}    : "
        f"{'MULTIPLE LINES PER DEPTH STEP' if wrap else 'ONE LINE PER DEPTH STEP'}",
    ] + ([
        # 3.0 declares its delimiter. lasio ignores DLM, which is why
        # las_reader honours it before handing the data over — a 3.0 file with
        # DLM. COMMA read as 8 rows of nan instead of 4 rows of data.
        " DLM .                 COMMA : DELIMITING CHARACTER",
    ] if _v3 else []) + [
        "",
        "~WELL INFORMATION",
        _hline("UWI", "", uwi_txt, "UNIQUE WELL IDENTIFIER", version),
        _hline("WELL", "", w.get("well_name", ""), "WELL NAME", version),
        _hline("COMP", "", w.get("operator_name", ""), "COMPANY", version),
        _hline("FLD", "", _field(w, rng), "FIELD", version),
        _hline("SRVC", "", rng.choice(SERVICE_COS).upper(), "SERVICE COMPANY",
               version),
        _hline("DATE", "", w.get("spud_date", ""), "LOG DATE", version),
        _hline("STRT", units, start, "START DEPTH", version),
        _hline("STOP", units, stop, "STOP DEPTH", version),
        _hline("STEP", units, step, "STEP", version),
        _hline("NULL", "", null_val, "NULL VALUE", version),
        _hline("CNTY", "", w.get("county", ""), "COUNTY", version),
        _hline("STAT", "", _abbr(w), "STATE", version),
        _hline("CTRY", "", "US", "COUNTRY", version),
        _hline("API", "", _dashed(w["uwi"]) if shown != "none" else "",
               "API NUMBER", version),
        _hline("LOG_ID", "", f"LOG_{w['uwi']}_1", "LOG ID", version),
        "",
        "~PARAMETER INFORMATION",
        _hline("RUN", "", rng.randint(1, 3), "RUN NUMBER", version),
        _hline("EKB", units, w.get("kb_elevation", ""),
               "KELLY BUSHING ELEVATION", version),
        _hline("EGL", units, w.get("ground_elevation", ""),
               "GROUND LEVEL ELEVATION", version),
        _hline("TDD", units, td, "DRILLER TOTAL DEPTH", version),
        _hline("BHT", "DEGF", rng.randint(140, 240),
               "BOTTOM HOLE TEMPERATURE", version),
        _hline("MUD", "", rng.choice(["WBM", "OBM", "SBM"]), "MUD TYPE",
               version),
        _hline("MDEN", "G/CC", f"{rng.uniform(1.05, 1.35):.2f}", "MUD DENSITY",
               version),
        _hline("MATR", "", rng.choice(["SAND", "LIME", "DOLO"]),
               "MATRIX FOR NEUTRON", version),
        "",
        "~CURVE INFORMATION",
    ]
    for m, u, d_ in curves:
        hdr.append(f" {m:<5}.{u:<16}: {d_}")
    # LAS 3.0 names the data section ~Ascii and may delimit it with something
    # other than whitespace. Comma is the common choice and the one that was
    # silently misread until las_reader honoured DLM, so the fixture uses it
    # deliberately — a 3.0 variant written with spaces would pass without
    # exercising the thing that was broken.
    if str(version).startswith("3"):
        # ~Ascii is the section header and NOTHING ELSE. LAS 2.0 puts the
        # curve mnemonics on the ~A line itself; 3.0 does not, and emitting
        # them anyway makes the first data row a row of TEXT — which turns the
        # whole array to <U32 and silently gives 56,826 string rows instead of
        # 6,313 numeric ones. Caught by the round-trip check, not by reading.
        hdr += ["", "~Ascii"]
    else:
        hdr += ["", "~A  " + "  ".join(m for m, _u, _d in curves)]

    # Curve state, seeded at the first bed's characteristics
    b0 = BEDS[bounds[0][2]]
    cur = {"GR": b0[0], "RHOB": b0[1], "NPHI": b0[2], "RT": b0[3],
           "DT": b0[4], "PEF": b0[5], "SP": b0[6], "CALI": 8.5}
    sd = {"GR": 4.0, "RHOB": 0.012, "NPHI": 0.006, "RT": 0.9, "DT": 1.2,
          "PEF": 0.06, "SP": 1.5, "CALI": 0.05}

    lines, d, bi = [], start, 0
    hole = 8.5
    while d <= stop + 1e-6:
        while bi + 1 < len(bounds) and d > bounds[bi][1]:
            bi += 1
        t = BEDS[bounds[bi][2]]
        tgt = {"GR": t[0], "RHOB": t[1], "NPHI": t[2], "RT": t[3],
               "DT": t[4], "PEF": t[5], "SP": t[6],
               # washout in shale, on gauge elsewhere — a caliper that never
               # moves is a giveaway that the file is synthetic
               "CALI": hole + (1.4 if bounds[bi][2] == "shale" else 0.05)}
        for k in cur:
            cur[k] = _walk(cur[k], tgt[k], step, rng, sd[k])
        cur["NPHI"] = max(0.005, cur["NPHI"])
        cur["RHOB"] = max(1.4, cur["RHOB"])
        cur["RT"] = max(0.2, cur["RT"])
        vals = [d, cur["GR"], cur["CALI"], cur["SP"], cur["RHOB"],
                cur["NPHI"], cur["RT"], cur["DT"], cur["PEF"]]
        if _rich:
            vals += _rich_values(cur, bounds[bi][2], d, td, rng)
        # A real log has gaps — tool off-bottom, bad hole, a curve not recorded
        if rng.random() < 0.004:
            k = rng.randint(1, len(vals) - 1)
            vals[k] = null_val
        if _v3:
            # Comma-delimited, matching the DLM this file declares. Written
            # unpadded because a delimited format does not use column
            # alignment — and because a reader that only ever meets padded
            # columns will pass on a file that happens to line up.
            lines.append(",".join(f"{v:.4f}" for v in vals))
        elif wrap:
            lines.append(f"{vals[0]:>10.4f}")
            lines.append("   " + "".join(f"{v:>12.4f}" for v in vals[1:]))
        else:
            lines.append("".join(f"{v:>12.4f}" for v in vals))
        d += step

    extra = (las3_extra_sections(w, rng, start, stop, td, units, null_val)
             if _rich else [])
    with open(path, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(hdr) + "\n" + "\n".join(lines) + "\n")
        if extra:
            f.write("\n".join(extra) + "\n")
    return len(lines)


# --------------------------------------------------------------------------- #
# Office documents
# --------------------------------------------------------------------------- #
def completion_docx(path, w, rng, shown, extra_uwi=None):
    from docx import Document
    doc = Document()
    doc.add_heading("WELL COMPLETION REPORT", 0)
    doc.add_paragraph("CONFIDENTIAL — For Operator Use Only")
    td = float(w.get("final_td") or 12000)
    tvd = round(td * rng.uniform(0.45, 0.62))
    zone = _strat(w)[-4][0]

    doc.add_heading("1. Well Identification", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Field", "Value"
    for k, v in (("Well Name", w.get("well_name", "")),
                 ("UWI", _long_uwi(w["uwi"]) if shown != "none" else ""),
                 ("API Number", _dashed(w["uwi"]) if shown != "none" else ""),
                 ("Operator", w.get("operator_name", "")),
                 ("Field", _field(w, rng)), ("County", w.get("county", "")),
                 ("State", _abbr(w)), ("Spud Date", w.get("spud_date", "")),
                 ("Completion Date", w.get("completion_date", "")),
                 ("Total Depth (ft MD)", f"{td:,.0f}"), ("TVD (ft)", f"{tvd:,}"),
                 ("KB Elevation (ft)", w.get("kb_elevation", ""))):
        r = t.add_row().cells
        r[0].text, r[1].text = k, str(v)

    doc.add_heading("2. Completion Summary", level=1)
    doc.add_paragraph(
        f"The {w.get('well_name','')} was completed as a horizontal producer in "
        f"the {zone} formation. The lateral was landed at {tvd:,} ft TVD and "
        f"drilled to {td:,.0f} ft MD. Production casing was set at total depth "
        f"and cemented. The well was stimulated in {rng.randint(20,44)} stages "
        f"using slickwater and hybrid fluid systems.")
    if extra_uwi:
        doc.add_paragraph(
            f"Offset well {_dashed(extra_uwi)} was used for correlation of the "
            f"{zone} landing zone and stage placement.")

    doc.add_heading("3. Perforation Intervals", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("Stage", "Top Depth (ft MD)", "Base Depth (ft MD)",
                           "Length (ft)", "Formation")):
        t.rows[0].cells[i].text = h
    top = tvd + rng.uniform(50, 250)
    for st in range(1, rng.randint(6, 12)):
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text = str(st), f"{top:,.0f}", f"{top+50:,.0f}"
        r[3].text, r[4].text = "50", zone
        top += rng.uniform(230, 280)

    doc.add_heading("4. Stimulation Parameters", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Parameter", "Value"
    for k, v in (("Total Fluid (bbl)", f"{rng.randint(70,160)*1000:,}"),
                 ("Total Proppant (lbs)", f"{rng.randint(10,24)*1000000:,}"),
                 ("Max Treating Pressure (psi)", f"{rng.randint(7800,8900):,}"),
                 ("Average Rate (bpm)", rng.randint(55, 70)),
                 ("Fluid System", "Slickwater / hybrid"),
                 ("Proppant Type", "100 mesh + 40/70 white sand")):
        r = t.add_row().cells
        r[0].text, r[1].text = k, str(v)

    doc.add_heading("5. Initial Production", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("Date", "Oil (bbl/d)", "Gas (Mcf/d)",
                           "Water (bbl/d)", "Choke")):
        t.rows[0].cells[i].text = h
    q = rng.uniform(800, 2000)
    try:
        d0 = date.fromisoformat(str(w.get("completion_date"))[:10])
    except Exception:
        d0 = date(2024, 6, 1)
    for m in range(4):
        r = t.add_row().cells
        r[0].text = (d0 + timedelta(days=30*m)).isoformat()
        r[1].text, r[2].text = f"{q:,.0f}", f"{q*rng.uniform(1.4,1.9):,.0f}"
        r[3].text = f"{q*rng.uniform(0.2,0.6):,.0f}"
        r[4].text = f"{rng.choice([32,40,48])}/64"
        q *= rng.uniform(0.78, 0.9)
    doc.save(path)


def tops_docx(path, wells, rng):
    """A multi-well report. Formation tops studies cover a project, not a single
    well — so this is the many-UWIs-in-one-document case, which is realistic
    rather than contrived."""
    from docx import Document
    doc = Document()
    doc.add_heading("FORMATION TOPS REPORT", 0)
    doc.add_paragraph(f"{_field(wells[0], rng)} — Sub-Basin Study")
    doc.add_paragraph("Prepared by: Data Wrangler Solutions  |  Date: "
                      f"{date(2024, rng.randint(1,12), 15).strftime('%B %Y')}")
    doc.add_heading("1. Geological Overview", level=1)
    doc.add_paragraph(
        f"The following formation tops were picked from petrophysical log "
        f"analysis of {len(wells)} horizontal wells. Picks were correlated "
        f"against regional type logs and calibrated to available core.")
    doc.add_heading("2. Formation Tops by Well", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("UWI", "Formation", "Top MD (ft)", "Base MD (ft)",
                           "Net Pay (ft)", "Fluid")):
        t.rows[0].cells[i].text = h
    for w in wells:
        td = float(w.get("final_td") or 12000)
        for name, nominal, _n in _strat(w)[-4:]:
            if nominal > td:
                continue
            top = nominal * rng.uniform(0.92, 1.08)
            r = t.add_row().cells
            r[0].text, r[1].text = _dashed(w["uwi"]), name
            r[2].text = f"{top:,.0f}"
            r[3].text = f"{top+rng.uniform(120,320):,.0f}"
            r[4].text, r[5].text = str(rng.randint(18, 62)), \
                rng.choice(["OIL", "OIL/GAS", "GAS"])
    doc.add_heading("3. Fluid Contacts", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("UWI", "Formation", "Contact Type", "Depth MD (ft)",
                           "Method")):
        t.rows[0].cells[i].text = h
    for w in wells[:4]:
        r = t.add_row().cells
        r[0].text, r[1].text = _dashed(w["uwi"]), _strat(w)[-4][0]
        r[2].text = rng.choice(["OWC", "GWC", "GOC"])
        r[3].text = f"{float(w.get('final_td') or 9000)*rng.uniform(0.5,0.7):,.0f}"
        r[4].text = rng.choice(["Log analysis", "Pressure data", "Test data"])
    doc.save(path)


def production_xlsx(path, w, rng):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Monthly Production"
    ws.append(["UWI", "WELL_NAME", "OPERATOR", "DATE", "OIL_BBL", "GAS_MCF",
               "WATER_BBL", "BOE", "TUBING_PRESS_PSI", "STATUS"])
    try:
        d0 = date.fromisoformat(str(w.get("completion_date"))[:10])
    except Exception:
        d0 = date(2022, 1, 1)
    q = rng.uniform(200, 900)
    n = rng.randint(18, 48)
    for m in range(n):
        d = d0 + timedelta(days=30*m)
        oil = q * math.exp(-0.03*m) * rng.uniform(0.85, 1.15)
        gas = oil * rng.uniform(1.2, 2.6)
        ws.append([_long_uwi(w["uwi"]), w.get("well_name", ""),
                   w.get("operator_name", ""), d.isoformat(), round(oil),
                   round(gas), round(oil*rng.uniform(0.4, 2.4)),
                   round(oil + gas/6), rng.randint(900, 2600),
                   "PRODUCING" if m < n-3 else "SHUT-IN"])
    ws2 = wb.create_sheet("Well Header")
    for k, v in (("UWI", _long_uwi(w["uwi"])), ("API", _dashed(w["uwi"])),
                 ("Well Name", w.get("well_name", "")),
                 ("Operator", w.get("operator_name", "")),
                 ("County", w.get("county", "")), ("State", _abbr(w)),
                 ("Total Depth", w.get("final_td", "")),
                 ("Spud Date", w.get("spud_date", ""))):
        ws2.append([k, str(v)])
    wb.save(path)


def core_xlsx(path, w, rng):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Core Analysis"
    ws.append(["CORE ANALYSIS REPORT"])
    ws.append(["UWI", _long_uwi(w["uwi"])])
    ws.append(["API Number", _dashed(w["uwi"])])
    ws.append(["Well Name", w.get("well_name", "")])
    ws.append(["Operator", w.get("operator_name", "")])
    ws.append(["Laboratory", rng.choice(["Core Lab", "Weatherford Labs",
                                         "Premier Oilfield Group"])])
    ws.append([])
    ws.append(["Sample", "Depth (ft)", "Porosity (%)", "Perm kair (md)",
               "Perm Klink (md)", "Grain Density (g/cc)", "So (%)", "Sw (%)",
               "Lithology", "Description"])
    base = float(w.get("final_td") or 9000) * rng.uniform(0.55, 0.75)
    for i in range(1, rng.randint(20, 60)):
        ka = rng.uniform(0.005, 480)
        ws.append([f"S{i:03d}", round(base + i*0.5, 1),
                   round(rng.uniform(2.5, 22), 2), round(ka, 3),
                   round(ka*rng.uniform(0.6, 0.95), 3),
                   round(rng.uniform(2.62, 2.74), 3),
                   round(rng.uniform(0, 45), 1), round(rng.uniform(15, 70), 1),
                   rng.choice(["Limestone", "Dolomite", "Sandstone", "Shale"]),
                   rng.choice(["vuggy", "fractured", "tight", "porous",
                               "argillaceous", "bioturbated"])])
    wb.save(path)


# --------------------------------------------------------------------------- #
# Orchestration
# --------------------------------------------------------------------------- #
def load_wells(csv_path):
    with open(csv_path, newline="", encoding="utf-8-sig") as f:
        return [r for r in csv.DictReader(f) if (r.get("uwi") or "").strip()]


def generate(wells, out_dir, per_well, seed=42, log=print):
    rng = random.Random(seed)
    pdf_dir = os.path.join(out_dir, "sample_pdfs")
    off_dir = os.path.join(out_dir, "sample_office")
    las_dir = os.path.join(out_dir, "las_files")
    for d in (pdf_dir, off_dir, las_dir):
        os.makedirs(d, exist_ok=True)
        for f in os.listdir(d):          # a stale document is as misleading here
            try:                          # as a stale CSV is on the data side
                os.remove(os.path.join(d, f))
            except OSError:
                pass

    have = {}
    for mod, name in (("reportlab", "pdf"), ("openpyxl", "xlsx"),
                      ("docx", "docx")):
        try:
            __import__(mod)
            have[name] = True
        except ImportError:
            have[name] = False
            log(f"  ⚠ {mod} not installed — skipping {name} documents")

    manifest = []
    PDF_KINDS = [("SCOUT", scout_ticket), ("EOW", end_of_well),
                 ("CASING_CEMENT", casing_cement), ("WELL_TEST", well_test),
                 ("PETROPHYSICS", petrophysics), ("DIRSRVY", dir_survey)]

    for w in wells:
        uwi = w["uwi"]
        for _ in range(per_well):
            case, shown, fname_uwi, expect, ww = "clean", "full", uwi, uwi, w
            if rng.random() < HARD_CASE_SHARE:
                case = rng.choice(["filename_only", "text_only", "dashed_api",
                                   "unknown_uwi", "no_uwi", "image_only"])
                if case == "filename_only":
                    shown = "none"
                elif case == "text_only":
                    fname_uwi = f"DOC{rng.randint(10000, 99999)}"
                elif case == "dashed_api":
                    shown = "dashed"
                elif case == "unknown_uwi":
                    ghost = f"{uwi[:5]}{rng.randint(90000, 99999)}0000"
                    ww, fname_uwi, expect = dict(w, uwi=ghost), ghost, ""
                elif case == "no_uwi":
                    shown, expect = "none", ""
                    fname_uwi = f"REPORT{rng.randint(1000, 9999)}"
            shown_txt = {"full": ww["uwi"], "dashed": _dashed(ww["uwi"]),
                         "none": ""}[shown]
            try:
                if case == "image_only" and have["pdf"]:
                    p = os.path.join(pdf_dir, f"SCOUT_SCAN_{fname_uwi}.pdf")
                    image_only_pdf(p, ww, rng)
                    expect = ""
                else:
                    pick = rng.random()
                    if pick < 0.62 and have["pdf"]:
                        tag, fn = rng.choice(PDF_KINDS)
                        p = os.path.join(pdf_dir, f"{tag}_{fname_uwi}.pdf")
                        fn(p, ww, rng, shown_txt)
                    elif pick < 0.78 and have["xlsx"]:
                        p = os.path.join(off_dir, f"PRODUCTION_{fname_uwi}.xlsx")
                        production_xlsx(p, ww, rng)
                    elif pick < 0.88 and have["xlsx"]:
                        p = os.path.join(off_dir, f"CORE_ANALYSIS_{fname_uwi}.xlsx")
                        core_xlsx(p, ww, rng)
                    elif have["docx"]:
                        p = os.path.join(off_dir, f"COMPLETION_{fname_uwi}.docx")
                        other = (rng.choice(wells)["uwi"]
                                 if rng.random() < 0.25 else None)
                        completion_docx(p, ww, rng, shown, other)
                        if other:
                            case = f"two_uwis (also {other})"
                    else:
                        continue
            except Exception as e:
                log(f"  !! {case} for {uwi}: {type(e).__name__}: {e}")
                continue
            manifest.append({"file": os.path.relpath(p, out_dir),
                             "expected_uwi": expect,
                             "well_name": ww.get("well_name", ""), "case": case})

    # LAS — one or two per well, plus format variants. A parser that only ever
    # meets unwrapped LAS 2.0 with a populated UWI is not a tested parser.
    for w in wells:
        for run in range(1, rng.randint(2, 3)):
            case, shown = "clean", "full"
            wrap, ver, units, nullv = False, "2.0", "FT", -999.25
            r = rng.random()
            if r < 0.10:
                case, shown = "no_uwi", "none"
            elif r < 0.18:
                case, shown = "dashed_api", "dashed"
            elif r < 0.26:
                case, wrap = "wrapped", True
            elif r < 0.33:
                case, ver = "las_1.2", "1.2"
            elif r < 0.40:
                case, units = "metric", "M"
            elif r < 0.46:
                case, nullv = "null_9999", -9999.00
            elif r < 0.52:
                # 3.0: ~Ascii, comma-delimited, DLM declared. lasio parses the
                # header fine and ignores DLM, so this variant reads as nan
                # columns unless it goes through las_reader — which is the
                # point of having it in the corpus.
                case, ver = "las_3.0", "3.0"
            p = os.path.join(las_dir, f"{w['uwi']}_{run}.las")
            try:
                las_file(p, w, rng, shown, wrap, ver, units, nullv)
            except Exception as e:
                log(f"  !! LAS for {w['uwi']}: {type(e).__name__}: {e}")
                continue
            manifest.append({"file": os.path.relpath(p, out_dir),
                             "expected_uwi": "" if shown == "none" else w["uwi"],
                             "well_name": w.get("well_name", ""),
                             "case": f"las_{case}"})

    if have["docx"] and len(wells) >= 4:
        p = os.path.join(off_dir, "Formation_Tops_Study.docx")
        grp = rng.sample(wells, min(8, len(wells)))
        try:
            tops_docx(p, grp, rng)
            manifest.append({"file": os.path.relpath(p, out_dir),
                             "expected_uwi": "|".join(g["uwi"] for g in grp),
                             "well_name": "(multiple)", "case": "multi_well"})
        except Exception as e:
            log(f"  !! tops study: {type(e).__name__}: {e}")

    mpath = os.path.join(out_dir, "MANIFEST.csv")
    with open(mpath, "w", newline="", encoding="utf-8") as f:
        wtr = csv.DictWriter(f, fieldnames=["file", "expected_uwi",
                                            "well_name", "case"])
        wtr.writeheader()
        wtr.writerows(manifest)

    from collections import Counter
    for case, n in sorted(Counter(m["case"].split(" ")[0]
                                  for m in manifest).items()):
        log(f"   {case:16} {n:>5}")
    log(f"-- {len(manifest)} document(s); ground truth in {mpath}")
    return len(manifest)


def _main() -> int:
    ap = argparse.ArgumentParser(
        description="Generate synthetic well documents for File Catalog testing")
    ap.add_argument("--wells-csv", required=True)
    ap.add_argument("--out", default="synth_docs")
    ap.add_argument("--per-well", type=int, default=2)
    ap.add_argument("--limit", type=int)
    ap.add_argument("--seed", type=int, default=42)
    a = ap.parse_args()
    wells = load_wells(a.wells_csv)
    if a.limit:
        wells = wells[:a.limit]
    print(f"-- {len(wells)} well(s) x {a.per_well} -> {a.out}")
    generate(wells, a.out, a.per_well, a.seed)
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
